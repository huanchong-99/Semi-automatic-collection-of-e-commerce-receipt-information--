#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
收货信息自动采集工具
单文件设计架构，所有功能模块都集成在一个main.py文件中
"""

import os
import sys
import time
import json
import logging
import subprocess
import threading
import copy

# TCL兼容性初始化函数
def init_tcl_compatibility():
    """初始化TCL兼容性设置，解决版本冲突问题"""
    try:
        # 获取当前Python环境的TCL路径
        import sys
        python_dir = os.path.dirname(sys.executable)
        
        # 尝试多个可能的TCL路径
        possible_tcl_paths = [
            os.path.join(python_dir, 'tcl', 'tcl8.6'),
            os.path.join(python_dir, 'lib', 'tcl8.6'),
            r"C:\Users\Administrator\AppData\Local\Programs\Python\Python311\tcl\tcl8.6",
            # Conda环境路径
            os.path.join(python_dir, '..', 'lib', 'tcl8.6'),
            os.path.join(python_dir, '..', 'Library', 'lib', 'tcl8.6')
        ]
        
        possible_tk_paths = [
            os.path.join(python_dir, 'tcl', 'tk8.6'),
            os.path.join(python_dir, 'lib', 'tk8.6'),
            r"C:\Users\Administrator\AppData\Local\Programs\Python\Python311\tcl\tk8.6",
            # Conda环境路径
            os.path.join(python_dir, '..', 'lib', 'tk8.6'),
            os.path.join(python_dir, '..', 'Library', 'lib', 'tk8.6')
        ]
        
        # 查找存在的TCL路径
        tcl_path = None
        for path in possible_tcl_paths:
            if os.path.exists(path):
                tcl_path = path
                break
        
        tk_path = None
        for path in possible_tk_paths:
            if os.path.exists(path):
                tk_path = path
                break
        
        if tcl_path and tk_path:
            # 设置环境变量
            os.environ['TCL_LIBRARY'] = tcl_path
            os.environ['TK_LIBRARY'] = tk_path
            
            # 禁用精确版本检查，允许版本不匹配
            os.environ['TCL_EXACT_VERSION'] = '0'
            os.environ['TK_EXACT_VERSION'] = '0'
            
            print(f"TCL兼容性设置成功: TCL={tcl_path}, TK={tk_path}")
            return True
        else:
            print(f"未找到有效的TCL/TK路径")
            return False
            
    except Exception as e:
        print(f"TCL兼容性设置失败: {e}")
        return False

# 在导入tkinter前初始化TCL兼容性
init_tcl_compatibility()

import tkinter as tk
from tkinter import ttk, messagebox, filedialog, simpledialog
from tkinter.scrolledtext import ScrolledText
from datetime import datetime
import json
import ctypes
from ctypes import windll
import tempfile
import traceback
# 依赖检查和安装函数
def install_package(package):
    """安装单个包"""
    try:
        print(f"正在安装 {package}...")
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', package])
        print(f"✓ {package} 安装成功")
        return True
    except subprocess.CalledProcessError as e:
        print(f"✗ {package} 安装失败: {e}")
        return False

def auto_install_dependencies(missing_packages):
    """自动安装缺失的依赖包"""
    print("收货信息自动采集工具 - 依赖安装")
    print("=" * 50)
    print(f"检测到缺失的依赖包: {', '.join(missing_packages)}")
    print("正在自动安装...")
    print("-" * 30)
    
    # 安装每个依赖包
    success_count = 0
    failed_packages = []
    
    for package in missing_packages:
        if install_package(package):
            success_count += 1
        else:
            failed_packages.append(package)
        print()
    
    # 显示安装结果
    print("=" * 50)
    print(f"安装完成！成功: {success_count}/{len(missing_packages)}")
    
    if failed_packages:
        print(f"\n安装失败的包: {', '.join(failed_packages)}")
        print("请手动安装这些包或检查网络连接。")
        return False
    else:
        print("\n所有依赖包安装成功！程序将重新启动...")
        return True

def check_and_install_dependencies():
    """检查并安装缺失的依赖包，避免重复安装"""
    import importlib
    try:
        from importlib.metadata import distribution
    except ImportError:
        # Python < 3.8 fallback
        from importlib_metadata import distribution
    
    # 定义依赖包映射（模块名 -> 安装名）
    dependencies = {
        'websocket': 'websocket-client',
        'requests': 'requests',
        'cv2': 'opencv-python',
        'numpy': 'numpy',
        'pyautogui': 'pyautogui',
        'PIL': 'pillow',
        'win32gui': 'pywin32',
        'selenium': 'selenium',
        'pandas': 'pandas',
        'docx': 'python-docx',
        'pyperclip': 'pyperclip'
    }
    
    missing_packages = []
    
    # 检查每个依赖
    for module_name, install_name in dependencies.items():
        try:
            # 使用importlib.metadata检查包是否已安装
            distribution(install_name)
        except Exception:
            missing_packages.append(install_name)
    
    # 如果有缺失的包，提供自动安装选项
    if missing_packages:
        print(f"检测到缺失的依赖包: {', '.join(missing_packages)}")
        
        # 尝试自动安装
        if auto_install_dependencies(missing_packages):
            # 安装成功，重新启动程序
            print("\n正在重新启动程序...")
            try:
                import time
                time.sleep(2)
                os.execv(sys.executable, [sys.executable] + sys.argv)
            except Exception as e:
                print(f"重新启动失败: {e}")
                print("请手动重新运行程序。")
        
        return False
    
    return True

# 在程序启动时检查依赖
dependencies_ok = check_and_install_dependencies()

# 初始化全局变量
websocket = None
requests = None
cv2 = None
np = None
webdriver = None
pd = None
Document = None
pyautogui = None
pyperclip = None

# 只有在依赖检查通过后才导入模块
if dependencies_ok:
    # 网络相关导入
    try:
        import websocket
        import requests
    except ImportError:
        websocket = None
        requests = None

    # 验证码检测相关导入
    try:
        import cv2
        import numpy as np
        import pyautogui
        from PIL import Image, ImageTk
        import winsound
        import glob
        import win32gui
        import win32ui
        import win32con
        import win32api
    except ImportError:
        cv2 = None
        np = None

    # Selenium相关导入
    try:
        from selenium import webdriver
        from selenium.webdriver.common.by import By
        from selenium.webdriver.edge.options import Options
        from selenium.common.exceptions import NoSuchElementException
        from selenium.webdriver.common.action_chains import ActionChains
        from selenium.webdriver.common.keys import Keys
    except ImportError:
        webdriver = None

    # 数据处理相关导入
    try:
        import pandas as pd
        from docx import Document
    except ImportError:
        pd = None
        Document = None

    # 其他自动化导入
    try:
        import pyautogui
        import pyperclip
    except ImportError:
        pyautogui = None
        pyperclip = None

# 其他必要导入
import re
from ctypes import Structure, wintypes

# 使用备选方案，通过临时禁用主窗口置顶实现
def askinteger_topmost(parent, title, prompt, **kw):
    """置顶版本的askinteger"""
    # 获取主窗口
    main_window = parent.winfo_toplevel()
    # 保存当前置顶状态
    topmost = main_window.attributes('-topmost')
    
    # 临时禁用置顶
    main_window.attributes('-topmost', False)
    
    try:
        # 使用标准askinteger
        result = simpledialog.askinteger(title, prompt, parent=parent, **kw)
        return result
    finally:
        # 恢复主窗口置顶状态
        main_window.attributes('-topmost', topmost)

# 置顶消息框
def messagebox_topmost(parent, type_name, title, message, **kwargs):
    """置顶版本的messagebox"""
    # 获取主窗口
    main_window = parent.winfo_toplevel()
    # 保存当前置顶状态
    topmost = main_window.attributes('-topmost')
    
    # 临时禁用置顶
    main_window.attributes('-topmost', False)
    
    try:
        # 根据类型调用不同的messagebox函数
        if type_name == "warning":
            result = messagebox.showwarning(title, message, parent=parent, **kwargs)
        elif type_name == "error":
            result = messagebox.showerror(title, message, parent=parent, **kwargs)
        elif type_name == "info":
            result = messagebox.showinfo(title, message, parent=parent, **kwargs)
        else:
            result = messagebox.showinfo(title, message, parent=parent, **kwargs)
        return result
    finally:
        # 恢复主窗口置顶状态
        main_window.attributes('-topmost', topmost)

# 定义POINT结构用于GetCursorPos
class POINT(Structure):
    _fields_ = [("x", wintypes.LONG), ("y", wintypes.LONG)]

# 配置日志记录
log_filename = "收货信息自动采集工具.log"
# 如果日志文件已存在，则删除它（每次启动都使用新的日志文件）
if os.path.exists(log_filename):
    try:
        os.remove(log_filename)
    except:
        pass  # 如果删除失败，继续执行

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(log_filename, encoding='utf-8'),
        logging.StreamHandler()
    ]
)

logger = logging.getLogger(__name__)

class ActionSelectionDialog(tk.Toplevel):
    """动作选择弹窗，用于选择对采集元素执行的动作"""
    
    def __init__(self, parent):
        super().__init__(parent)
        self.parent = parent
        self.result = None
        self.custom_name = ""  # 添加自定义名称变量
        
        # 设置窗口属性
        self.title("选择动作")
        self.geometry("300x250")  # 增加高度以容纳新控件
        self.resizable(False, False)
        self.transient(parent)
        self.grab_set()
        
        # 居中显示
        self.update_idletasks()
        width = self.winfo_width()
        height = self.winfo_height()
        x = (self.winfo_screenwidth() // 2) - (width // 2)
        y = (self.winfo_screenheight() // 2) - (height // 2)
        self.geometry(f"+{x}+{y}")
        
        # 创建控件
        ttk.Label(self, text="请选择要执行的动作:").pack(pady=10)
        
        # 单选按钮
        self.action_var = tk.StringVar(value="getText")
        ttk.Radiobutton(self, text="获取文本", variable=self.action_var, value="getText").pack(anchor=tk.W, padx=20, pady=5)
        ttk.Radiobutton(self, text="点击元素", variable=self.action_var, value="click").pack(anchor=tk.W, padx=20, pady=5)
        ttk.Radiobutton(self, text="点击并获取剪贴板内容", variable=self.action_var, value="clickAndGetClipboard").pack(anchor=tk.W, padx=20, pady=5)
        
        # 添加自定义名称输入框
        name_frame = ttk.Frame(self)
        name_frame.pack(fill=tk.X, padx=20, pady=10)
        
        ttk.Label(name_frame, text="元素名称(必填):").pack(side=tk.LEFT)
        self.name_entry = ttk.Entry(name_frame, width=25)
        self.name_entry.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        self.name_entry.insert(0, "")  # 默认为空
        
        # 错误提示标签（初始隐藏）
        self.error_label = ttk.Label(self, text="请输入元素名称", foreground="red")
        self.error_label.pack(fill=tk.X, padx=20, pady=0)
        self.error_label.pack_forget()  # 初始隐藏
        
        # 按钮区域
        btn_frame = ttk.Frame(self)
        btn_frame.pack(fill=tk.X, pady=15)
        
        ttk.Button(btn_frame, text="确定", command=self._on_ok).pack(side=tk.RIGHT, padx=5)
        ttk.Button(btn_frame, text="取消", command=self._on_cancel).pack(side=tk.RIGHT, padx=5)
        
        # 设置焦点并等待
        self.name_entry.focus_set()  # 将焦点设置到名称输入框
        self.wait_window(self)
    
    def _on_ok(self):
        """确定按钮回调"""
        self.custom_name = self.name_entry.get().strip()  # 获取并清除首尾空格
        
        # 验证名称是否已填写
        if not self.custom_name:
            self.error_label.pack()  # 显示错误提示
            self.name_entry.focus_set()  # 将焦点设回输入框
            return
            
        self.result = self.action_var.get()
        self.destroy()
    
    def _on_cancel(self):
        """取消按钮回调"""
        self.result = None
        self.custom_name = ""
        self.destroy()


class OperationSequenceDialog(tk.Toplevel):
    """操作选择与排序对话框，用于配置要执行的操作及其顺序"""
    
    def __init__(self, parent, elements_data=None):
        """
        初始化对话框
        parent: 父窗口
        elements_data: 从JSON文件加载的元素数据列表
        """
        super().__init__(parent)
        self.parent = parent
        self.elements_data = elements_data or []
        self.result = None
        
        # 设置窗口属性
        self.title("操作选择与顺序配置")
        self.geometry("800x500")  # 较大的窗口以显示表格和控件
        self.resizable(True, True)
        self.transient(parent)
        self.grab_set()
        
        # 居中显示
        self.update_idletasks()
        width = self.winfo_width()
        height = self.winfo_height()
        x = (self.winfo_screenwidth() // 2) - (width // 2)
        y = (self.winfo_screenheight() // 2) - (height // 2)
        self.geometry(f"+{x}+{y}")
        
        # 创建主框架
        main_frame = ttk.Frame(self)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 创建操作表格
        self._create_operations_table(main_frame)
        
        # 指定一个元素作为订单数量元素的框架
        count_frame = ttk.LabelFrame(main_frame, text="订单数量元素选择")
        count_frame.pack(fill=tk.X, pady=10)
        
        ttk.Label(count_frame, text="选择一个包含订单数量的元素(可选):").pack(side=tk.LEFT, padx=5)
        self.count_element_var = tk.StringVar(value="")
        self.count_element_combo = ttk.Combobox(
            count_frame, 
            textvariable=self.count_element_var,
            state="readonly",
            width=30
        )
        # 填充下拉框选项
        count_options = ["(不使用自动检测)"] + [elem["name"] for elem in self.elements_data]
        self.count_element_combo["values"] = count_options
        self.count_element_combo.current(0)  # 默认选择"不使用自动检测"
        self.count_element_combo.pack(side=tk.LEFT, padx=5)
        
        # 按钮区域
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=10)
        
        ttk.Button(btn_frame, text="全部选中", command=self._select_all).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="全部取消", command=self._deselect_all).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="按名称排序", command=self._sort_by_name).pack(side=tk.LEFT, padx=5)
        
        ttk.Button(btn_frame, text="确定", command=self._on_ok).pack(side=tk.RIGHT, padx=5)
        ttk.Button(btn_frame, text="取消", command=self._on_cancel).pack(side=tk.RIGHT, padx=5)
        
        # 创建元素变量
        self._create_element_variables()
        
        # 填充表格
        self._populate_table()
        
        # 等待窗口关闭
        self.wait_window(self)
    
    def _create_operations_table(self, parent):
        """创建操作表格"""
        # 创建表格框架
        table_frame = ttk.LabelFrame(parent, text="操作配置")
        table_frame.pack(fill=tk.BOTH, expand=True)
        
        # 创建表格视图
        columns = ("enabled", "order", "name", "action", "loop_mode", "preview")  # 添加loop_mode列
        self.tree = ttk.Treeview(table_frame, columns=columns, show="headings", selectmode="browse")
        
        # 设置列标题
        self.tree.heading("enabled", text="启用")
        self.tree.heading("order", text="顺序")
        self.tree.heading("name", text="元素名称")
        self.tree.heading("action", text="操作类型")
        self.tree.heading("loop_mode", text="循环模式")  # 新增列标题
        self.tree.heading("preview", text="预览")
        
        # 设置列宽度
        self.tree.column("enabled", width=50, anchor=tk.CENTER)
        self.tree.column("order", width=50, anchor=tk.CENTER)
        self.tree.column("name", width=200, anchor=tk.W)
        self.tree.column("action", width=150, anchor=tk.CENTER)
        self.tree.column("loop_mode", width=100, anchor=tk.CENTER)  # 设置新列宽度
        self.tree.column("preview", width=80, anchor=tk.CENTER)
        
        # 添加滚动条
        scrollbar = ttk.Scrollbar(table_frame, orient=tk.VERTICAL, command=self.tree.yview)
        self.tree.configure(yscrollcommand=scrollbar.set)
        
        # 布局表格和滚动条
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 绑定事件处理
        self.tree.bind("<ButtonRelease-1>", self._on_tree_click)
        self.tree.bind("<Double-1>", self._on_tree_double_click)
    
    def _create_element_variables(self):
        """为每个元素创建变量"""
        # 存储每个元素的启用状态、顺序、操作类型和循环模式
        self.element_vars = []
        
        for i, elem in enumerate(self.elements_data):
            enabled_var = tk.BooleanVar(value=elem.get("enabled", True))
            order_var = tk.IntVar(value=elem.get("order", i+1))
            action_var = tk.StringVar(value=elem.get("action", "getText"))
            loop_mode_var = tk.StringVar(value=elem.get("loop_mode", "always"))  # 默认为"始终循环"
            
            self.element_vars.append({
                "element_id": elem["element_id"],
                "enabled_var": enabled_var,
                "order_var": order_var,
                "action_var": action_var,
                "loop_mode_var": loop_mode_var  # 添加循环模式变量
            })
    
    def _populate_table(self):
        """填充表格数据"""
        # 先清空表格
        for item in self.tree.get_children():
            self.tree.delete(item)
        
        # 按顺序排序元素
        sorted_elements = sorted(self.elements_data, key=lambda x: x.get("order", 0))
        
        # 添加元素到表格
        for elem in sorted_elements:
            # 获取对应的变量
            vars_dict = next((v for v in self.element_vars if v["element_id"] == elem["element_id"]), None)
            if vars_dict:
                # 创建表格项
                enabled_text = "✓" if vars_dict["enabled_var"].get() else ""
                order = vars_dict["order_var"].get()
                name = elem.get("name", "未命名")
                action = vars_dict["action_var"].get()
                action_text = {
                    "getText": "获取文本",
                    "click": "点击元素",
                    "clickAndGetClipboard": "点击并获取剪贴板"
                }.get(action, action)
                
                # 获取循环模式文本
                loop_mode = vars_dict["loop_mode_var"].get()
                loop_mode_text = "单次循环" if loop_mode == "once" else "始终循环"
                
                item_values = (enabled_text, order, name, action_text, loop_mode_text, "预览")
                item_id = self.tree.insert("", "end", values=item_values)
                
                # 存储元素ID与表格项的映射
                self.tree.item(item_id, tags=(str(elem["element_id"]),))
    
    def _on_tree_click(self, event):
        """处理表格点击事件"""
        # 获取点击的列
        region = self.tree.identify_region(event.x, event.y)
        column = self.tree.identify_column(event.x)
        
        if region == "cell":
            item_id = self.tree.identify_row(event.y)
            if not item_id:
                return
                
            # 获取元素ID
            elem_id = int(self.tree.item(item_id)["tags"][0])
            
            # 根据点击的列执行不同操作
            if column == "#1":  # 启用列
                self._toggle_element_enabled(elem_id, item_id)
            elif column == "#2":  # 顺序列
                self._edit_element_order(elem_id, item_id)
            elif column == "#4":  # 操作类型列
                self._edit_element_action(elem_id, item_id)
            elif column == "#5":  # 循环模式列
                self._edit_element_loop_mode(elem_id, item_id)
            elif column == "#6":  # 预览列
                self._preview_element(elem_id)
    
    def _on_tree_double_click(self, event):
        """处理表格双击事件"""
        region = self.tree.identify_region(event.x, event.y)
        if region == "cell":
            item_id = self.tree.identify_row(event.y)
            if item_id:
                # 双击时编辑元素顺序
                elem_id = int(self.tree.item(item_id)["tags"][0])
                self._edit_element_order(elem_id, item_id)
    
    def _toggle_element_enabled(self, elem_id, item_id):
        """切换元素启用状态"""
        # 查找对应的变量
        for var in self.element_vars:
            if var["element_id"] == elem_id:
                # 切换状态
                new_state = not var["enabled_var"].get()
                var["enabled_var"].set(new_state)
                
                # 更新表格显示
                current_values = list(self.tree.item(item_id, "values"))
                current_values[0] = "✓" if new_state else ""
                self.tree.item(item_id, values=current_values)
                break
    
    def _edit_element_order(self, elem_id, item_id):
        """编辑元素顺序并自动调整其他元素的顺序"""
        # 查找对应的变量
        for var in self.element_vars:
            if var["element_id"] == elem_id:
                # 当前顺序
                current_order = var["order_var"].get()
                
                # 创建一个简单的输入对话框
                new_order = askinteger_topmost(self, "编辑顺序", f"请输入元素的执行顺序 (1-{len(self.elements_data)}):", initialvalue=current_order, minvalue=1, maxvalue=len(self.elements_data))
                
                # 如果用户取消或输入相同的值，则不进行任何更改
                if new_order is None or new_order == current_order:
                    return
                
                # 更新当前元素的顺序
                var["order_var"].set(new_order)
                
                # 调整其他元素的顺序以避免重复
                self._adjust_other_elements_order(elem_id, current_order, new_order)
                
                # 更新表格显示
                current_values = list(self.tree.item(item_id, "values"))
                current_values[1] = str(new_order)
                self.tree.item(item_id, values=tuple(current_values))
                
                # 重新排序表格
                self._resort_table()
                break
    
    def _edit_element_action(self, elem_id, item_id):
        """编辑元素操作类型"""
        # 查找对应的变量
        for var in self.element_vars:
            if var["element_id"] == elem_id:
                # 当前操作类型
                current_action = var["action_var"].get()
                
                # 创建操作类型选择对话框
                dialog = tk.Toplevel(self)
                dialog.title("选择操作类型")
                dialog.geometry("300x200")
                dialog.transient(self)
                dialog.grab_set()
                
                # 设置窗口置顶
                dialog.attributes('-topmost', True)
                
                # 居中显示
                dialog.update_idletasks()
                x = self.winfo_rootx() + (self.winfo_width() // 2) - (300 // 2)
                y = self.winfo_rooty() + (self.winfo_height() // 2) - (200 // 2)
                dialog.geometry(f"+{x}+{y}")
                
                # 创建单选按钮
                action_var = tk.StringVar(value=current_action)
                ttk.Label(dialog, text="请选择操作类型:").pack(pady=10)
                
                ttk.Radiobutton(
                    dialog, 
                    text="获取文本", 
                    variable=action_var, 
                    value="getText"
                ).pack(anchor=tk.W, padx=20, pady=5)
                
                ttk.Radiobutton(
                    dialog, 
                    text="点击元素", 
                    variable=action_var, 
                    value="click"
                ).pack(anchor=tk.W, padx=20, pady=5)
                
                ttk.Radiobutton(
                    dialog, 
                    text="点击并获取剪贴板内容", 
                    variable=action_var, 
                    value="clickAndGetClipboard"
                ).pack(anchor=tk.W, padx=20, pady=5)
                
                # 按钮区域
                btn_frame = ttk.Frame(dialog)
                btn_frame.pack(fill=tk.X, pady=15)
                
                # 确定按钮回调
                def on_ok():
                    new_action = action_var.get()
                    if new_action != current_action:
                        var["action_var"].set(new_action)
                        
                        # 更新表格显示
                        current_values = list(self.tree.item(item_id, "values"))
                        action_text = {
                            "getText": "获取文本",
                            "click": "点击元素",
                            "clickAndGetClipboard": "点击并获取剪贴板"
                        }.get(new_action, new_action)
                        current_values[3] = action_text
                        self.tree.item(item_id, values=current_values)
                    
                    dialog.destroy()
                    # --- 焦点恢复 ---
                    try:
                        if hasattr(self.parent, '_manage_focus'):
                            self.parent._manage_focus()
                            self.parent.root.after(100, self.parent._manage_focus)
                            self.parent._log_info("已尝试恢复主窗口焦点", "blue")
                    except Exception as e:
                        print(f"焦点恢复失败: {e}")
                
                ttk.Button(btn_frame, text="确定", command=on_ok).pack(side=tk.RIGHT, padx=5)
                ttk.Button(btn_frame, text="取消", command=dialog.destroy).pack(side=tk.RIGHT, padx=5)
                
                # 等待对话框关闭
                self.wait_window(dialog)
                break
    
    def _edit_element_loop_mode(self, elem_id, item_id):
        """编辑元素循环模式"""
        # 查找对应的变量
        for var in self.element_vars:
            if var["element_id"] == elem_id:
                # 当前循环模式
                current_mode = var["loop_mode_var"].get()
                
                # 创建循环模式选择对话框
                dialog = tk.Toplevel(self)
                dialog.title("选择循环模式")
                dialog.geometry("300x200")
                dialog.transient(self)
                dialog.grab_set()
                
                # 设置窗口置顶
                dialog.attributes('-topmost', True)
                
                # 居中显示
                dialog.update_idletasks()
                x = self.winfo_rootx() + (self.winfo_width() // 2) - (300 // 2)
                y = self.winfo_rooty() + (self.winfo_height() // 2) - (200 // 2)
                dialog.geometry(f"+{x}+{y}")
                
                # 创建单选按钮
                mode_var = tk.StringVar(value=current_mode)
                ttk.Label(dialog, text="请选择循环模式:").pack(pady=10)
                
                ttk.Radiobutton(
                    dialog, 
                    text="单次循环 (仅在第一个订单执行)", 
                    variable=mode_var, 
                    value="once"
                ).pack(anchor=tk.W, padx=20, pady=5)
                
                ttk.Radiobutton(
                    dialog, 
                    text="始终循环 (每个订单都执行)", 
                    variable=mode_var, 
                    value="always"
                ).pack(anchor=tk.W, padx=20, pady=5)
                
                # 按钮区域
                btn_frame = ttk.Frame(dialog)
                btn_frame.pack(fill=tk.X, pady=15)
                
                # 确定按钮回调
                def on_ok():
                    new_mode = mode_var.get()
                    if new_mode != current_mode:
                        var["loop_mode_var"].set(new_mode)
                        
                        # 更新表格显示
                        current_values = list(self.tree.item(item_id, "values"))
                        loop_mode_text = "单次循环" if new_mode == "once" else "始终循环"
                        current_values[4] = loop_mode_text
                        self.tree.item(item_id, values=current_values)
                    
                    dialog.destroy()
                    # --- 焦点恢复 ---
                    try:
                        if hasattr(self.parent, '_manage_focus'):
                            self.parent._manage_focus()
                            self.parent.root.after(100, self.parent._manage_focus)
                            self.parent._log_info("已尝试恢复主窗口焦点", "blue")
                    except Exception as e:
                        print(f"焦点恢复失败: {e}")
                
                ttk.Button(btn_frame, text="确定", command=on_ok).pack(side=tk.RIGHT, padx=5)
                ttk.Button(btn_frame, text="取消", command=dialog.destroy).pack(side=tk.RIGHT, padx=5)
                
                # 等待对话框关闭
                self.wait_window(dialog)
                break
    
    def _preview_element(self, elem_id):
        """预览元素（触发外部预览回调）"""
        # 查找对应的元素数据
        elem = next((e for e in self.elements_data if e["element_id"] == elem_id), None)
        if elem and "xpath" in elem:
            # 通知父窗口进行预览
            # 注意：这里需要父窗口提供一个预览方法
            if hasattr(self.parent, "_preview_element"):
                xpath = elem["xpath"]
                success = self.parent._preview_element(xpath)
                if not success:
                    messagebox_topmost(self, "warning", "预览失败", "无法定位或高亮显示元素，请检查XPath路径是否正确")
    
    def _resort_table(self):
        """根据顺序重新排序表格"""
        # 获取当前表格中的所有项
        items = []
        for item_id in self.tree.get_children():
            elem_id = int(self.tree.item(item_id)["tags"][0])
            var = next((v for v in self.element_vars if v["element_id"] == elem_id), None)
            if var:
                order = var["order_var"].get()
                items.append((item_id, elem_id, order))
        
        # 按顺序排序
        items.sort(key=lambda x: x[2])
        
        # 重新排列表格
        for index, (item_id, _, _) in enumerate(items):
            self.tree.move(item_id, "", index)
    
    def _adjust_other_elements_order(self, changed_elem_id, old_order, new_order):
        """调整其他元素的顺序以避免重复"""
        # 如果顺序没有变化，则不需要调整
        if old_order == new_order:
            return
            
        # 确定调整方向和范围
        if old_order < new_order:
            # 向下移动：将中间的元素顺序值减1
            for var in self.element_vars:
                if var["element_id"] != changed_elem_id:
                    current_order = var["order_var"].get()
                    if old_order < current_order <= new_order:
                        var["order_var"].set(current_order - 1)
        else:
            # 向上移动：将中间的元素顺序值加1
            for var in self.element_vars:
                if var["element_id"] != changed_elem_id:
                    current_order = var["order_var"].get()
                    if new_order <= current_order < old_order:
                        var["order_var"].set(current_order + 1)
        
        # 更新所有表格项的显示
        for item_id in self.tree.get_children():
            elem_id = int(self.tree.item(item_id)["tags"][0])
            var = next((v for v in self.element_vars if v["element_id"] == elem_id), None)
            if var:
                current_values = list(self.tree.item(item_id, "values"))
                current_values[1] = str(var["order_var"].get())
                self.tree.item(item_id, values=tuple(current_values))
    
    def _select_all(self):
        """选择所有元素"""
        for var in self.element_vars:
            var["enabled_var"].set(True)
        
        # 更新表格显示
        for item_id in self.tree.get_children():
            current_values = list(self.tree.item(item_id, "values"))
            current_values[0] = "✓"
            self.tree.item(item_id, values=current_values)
    
    def _deselect_all(self):
        """取消选择所有元素"""
        for var in self.element_vars:
            var["enabled_var"].set(False)
        
        # 更新表格显示
        for item_id in self.tree.get_children():
            current_values = list(self.tree.item(item_id, "values"))
            current_values[0] = ""
            self.tree.item(item_id, values=current_values)
    
    def _sort_by_name(self):
        """按名称排序元素顺序"""
        # 按名称排序元素
        named_elements = [(elem["element_id"], elem.get("name", "")) for elem in self.elements_data]
        named_elements.sort(key=lambda x: x[1])
        
        # 保存原始顺序，以便后续调整
        original_orders = {}
        for var in self.element_vars:
            original_orders[var["element_id"]] = var["order_var"].get()
        
        # 按名称顺序设置新的顺序值
        for i, (elem_id, _) in enumerate(named_elements):
            var = next((v for v in self.element_vars if v["element_id"] == elem_id), None)
            if var:
                # 记录原始顺序
                old_order = var["order_var"].get()
                # 设置新顺序
                new_order = i + 1
                
                if old_order != new_order:
                    var["order_var"].set(new_order)
        
        # 更新表格显示
        for item_id in self.tree.get_children():
            elem_id = int(self.tree.item(item_id)["tags"][0])
            var = next((v for v in self.element_vars if v["element_id"] == elem_id), None)
            if var:
                current_values = list(self.tree.item(item_id, "values"))
                current_values[1] = var["order_var"].get()
                self.tree.item(item_id, values=current_values)
        
        # 重新排序表格
        self._resort_table()
    
    def _on_ok(self):
        """确定按钮回调"""
        # 收集结果
        result = []
        
        # 处理订单数量元素选择
        count_element_name = self.count_element_var.get()
        
        for i, elem in enumerate(self.elements_data):
            var = next((v for v in self.element_vars if v["element_id"] == elem["element_id"]), None)
            if var:
                # 检查是否为订单数量元素
                # 如果选择了"(不使用自动检测)"，则所有元素的is_order_count都为False
                if count_element_name == "(不使用自动检测)":
                    is_order_count = False
                else:
                    is_order_count = (count_element_name == elem.get("name", ""))
                
                result.append({
                    "element_id": elem["element_id"],
                    "name": elem.get("name", f"元素{i+1}"),
                    "xpath": elem.get("xpath", ""),
                    "action": var["action_var"].get(),
                    "order": var["order_var"].get(),
                    "enabled": var["enabled_var"].get(),
                    "is_order_count": is_order_count,
                    "loop_mode": var["loop_mode_var"].get()  # 添加循环模式
                })
        
        # 保存结果
        self.result = result
        self.destroy()
    
    def _on_cancel(self):
        """取消按钮回调"""
        self.result = None
        self.destroy()


class ShippingInfoCollector:
    """收货信息自动采集工具主类，集成所有功能模块"""
    
    def __init__(self, root):
        """初始化应用程序"""
        self.root = root
        self.root.title("收货信息自动采集工具")
        self.root.geometry("800x600")
        
        # 设置应用程序图标（如果有的话）
        # self.root.iconbitmap("icon.ico")
        
        # 设置自动执行动作的间隔时间（秒）
        self.auto_action_interval = 1.0
        
        # 状态变量
        self.is_running = False  # 是否正在运行
        self.is_paused = False   # 是否暂停
        self.is_browser_connected = False  # 浏览器连接状态
        self.collected_data = []  # 收集的数据
        self.browser_process = None  # 浏览器进程
        self.ws = None  # WebSocket连接
        self.session_id = None  # 会话ID
        self.request_id = 1  # CDP请求ID
        self.collection_mode = tk.StringVar(value="正常模式")  # 默认为正常模式
        
        # 辅助定位相关状态变量
        self.is_distance_learning = False  # 是否处于距离学习模式
        self.distance_learning_step = 0    # 当前学习步骤（0未开始，1第一点已采集，2完成）
        self.first_element_position = None # 第一个元素位置信息
        self.second_element_position = None # 第二个元素位置信息
        self.scroll_distance = None        # 计算出的滚动距离
        
        # WASD微调相关状态变量
        self.element_offsets = {}        # 按元素名称保存的偏移量字典
        self.consecutive_same_order = 0    # 连续重复订单的计数
        self.scroll_distance_multiplier = 1.0  # 滚动距离调整系数
        
        # 配置变量
        self.always_on_top = tk.BooleanVar(value=True)  # 窗口置顶
        self.confirm_click = tk.BooleanVar(value=False)  # 点击前确认
        self.debug_port = 9222  # 调试端口
        self.user_data_dir = os.path.join(tempfile.gettempdir(), "edge_user_data")  # 本地用户数据目录，确保数据复用
        self.driver = None  # Selenium WebDriver实例
        
        # 添加操作序列变量
        self.operation_sequence = []  # 包含多个操作对象的列表，按order排序
        
        # 初始化元素特定的偏移量
        self.element_offsets = {}
        
        # 初始化全局变量用于保存最后一次成功的剪贴板内容
        self.last_clipboard_content = ""
        
        # 初始化订单ID与收货信息的映射字典
        self.order_clipboard_contents = {}
        
        # 初始化收货信息验证相关变量
        self.orders_need_review = set()  # 需要人工审核的订单ID集合
        self.content_validation_results = {}  # 内容验证结果
        self.last_captured_order_id = None  # 最后捕获的订单ID
        
        # 创建GUI组件
        self._create_gui()
        
        # 设置窗口置顶状态
        self._update_always_on_top()
        
        # 绑定键盘事件
        self.root.bind("<Key>", self._handle_key_event)
        
        # 加载偏移量配置（按元素名称保存的WASD微调配置）
        self._load_offset_config()
        
        # 每次启动时清空所有历史映射数据，确保收集全新信息
        self.order_clipboard_contents = {}
        # 同时清空映射文件，确保不会重新加载旧数据
        self._save_clipboard_mappings()
        self._log_info("已清空所有历史映射数据和文件，准备收集全新信息", "green")
        
        # 启动时记录日志
        self._log_info("程序已启动，等待操作...", "blue")
        self._log_info(f"订单ID与收货信息映射字典初始化完成，包含 {len(self.order_clipboard_contents)} 个映射", "blue")
        print(f"DEBUG-INIT: 初始化订单ID与收货信息映射字典，包含 {len(self.order_clipboard_contents)} 个映射")
        
        # 添加剪贴板监听状态变量
        self.clipboard_monitor_active = False
        self.last_known_clipboard = ""
        self.clipboard_monitor_thread = None
        
        # 添加订单ID缓存
        self.current_order_id = None
        self.last_order_ids = []  # 保存最近的几个订单ID，用于去重和恢复
        
        # 验证码检测相关变量
        self.captcha_running = False
        self.captcha_detected = False
        self.last_detection_time = 0
        self.captcha_detection_thread = None
        self.template_images = []
        self.template_paths = []
        self.monitor_area = None  # 监控区域 (x1, y1, x2, y2)
        self.detection_interval = 0.25  # 检测间隔（秒）
        self.similarity_threshold = 0.7  # 相似度阈值
        self.consecutive_frames = 3  # 连续检测帧数
        self.consecutive_detections = 0  # 连续检测到验证码的次数
        self.consecutive_non_detections = 0  # 连续未检测到验证码的次数
        self.use_mask_detection = True  # 是否使用遮罩层检测
        self.mask_threshold = 0.1  # 遮罩层检测灰度差异阈值
        self.temp_dir = tempfile.mkdtemp()  # 创建临时目录用于存储截图
        self.screenshot_files = []  # 存储截图文件路径
        self.target_window_handle = None  # 目标窗口句柄
        self.target_window_title = ""  # 目标窗口标题
        self.use_window_capture = True  # 是否使用窗口截图（默认启用）
        self.captcha_force_stop = False  # 验证码强制停止标志
        self.force_stop_flag = False  # 操作强制停止标志（用于验证码检测时暂停操作）
        
        # 程序启动时自动加载验证码模板（必须在template_images初始化之后）
        self._auto_load_captcha_templates()
        
        # 翻页功能相关变量
        self.next_page_xpath = None  # 翻页按钮的XPath
        self.next_page_collected = False  # 是否已采集翻页按钮
        self.page_turn_count = 0  # 当前翻页计数
        self.target_page_count = 20  # 目标翻页页数（默认20）
        self.screenshot_dir = "page_screenshots"  # 截图保存目录
        self.collecting_page_turn = False  # 是否正在采集翻页元素
        
        # 创建截图目录
        if not os.path.exists(self.screenshot_dir):
            os.makedirs(self.screenshot_dir)
    
    def _create_gui(self):
        """创建GUI界面"""
        # 创建主框架并使用网格布局
        self.main_frame = ttk.Frame(self.root)
        self.main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # ====== 顶部区域：浏览器连接状态与模式切换 ======
        self.top_frame = ttk.Frame(self.main_frame)
        self.top_frame.pack(fill=tk.X, pady=(0, 10))
        
        # 浏览器连接按钮和状态
        self.browser_frame = ttk.Frame(self.top_frame)
        self.browser_frame.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        self.browser_button = ttk.Button(
            self.browser_frame, 
            text="打开浏览器", 
            command=self._start_browser
        )
        self.browser_button.pack(side=tk.LEFT, padx=(0, 10))
        
        ttk.Label(self.browser_frame, text="浏览器状态:").pack(side=tk.LEFT)
        self.browser_status = ttk.Label(
            self.browser_frame, 
            text="未连接",
            foreground="red"
        )
        self.browser_status.pack(side=tk.LEFT, padx=5)
        
        # 浏览器连接进度条
        self.connection_progress = ttk.Progressbar(
            self.browser_frame, 
            orient=tk.HORIZONTAL, 
            length=200, 
            mode='determinate'
        )
        self.connection_progress.pack(side=tk.LEFT, padx=10)
        
        # 模式切换下拉框
        self.mode_frame = ttk.Frame(self.top_frame)
        self.mode_frame.pack(side=tk.RIGHT)
        
        ttk.Label(self.mode_frame, text="操作模式:").pack(side=tk.LEFT)
        self.mode_combobox = ttk.Combobox(
            self.mode_frame, 
            textvariable=self.collection_mode,
            values=["正常模式", "采集模式"],
            state="readonly",
            width=10
        )
        self.mode_combobox.pack(side=tk.LEFT, padx=5)
        self.mode_combobox.bind("<<ComboboxSelected>>", self._mode_changed)
        
        # ====== 中间区域：控制按钮和状态显示 ======
        self.middle_frame = ttk.Frame(self.main_frame)
        self.middle_frame.pack(fill=tk.BOTH, expand=True, pady=10)
        
        # 控制按钮区
        self.control_frame = ttk.Frame(self.middle_frame)
        self.control_frame.pack(fill=tk.X)
        
        self.start_button = ttk.Button(
            self.control_frame, 
            text="开始", 
            command=self._start_collection,
            state=tk.DISABLED  # 初始禁用，直到浏览器连接
        )
        self.start_button.pack(side=tk.LEFT, padx=5)
        
        self.pause_button = ttk.Button(
            self.control_frame, 
            text="暂停", 
            command=self._pause_collection,
            state=tk.DISABLED
        )
        self.pause_button.pack(side=tk.LEFT, padx=5)
        
        self.continue_button = ttk.Button(
            self.control_frame, 
            text="继续", 
            command=self._continue_collection,
            state=tk.DISABLED
        )
        self.continue_button.pack(side=tk.LEFT, padx=5)
        
        self.stop_button = ttk.Button(
            self.control_frame, 
            text="终止", 
            command=self._stop_collection,
            state=tk.DISABLED
        )
        self.stop_button.pack(side=tk.LEFT, padx=5)
        
        # 添加配置操作按钮
        self.configure_button = ttk.Button(
            self.control_frame, 
            text="配置操作", 
            command=self._configure_operations,
            state=tk.DISABLED  # 初始禁用，直到浏览器连接
        )
        self.configure_button.pack(side=tk.LEFT, padx=5)
        
        # 添加验证码管理按钮
        self.captcha_manage_button = ttk.Button(
            self.control_frame,
            text="验证码管理",
            command=self._show_captcha_manager
        )
        self.captcha_manage_button.pack(side=tk.LEFT, padx=5)
        
        # 添加验证码选择目标窗口按钮
        self.captcha_window_button = ttk.Button(
            self.control_frame,
            text="选择目标窗口",
            command=self._select_captcha_target_window
        )
        self.captcha_window_button.pack(side=tk.LEFT, padx=5)
        
        # 添加翻页功能相关按钮
        self.collect_page_btn = ttk.Button(
            self.control_frame,
            text="采集翻页元素",
            command=self.collect_page_turn_element,
            state=tk.DISABLED  # 初始禁用，直到浏览器连接
        )
        self.collect_page_btn.pack(side=tk.LEFT, padx=5)
        
        # 翻页页数选择下拉框
        ttk.Label(self.control_frame, text="翻页页数:").pack(side=tk.LEFT, padx=(10, 0))
        self.page_count_var = tk.StringVar(value="20")
        self.page_count_combo = ttk.Combobox(
            self.control_frame,
            textvariable=self.page_count_var,
            values=["20", "50"],
            state="readonly",
            width=5
        )
        self.page_count_combo.pack(side=tk.LEFT, padx=5)
        self.page_count_combo.current(0)  # 默认选择20
        
        # 绑定翻页页数变化事件
        self.page_count_var.trace('w', self.on_page_count_changed)
        
        # 进度显示区
        self.progress_frame = ttk.Frame(self.middle_frame)
        self.progress_frame.pack(fill=tk.X, pady=10)
        
        ttk.Label(self.progress_frame, text="进度:").pack(side=tk.LEFT)
        self.progress_label = ttk.Label(self.progress_frame, text="0/0")
        self.progress_label.pack(side=tk.LEFT, padx=5)
        
        self.progress_bar = ttk.Progressbar(
            self.progress_frame, 
            orient=tk.HORIZONTAL, 
            length=400, 
            mode='determinate'
        )
        self.progress_bar.pack(side=tk.LEFT, padx=10, fill=tk.X, expand=True)
        
        # 状态文本显示区（带颜色）
        self.status_frame = ttk.LabelFrame(self.middle_frame, text="状态信息")
        self.status_frame.pack(fill=tk.BOTH, expand=True, pady=10)
        
        self.status_text = ScrolledText(
            self.status_frame, 
            wrap=tk.WORD,
            height=10
        )
        self.status_text.pack(fill=tk.BOTH, expand=True)
        self.status_text.config(state=tk.DISABLED)  # 初始设置为只读
        
        # 配置状态文本的颜色标签
        self.status_text.tag_config("color_red", foreground="red")
        self.status_text.tag_config("color_green", foreground="green")
        self.status_text.tag_config("color_blue", foreground="blue")
        self.status_text.tag_config("color_orange", foreground="orange")
        
        # 验证码状态显示区
        self.captcha_status_frame = ttk.LabelFrame(self.middle_frame, text="验证码检测状态")
        self.captcha_status_frame.pack(fill=tk.X, pady=5)
        
        # 验证码状态标签
        self.captcha_status_label = ttk.Label(
            self.captcha_status_frame, 
            text="验证码检测: 未启动", 
            font=("Arial", 12)
        )
        self.captcha_status_label.pack(side=tk.LEFT, padx=10, pady=5)
        
        # 验证码状态指示器
        self.captcha_status_indicator = tk.Canvas(
            self.captcha_status_frame, 
            width=20, 
            height=20, 
            bg=self.root.cget('bg'), 
            highlightthickness=0
        )
        self.captcha_status_indicator.pack(side=tk.LEFT, padx=5, pady=5)
        self.captcha_status_indicator.create_oval(2, 2, 18, 18, fill="gray", outline="")
        
        # 目标窗口状态标签
        self.target_window_label = ttk.Label(
            self.captcha_status_frame,
            text="目标窗口: 未设置"
        )
        self.target_window_label.pack(side=tk.LEFT, padx=10, pady=5)
        
        # 模板数量标签
        self.template_count_label = ttk.Label(
            self.captcha_status_frame,
            text="模板数量: 0"
        )
        self.template_count_label.pack(side=tk.LEFT, padx=10, pady=5)
        

        
        # ====== 底部区域：设置与导出 ======
        self.bottom_frame = ttk.Frame(self.main_frame)
        self.bottom_frame.pack(fill=tk.X, pady=(10, 0))
        
        # 设置区
        self.settings_frame = ttk.Frame(self.top_frame)
        self.settings_frame.pack(side=tk.RIGHT, padx=5)
        
        self.always_on_top_check = ttk.Checkbutton(
            self.settings_frame,
            text="窗口置顶",
            variable=self.always_on_top,
            command=self._update_always_on_top
        )
        self.always_on_top_check.pack(side=tk.LEFT, padx=10)
        
        self.confirm_click_check = ttk.Checkbutton(
            self.settings_frame,
            text="点击前确认",
            variable=self.confirm_click
        )
        self.confirm_click_check.pack(side=tk.LEFT, padx=10)
        
        # 添加操作间隔设置
        self.interval_frame = ttk.Frame(self.settings_frame)
        self.interval_frame.pack(side=tk.LEFT, padx=10)
        
        ttk.Label(self.interval_frame, text="操作间隔(秒):").pack(side=tk.LEFT)
        self.interval_var = tk.StringVar(value=str(self.auto_action_interval))
        self.interval_entry = ttk.Spinbox(
            self.interval_frame, 
            from_=0.1, 
            to=10.0, 
            increment=0.1, 
            width=5, 
            textvariable=self.interval_var,
            command=self._update_interval
        )
        self.interval_entry.pack(side=tk.LEFT, padx=5)
        self.interval_entry.bind('<Return>', self._update_interval)
        self.interval_entry.bind('<FocusOut>', self._update_interval)
        
        # 添加偏移量显示和管理按钮
        self.offset_frame = ttk.Frame(self.settings_frame)
        self.offset_frame.pack(side=tk.LEFT, padx=10)
        
        self.offset_label = ttk.Label(self.offset_frame, text="元素偏移量")
        self.offset_label.pack(side=tk.LEFT, padx=5)
        
        self.offset_manage_button = ttk.Button(
            self.offset_frame,
            text="管理偏移量",
            command=self._show_offset_manager
        )
        self.offset_manage_button.pack(side=tk.LEFT, padx=5)
        
        # 导出区
        self.export_frame = ttk.LabelFrame(self.bottom_frame, text="数据导出")
        self.export_frame.pack(side=tk.RIGHT)
        
        # 正常模式下使用的按钮
        self.excel_button = ttk.Button(
            self.export_frame,
            text="导出Excel",
            command=self._export_excel,
            state=tk.DISABLED  # 初始禁用，直到有数据
        )
        self.excel_button.pack(side=tk.LEFT, padx=5, pady=5)
        
        self.word_button = ttk.Button(
            self.export_frame,
            text="导出Word",
            command=self._export_word,
            state=tk.DISABLED  # 初始禁用，直到有数据
        )
        self.word_button.pack(side=tk.LEFT, padx=5, pady=5)
        
        # 采集模式下使用的按钮
        self.json_button = ttk.Button(
            self.export_frame,
            text="导出JSON",
            command=self._export_json,
            state=tk.DISABLED  # 初始禁用
        )
        # 初始不显示，根据模式切换显示
    
        # 偏移量管理区
        self.offset_frame = ttk.Frame(self.settings_frame)
        self.offset_frame.pack(side=tk.LEFT, padx=10)
        self.offset_label = ttk.Label(self.offset_frame, text="元素偏移量")
        self.offset_label.pack(side=tk.LEFT, padx=5)
        self.offset_manage_button = ttk.Button(
            self.offset_frame,
            text="管理偏移量",
            command=self._show_offset_manager
        )
        self.offset_manage_button.pack(side=tk.LEFT, padx=5)
        
        # 剪贴板映射管理区
        self.clipboard_frame = ttk.Frame(self.settings_frame)
        self.clipboard_frame.pack(side=tk.LEFT, padx=10)
        ttk.Button(
            self.clipboard_frame,
            text="保存剪贴板映射",
            command=self._save_clipboard_mappings
        ).pack(side=tk.LEFT, padx=5)
        ttk.Button(
            self.clipboard_frame,
            text="手动关联订单ID",
            command=lambda: self._manually_associate_clipboard_with_order_id()
        ).pack(side=tk.LEFT, padx=5)

        # 智能循环与滚动设置区
        self.smart_loop_frame = ttk.LabelFrame(self.main_frame, text="智能循环与滚动设置")
        self.smart_loop_frame.pack(fill=tk.X, padx=10, pady=5)
        self.ref1_xpath = None
        self.ref2_xpath = None
        self.scroll_container_xpath = None  # 兼容性保留，但不再采集
        self.scroll_step = None
        ttk.Button(self.smart_loop_frame, text="采集第1个订单参照点", command=self._collect_ref1_xpath).pack(side=tk.LEFT, padx=5)
        ttk.Button(self.smart_loop_frame, text="采集第2个订单参照点", command=self._collect_ref2_xpath).pack(side=tk.LEFT, padx=5)
        self.ref1_label = ttk.Label(self.smart_loop_frame, text="第1个订单参照点: 未设置")
        self.ref1_label.pack(side=tk.LEFT, padx=5)
        self.ref2_label = ttk.Label(self.smart_loop_frame, text="第2个订单参照点: 未设置")
        self.ref2_label.pack(side=tk.LEFT, padx=5)
    
    def _update_always_on_top(self):
        """更新窗口置顶状态"""
        if self.always_on_top.get():
            self.root.attributes('-topmost', True)
        else:
            self.root.attributes('-topmost', False)
            
    def _update_interval(self, event=None):
        """更新操作间隔设置"""
        try:
            value = float(self.interval_var.get())
            if value < 0.1:
                value = 0.1
            elif value > 10.0:
                value = 10.0
            self.auto_action_interval = value
            self.interval_var.set(str(value))
            self._log_info(f"已更新操作间隔为 {value} 秒", "blue")
        except ValueError:
            self.interval_var.set(str(self.auto_action_interval))
            self._log_info("操作间隔必须是有效的数字", "red")
    
    def _mode_changed(self, event=None):
        """处理模式切换"""
        mode = self.collection_mode.get()
        self._log_info(f"已切换到{mode}", "blue")
        
        # 根据不同模式更新UI和行为
        if mode == "正常模式":
            # 正常模式下的UI调整
            self.start_button.config(text="开始")
            
            # 显示Excel和Word导出按钮，隐藏JSON导出按钮
            self.excel_button.pack(side=tk.LEFT, padx=5, pady=5)
            self.word_button.pack(side=tk.LEFT, padx=5, pady=5)
            self.json_button.pack_forget()
            
        else:  # 采集模式
            # 采集模式下的UI调整
            self.start_button.config(text="开始采集")
            
            # 隐藏Excel和Word导出按钮，显示JSON导出按钮
            self.excel_button.pack_forget()
            self.word_button.pack_forget()
            self.json_button.pack(side=tk.LEFT, padx=5, pady=5)
    
    def _log_info(self, message, color=None):
        print(f"LOG: {message}")  # 强制输出到控制台，便于调试
        logger.info(message)
        # 添加到UI文本框
        self.status_text.config(state=tk.NORMAL)  # 临时设置为可写
        self.status_text.insert(tk.END, f"[{datetime.now().strftime('%H:%M:%S')}] {message}\n")
        # 应用颜色标签
        if color:
            last_line_start = self.status_text.index(f"end-2l")
            self.status_text.tag_add(f"color_{color}", last_line_start, "end-1c")
        self.status_text.see(tk.END)  # 滚动到最新内容
        self.status_text.config(state=tk.DISABLED)  # 恢复只读
    
    def _handle_key_event(self, event):
        key = event.char
        keysym = event.keysym
        if hasattr(self, '_pending_collect') and self._pending_collect:
            if not self.driver:
                self._log_info("浏览器未连接，无法采集", "red")
                messagebox.showerror("错误", "浏览器未连接，无法采集")
                self._pending_collect = None
                return
            try:
                self._inject_hover_listener()
                self.driver.switch_to.default_content()
                xpath = self._get_hovered_xpath_recursive()
                if not xpath:
                    self._log_info("未检测到悬停元素，请确保鼠标已悬停在目标元素上，且页面未刷新。", "orange")
                    messagebox.showerror("采集失败", "未能获取到元素信息，请确保鼠标悬停在浏览器页面内的有效元素上。")
                    self._pending_collect = None
                    return
                if self._pending_collect == 'ref1':
                    self.ref1_xpath = xpath
                    self.ref1_label.config(text=f"第1个订单参照点: {xpath}")
                    self._log_info(f"已采集第1个订单参照点XPath: {xpath}", "green")
                    messagebox.showinfo("采集成功", "第1个订单参照点采集成功！")
                elif self._pending_collect == 'ref2':
                    self.ref2_xpath = xpath
                    self.ref2_label.config(text=f"第2个订单参照点: {xpath}")
                    self._log_info(f"已采集第2个订单参照点XPath: {xpath}", "green")
                    messagebox.showinfo("采集成功", "第2个订单参照点采集成功！")
            except Exception as e:
                self._log_info(f"采集参照点异常: {str(e)}", "red")
                import traceback
                self._log_info(traceback.format_exc(), "red")
                messagebox.showerror("采集失败", f"采集过程中发生异常：{str(e)}")
            finally:
                self._pending_collect = None
            return
            
        # 处理.键：采集模式下采集，正常模式下暂停，翻页采集模式下采集翻页元素
        if key == '.':
            # 如果正在采集翻页元素
            if hasattr(self, 'collecting_page_turn') and self.collecting_page_turn:
                self._log_info("检测到.键，开始采集翻页元素...", "blue")
                self._handle_page_turn_collection()
                return
            elif self.collection_mode.get() == "采集模式":
                if self.is_running and not self.is_paused:
                    self._log_info("检测到.键，开始采集元素...", "blue")
                    self._collect_element()
                    return
            else:  # 正常模式
                if self.is_running and not self.is_paused:
                    self._log_info("检测到.键，暂停操作...", "orange")
                    self._pause_collection()
                    return
        
        # 处理-键：继续操作
        if key == '-':
            if self.is_running and self.is_paused:
                self._log_info("检测到-键，继续操作...", "green")
                self._continue_collection()
                return
        
        # 处理*键：终止操作并立即启用数据导出
        if key == '*':
            if self.is_running:
                self._log_info("检测到*键，终止操作...", "red")
                self._stop_collection()
                # 立即启用数据导出功能
                if hasattr(self, 'collected_data') and len(self.collected_data) > 0:
                    self._log_info("操作已终止，数据导出功能已启用", "green")
                    messagebox.showinfo("操作终止", "操作已终止，现在可以进行数据导出。")
                return
    
    def _start_browser(self):
        """重构：启动浏览器并通过远程调试端口连接，递归注入悬停监听脚本，断线重连，完全对齐代码逻辑.md"""
        # 检查Selenium依赖
        if webdriver is None:
            self._log_info("Selenium模块未正确导入，请安装: pip install selenium", "red")
            messagebox.showerror("依赖缺失", "Selenium模块未正确导入，请安装: pip install selenium")
            return
            
        self._log_info("正在启动浏览器...", "blue")
        self.connection_progress["value"] = 25

        def connect_browser():
            try:
                BROWSER_PATH = r"C:\\Program Files (x86)\\Microsoft\\Edge\\Application\\msedge.exe"
                DEBUG_PORT = str(self.debug_port)
                USER_DATA_DIR = self.user_data_dir
                if not os.path.exists(USER_DATA_DIR):
                    os.makedirs(USER_DATA_DIR)
                cmd = [BROWSER_PATH, f'--remote-debugging-port={DEBUG_PORT}', f'--user-data-dir={USER_DATA_DIR}']
                cmd_str = ' '.join(cmd)
                self._log_info(f"正在启动Edge浏览器: {cmd_str}", "blue")
                subprocess.Popen(cmd, creationflags=0x08000000)
                self._log_info('等待浏览器启动...')
                time.sleep(2)
                max_retries = 15
                driver = None
                for i in range(max_retries):
                    self._log_info(f'尝试连接到浏览器 (第 {i + 1}/{max_retries} 次)...')
                    try:
                        options = Options()
                        options.add_experimental_option('debuggerAddress', f'localhost:{DEBUG_PORT}')
                        driver = webdriver.Edge(options=options)
                        if driver.window_handles:
                            driver.switch_to.window(driver.window_handles[-1])
                            self.driver = driver
                            break
                    except Exception:
                        driver = None
                    time.sleep(1)
                if not driver:
                    raise ConnectionError('无法连接到浏览器。')
                self._log_info('成功连接到Edge浏览器。', "green")
                self.is_browser_connected = True
                self.browser_status.config(text="已连接", foreground="green")
                self.root.after(0, lambda: self.start_button.config(state=tk.NORMAL))
                self.root.after(0, lambda: self.configure_button.config(state=tk.NORMAL))
                # 启用翻页按钮
                self.root.after(0, lambda: self.collect_page_btn.config(state=tk.NORMAL))
                self._inject_hover_listener()
            except Exception as e:
                self.driver = None
                self._log_info(f'连接浏览器失败: {e}', "red")
                self.browser_status.config(text="连接失败", foreground="red")
                self.connection_progress["value"] = 0

        threading.Thread(target=connect_browser, daemon=True).start()

    def _inject_hover_listener(self):
        """递归注入悬停监听脚本到所有frame，便于采集鼠标悬停元素，对齐代码逻辑.md"""
        if not self.driver:
            return
        self._log_info('准备向所有框架（包括嵌套框架）注入悬停监听脚本...')
        try:
            self.driver.switch_to.default_content()
            self._inject_listener_recursive()
            self._log_info('悬停监听脚本注入完成。')
        except Exception as e:
            self._log_info(f'注入悬停监听脚本时发生主错误: {e}', 'red')
        finally:
            if self.driver:
                self.driver.switch_to.default_content()

    def _inject_listener_recursive(self):
        if not self.driver:
            self._log_info('无法注入监听脚本：浏览器未连接', 'red')
            return
        setup_script = '''
        if (window.pddToolListenerInjected) { return; }
        window.pddToolListenerInjected = true;
        window.lastHoveredElement = null;
        document.addEventListener('mouseover', function(e) {
            window.lastHoveredElement = e.target;
        }, true);
        '''
        self.driver.execute_script(setup_script)
        iframes = self.driver.find_elements(By.TAG_NAME, 'iframe')
        for i in range(len(iframes)):
            try:
                self.driver.switch_to.frame(i)
                self._inject_listener_recursive()
            except Exception as e:
                self._log_info(f'递归注入时无法切换到某个iframe: {e}', 'orange')
            finally:
                self.driver.switch_to.parent_frame()
    
    def _send_cdp_command(self, method, params={}):
        """发送CDP命令（远程调用）"""
        if self.ws is None:
            raise ValueError("No WebSocket connection established")
        payload = {
            "id": self.request_id,
            "method": method,
            "params": params
        }
        self.request_id += 1
        self.ws.send(json.dumps(payload))
        response = json.loads(self.ws.recv())
        return response.get("result", {})
    
    def _close_browser(self):
        """关闭浏览器连接和进程"""
        if self.driver:
            self.driver.quit()
            self.driver = None
        if self.ws:
            self.ws.close()
        if self.browser_process:
            self.browser_process.terminate()
        self.is_browser_connected = False
        self.browser_status.config(text="未连接", foreground="red")
        self._log_info("浏览器已关闭", "blue")
    
    def _execute_operation(self, operation):
        """重构：执行单个操作，采用pyautogui移动+WASD微调+剪贴板采集，支持用户验证，对齐代码逻辑.md"""
        import time  # 添加time模块导入，修复UnboundLocalError
        
        # 检查pyautogui依赖
        if 'pyautogui' not in globals():
            self._log_info("pyautogui模块未正确导入，请安装: pip install pyautogui", "red")
            return None
            
        try:
            if not self.driver:
                self._log_info("浏览器未连接，无法执行操作", "red")
                return None
            xpath = operation.get("smart_xpath") or operation["xpath"]
            action = operation["action"]
            name = operation["name"]
            
            # 使用智能定位查找元素
            element = self._find_element_smart(name, xpath)
                
            if not element:
                self._log_info(f"未找到元素: {name}", "red")
                return None
                
            # 获取元素特定的偏移量
            element_offset_x = 0
            element_offset_y = 0
            
            # 尝试从元素上获取名称属性，这是为了确保使用相对XPath等方法找到的元素也能正确应用偏移量
            try:
                element_name = self.driver.execute_script("return arguments[0].getAttribute('data-element-name');", element)
                if element_name and element_name in self.element_offsets:
                    self._log_info(f"使用元素'{element_name}'的偏移量配置", "blue")
                    element_offset_x = self.element_offsets[element_name].get("x", 0)
                    element_offset_y = self.element_offsets[element_name].get("y", 0)
                elif name in self.element_offsets:
                    self._log_info(f"使用元素'{name}'的偏移量配置", "blue")
                    element_offset_x = self.element_offsets[name].get("x", 0)
                    element_offset_y = self.element_offsets[name].get("y", 0)
                else:
                    self._log_info(f"元素'{name}'没有偏移量配置，使用默认值(0,0)", "blue")
            except Exception as e:
                # 如果获取属性失败，回退到使用操作名称
                if name in self.element_offsets:
                    element_offset_x = self.element_offsets[name].get("x", 0)
                    element_offset_y = self.element_offsets[name].get("y", 0)
                    self._log_info(f"从属性获取元素名称失败，使用操作名称'{name}'的偏移量: X={element_offset_x}, Y={element_offset_y}", "orange")
            
            if action == "getText":
                import re
                text = element.text.strip()
                self._log_info(f"获取文本 '{name}': {text}", "green")
                
                # 如果是订单编号元素，解析并保存订单ID
                if name == "订单编号" or "订单编号" in name:
                    match = re.search(r"订单编号[：: ]*([0-9a-zA-Z\-]+)", text)
                    if match:
                        self.last_captured_order_id = match.group(1)
                        self._log_info(f"已保存订单编号: {self.last_captured_order_id}", "blue")
                    else:
                        self.last_captured_order_id = None
                        self._log_info("未能解析订单编号", "red")
                return text
            elif action in ["click", "clickAndGetClipboard"]:
                # 滚动到元素可见
                self.driver.execute_script('arguments[0].scrollIntoView({behavior: "smooth", block: "center"});', element)
                time.sleep(0.5)
                
                # 获取浏览器窗口位置和尺寸
                window_pos = self.driver.get_window_position()
                window_size = self.driver.get_window_size()
                
                # 获取元素在视口中的位置和尺寸
                rect = self.driver.execute_script('return arguments[0].getBoundingClientRect();', element)
                
                # 记录元素的原始位置信息
                self._log_info(f"元素'{name}'的原始位置: left={rect['left']}, top={rect['top']}, width={rect['width']}, height={rect['height']}", "blue")
                
                # 获取浏览器内容区域的偏移量
                js = """
                return {
                    screenX: window.screenX, 
                    screenY: window.screenY, 
                    outerHeight: window.outerHeight, 
                    innerHeight: window.innerHeight,
                    outerWidth: window.outerWidth,
                    innerWidth: window.innerWidth
                };
                """
                win_metrics = self.driver.execute_script(js)
                
                # 计算内容区域的左上角位置
                content_left = win_metrics['screenX'] + (win_metrics['outerWidth'] - win_metrics['innerWidth']) / 2
                content_top = win_metrics['screenY'] + (win_metrics['outerHeight'] - win_metrics['innerHeight'])
                
                self._log_info(f"浏览器内容区域: left={content_left}, top={content_top}", "blue")
                
                # 计算元素的绝对屏幕位置（取元素中心点）
                screen_x = content_left + rect['left'] + rect['width'] / 2
                screen_y = content_top + rect['top'] + rect['height'] / 2
                
                self._log_info(f"元素'{name}'的中心点位置: X={screen_x}, Y={screen_y}", "blue")
                
                # 检查元素位置是否合理
                if rect['left'] < 0 or rect['top'] < 0 or rect['width'] <= 0 or rect['height'] <= 0:
                    self._log_info(f"警告: 元素'{name}'的位置或尺寸不合理，尝试重新获取", "orange")
                    
                    # 尝试使用JavaScript重新获取元素位置
                    try:
                        js_rect = """
                        var element = arguments[0];
                        var rect = element.getBoundingClientRect();
                        
                        // 确保元素可见
                        element.scrollIntoView({behavior: 'auto', block: 'center'});
                        
                        // 等待一下让滚动完成
                        return new Promise(resolve => {
                            setTimeout(() => {
                                var updatedRect = element.getBoundingClientRect();
                                resolve({
                                    left: updatedRect.left,
                                    top: updatedRect.top,
                                    width: updatedRect.width,
                                    height: updatedRect.height
                                });
                            }, 500);
                        });
                        """
                        
                        updated_rect = self.driver.execute_script(js_rect, element)
                        self._log_info(f"重新获取元素'{name}'的位置: left={updated_rect['left']}, top={updated_rect['top']}, width={updated_rect['width']}, height={updated_rect['height']}", "blue")
                        
                        # 使用更新后的位置
                        if updated_rect['width'] > 0 and updated_rect['height'] > 0:
                            rect = updated_rect
                            # 重新计算元素的绝对屏幕位置
                            screen_x = content_left + rect['left'] + rect['width'] / 2
                            screen_y = content_top + rect['top'] + rect['height'] / 2
                            self._log_info(f"更新后元素'{name}'的中心点位置: X={screen_x}, Y={screen_y}", "blue")
                    except Exception as e:
                        self._log_info(f"重新获取元素位置失败: {str(e)}", "red")
                
                # 应用元素特定的偏移量
                screen_x += element_offset_x
                screen_y += element_offset_y
                
                # 记录应用的偏移量
                self._log_info(f"应用元素'{name}'的偏移量: X={element_offset_x}, Y={element_offset_y}", "blue")
                self._log_info(f"最终点击位置: X={screen_x}, Y={screen_y}", "blue")
                
                # 检查点击位置是否在屏幕范围内
                screen_width, screen_height = pyautogui.size()
                if screen_x < 0 or screen_x > screen_width or screen_y < 0 or screen_y > screen_height:
                    self._log_info(f"警告: 计算的点击位置({screen_x}, {screen_y})超出屏幕范围({screen_width}x{screen_height})", "red")
                    # 调整到屏幕范围内
                    screen_x = max(0, min(screen_x, screen_width))
                    screen_y = max(0, min(screen_y, screen_height))
                    self._log_info(f"调整后的点击位置: X={screen_x}, Y={screen_y}", "orange")
                
                # 支持WASD微调
                if self.confirm_click.get():
                    ok = self._show_verification_dialog_and_wait(name, xpath, screen_x, screen_y, element_offset_x, element_offset_y)
                    if not ok:
                        self._log_info(f"用户取消了点击: {name}", "orange")
                        return None
                    # 在_show_verification_dialog_and_wait中已经执行了点击，直接返回
                    return True
                else:
                    # 没有确认对话框，直接在应用了元素偏移量的位置点击
                    
                    # 点击前确保浏览器窗口有焦点
                    self._switch_focus_to_browser()
                    time.sleep(0.3)
                    
                    pyautogui.moveTo(int(screen_x), int(screen_y))
                    pyautogui.click()
                    self._log_info(f"已通过 PyAutoGUI 点击 '{name}'", "blue")
                    
                    # 点击后延迟一下，让浏览器有时间响应
                    time.sleep(0.5)
                    
                    # 对于'复制完整的收货信息'元素，跳过额外点击以避免剪贴板内容重复
                    if name != '复制完整的收货信息':
                        # 执行一次额外的原地点击，与点击前确认模式行为一致
                        pyautogui.click()
                        self._log_info(f"已执行额外的原地点击 '{name}'", "blue")
                        time.sleep(0.3)
                    else:
                        self._log_info(f"跳过'{name}'的额外点击，避免剪贴板内容重复", "blue")
                    
                    # 恢复焦点到采集工具窗口
                    self._manage_focus()
                    
                # 统一处理clickAndGetClipboard动作
                if action == "clickAndGetClipboard":
                     # 对于'复制完整的收货信息'元素，如果跳过了额外点击，直接获取剪贴板内容
                     if name == '复制完整的收货信息':
                         self._log_info(f"直接获取剪贴板内容，无需等待更新", "blue")
                         time.sleep(0.5)
                         self._manage_focus()
                         clipboard_content = pyperclip.paste()
                     else:
                         self._log_info(f"点击后等待剪贴板内容更新...", "blue")
                         time.sleep(1.5)
                         self._manage_focus()
                         time.sleep(0.5)
                         clipboard_content = self._wait_for_clipboard_content(
                             timeout=12.0,
                             check_interval=0.5,
                             min_length=10
                         )
                     # 使用之前采集的订单ID
                     current_order_id = getattr(self, 'last_captured_order_id', None)
                     if not current_order_id:
                         # 弹窗要求用户输入订单ID
                         from tkinter import simpledialog
                         current_order_id = simpledialog.askstring(
                             "订单ID缺失", "未能自动提取订单ID，请手动输入当前订单ID：", parent=self.root)
                         if not current_order_id or not current_order_id.strip():
                             self._log_info("用户未输入订单ID，跳过本次映射", "red")
                             return clipboard_content
                     if clipboard_content and clipboard_content.strip():
                         self._store_clipboard_content(name, clipboard_content, current_order_id)
                         self._log_info(f"[映射写入] 订单ID: {current_order_id}, 长度: {len(clipboard_content)}, 内容: {clipboard_content[:30]}...", "green")
                         self._save_clipboard_mappings()
                     else:
                         self._log_info(f"[映射跳过] 订单ID: {current_order_id}, 内容无效或为空", "orange")
                     return clipboard_content
                    
        except Exception as e:
            self._log_info(f"执行'{name}'操作失败: {str(e)}", "red")
            import traceback
            self._log_info(traceback.format_exc(), "red")
            return None

    def _show_verification_dialog_and_wait(self, remark, xpath, screen_x, screen_y, offset_x, offset_y):
        """弹窗让用户验证鼠标位置，支持WASD微调，保存元素特定偏移量"""
        self._last_offset_x = offset_x  # 初始化为当前元素的偏移量
        self._last_offset_y = offset_y
        result = {'ok': -1}
        
        def show_dialog():
            win = tk.Toplevel(self.root)
            win.title('请验证鼠标位置')
            win.geometry('350x220')
            win.transient(self.root)
            win.grab_set()
            msg = f'已将鼠标移动到 "{remark}" 的目标位置。\n\n请检查屏幕上的鼠标指针是否准确。\n\nA/D键: 左/右移1像素，W/S键: 上/下移1像素。\n点击"位置准确"后将保存此元素的偏移量并执行点击。\n点击"位置不准"将跳过本次点击并记录日志。'
            label = ttk.Label(win, text=msg, wraplength=320, justify=tk.LEFT)
            label.pack(pady=10, padx=10)
            offset_label = ttk.Label(win, text=f'当前偏移: X={self._last_offset_x}, Y={self._last_offset_y}')
            offset_label.pack(pady=2)
            
            def update_offset_label():
                offset_label.config(text=f'当前偏移: X={self._last_offset_x}, Y={self._last_offset_y}')
            
            def on_ok():
                result['ok'] = 1
                
                # 保存当前偏移量为元素特定的偏移量
                if remark not in self.element_offsets:
                    self.element_offsets[remark] = {}
                self.element_offsets[remark]["x"] = self._last_offset_x
                self.element_offsets[remark]["y"] = self._last_offset_y
                
                # 在当前位置执行点击
                pyautogui.click()
                self._log_info(f"已保存元素'{remark}'的偏移量 X={self._last_offset_x}, Y={self._last_offset_y} 并执行点击", "green")
                
                # 保存偏移量到配置文件
                self._save_offset_config()
                win.destroy()
                # --- 焦点恢复 ---
                self._manage_focus()
                self.root.after(100, self._manage_focus)
                self._log_info("已尝试恢复主窗口焦点", "blue")
                
            def on_fail():
                result['ok'] = 0
                win.destroy()
                # --- 焦点恢复 ---
                self._manage_focus()
                self.root.after(100, self._manage_focus)
                self._log_info("已尝试恢复主窗口焦点", "blue")
                
            btn_frame = ttk.Frame(win)
            btn_frame.pack(pady=10)
            ok_btn = ttk.Button(btn_frame, text='位置准确', command=on_ok)
            ok_btn.pack(side=tk.LEFT, padx=10)
            fail_btn = ttk.Button(btn_frame, text='位置不准', command=on_fail)
            fail_btn.pack(side=tk.LEFT, padx=10)
            win.bind('<Return>', lambda event: on_ok())
            win.bind('<Escape>', lambda event: on_fail())
            
            def on_key(event):
                moved = False
                if event.keysym.lower() == 'a':
                    self._last_offset_x -= 1
                    pyautogui.moveRel(-1, 0, duration=0)
                    moved = True
                elif event.keysym.lower() == 'd':
                    self._last_offset_x += 1
                    pyautogui.moveRel(1, 0, duration=0)
                    moved = True
                elif event.keysym.lower() == 'w':
                    self._last_offset_y -= 1
                    pyautogui.moveRel(0, -1, duration=0)
                    moved = True
                elif event.keysym.lower() == 's':
                    self._last_offset_y += 1
                    pyautogui.moveRel(0, 1, duration=0)
                    moved = True
                if moved:
                    update_offset_label()
                    
            win.bind('<Key>', on_key)
            
            # 移动鼠标到初始位置
            pyautogui.moveTo(int(screen_x), int(screen_y))
            win.focus_set()
            
        self.root.after(0, show_dialog)
        # 阻塞等待用户操作
        while result['ok'] == -1:
            self.root.update()
            time.sleep(0.05)
        
        return result['ok'] == 1
    
    def _process_orders(self):
        """智能循环主流程：动态XPath采集+智能滚动+订单编号去重"""
        # 检查pyperclip依赖
        if pyperclip is None:
            self._log_info("pyperclip模块未正确导入，请安装: pip install pyperclip", "red")
            self._stop_collection()
            return
            
        try:
            total_orders = self._get_order_count()
            if not total_orders:
                self._log_info("无法获取订单数量，操作终止", "red")
                self._stop_collection()
                return
            self._log_info(f"检测到 {total_orders} 个待处理订单", "blue")
            self.progress_bar["maximum"] = total_orders
            self.progress_bar["value"] = 0
            self.progress_label.config(text=f"0/{total_orders}")
            self.collected_data = []
            if self.ref1_xpath and self.ref2_xpath:
                pattern = self._learn_xpath_pattern(self.ref1_xpath, self.ref2_xpath)
                if not pattern:
                    self._log_info("未能推算出循环XPath规律，回退到普通模式", "red")
                    return self._process_orders_fallback()
                self._log_info(f"已推算出循环XPath模板: {pattern['template']}", "green")
                prev_order_id = None
                scroll_fail_count = 0
                for i in range(1, total_orders + 1):
                    if not self.is_running or self.is_paused:
                        break
                    self._log_info(f"[智能循环] 正在处理第 {i}/{total_orders} 个订单", "blue")
                    order_data = {}
                    # 先定位到本订单参照元素并滚动到可见
                    try:
                        item_xpath = self._generate_xpath_for_index(pattern, i)
                        element = self.driver.find_element(By.XPATH, item_xpath)
                        self.driver.execute_script('arguments[0].scrollIntoView({behavior: "smooth", block: "center"});', element)
                        time.sleep(0.3)
                    except Exception as e:
                        self._log_info(f"滚动到第{i}个订单参照点失败: {str(e)}", "red")
                    # 获取当前订单ID
                    order_id = self._extract_current_order_id()
                    if order_id:
                        self._log_info(f"已提取到当前订单ID: {order_id}", "green")
                    else:
                        self._log_info(f"未能提取到当前订单ID，将在处理过程中获取", "orange")
                        order_id = None
                        
                    for op in self.operation_sequence:
                        try:
                            op_xpath = op.get('xpath', '')
                            op_pattern = self._learn_xpath_pattern(self.ref1_xpath, op_xpath) if op_xpath else None
                            if op_pattern:
                                op_item_xpath = self._generate_xpath_for_index(op_pattern, i)
                            else:
                                op_item_xpath = op_xpath
                            op_copy = op.copy()
                            op_copy['xpath'] = op_item_xpath
                            result = self._execute_operation(op_copy)
                            order_data[op['name']] = result
                            
                            # 如果"点击前确认"未勾选，每个操作后添加延迟
                            if not self.confirm_click.get():
                                time.sleep(self.auto_action_interval)
                                self._log_info(f"自动执行模式：操作间延迟 {self.auto_action_interval} 秒", "blue")

                            # 每次采集到关键字段后，立即尝试回填映射
                            order_id_for_map = order_data.get('订单编号')
                            shipping_info_for_map = order_data.get('复制完整收货信息')
                            if order_id_for_map and shipping_info_for_map:
                                clean_order_id = order_id_for_map.replace("订单编号：", "") if isinstance(order_id_for_map, str) else str(order_id_for_map)
                                self.order_clipboard_contents[clean_order_id] = shipping_info_for_map
                                print(f"DEBUG-MAP-ORDER-ID: '{clean_order_id}' -> 收货信息长度: {len(shipping_info_for_map)}")
                                print(f"DEBUG-MAP-SHIPPING-INFO-START: '{shipping_info_for_map[:50]}'")
                                self._log_info(f"已将收货信息与订单ID {clean_order_id} 关联", "blue")
                        except Exception as e:
                            self._log_info(f"执行'{op['name']}'操作失败: {str(e)}", "red")
                            import traceback
                            self._log_info(traceback.format_exc(), "red")
                    # 采集完所有字段后再兜底检查一次
                    order_id_for_map = order_data.get('订单编号')
                    shipping_info_for_map = order_data.get('复制完整收货信息')
                    
                    if order_id_for_map and shipping_info_for_map:
                        clean_order_id = order_id_for_map.replace("订单编号：", "") if isinstance(order_id_for_map, str) else str(order_id_for_map)
                        
                        # 检查是否已有映射
                        if clean_order_id in self.order_clipboard_contents:
                            existing_content = self.order_clipboard_contents[clean_order_id]
                            # 比较内容长度，保留更长的内容
                            if isinstance(shipping_info_for_map, str) and len(shipping_info_for_map) > len(existing_content):
                                self.order_clipboard_contents[clean_order_id] = shipping_info_for_map
                                print(f"DEBUG-UPDATE-MAP-END: 订单ID '{clean_order_id}' -> 更新为更长的收货信息 ({len(shipping_info_for_map)} > {len(existing_content)})")
                                self._log_info(f"采集结束后更新映射: 订单ID {clean_order_id} <-> 收货信息 (新长度: {len(shipping_info_for_map)})", "blue")
                            else:
                                print(f"DEBUG-KEEP-MAP-END: 订单ID '{clean_order_id}' -> 保留现有收货信息 ({len(existing_content)} 字符)")
                                self._log_info(f"保留现有映射: 订单ID {clean_order_id}", "blue")
                        else:
                            # 创建新映射
                            self.order_clipboard_contents[clean_order_id] = shipping_info_for_map
                            print(f"DEBUG-NEW-MAP-END: 订单ID '{clean_order_id}' -> 新建收货信息映射 ({len(shipping_info_for_map)} 字符)")
                            self._log_info(f"采集结束后补充映射: 订单ID {clean_order_id} <-> 收货信息", "blue")
                    elif not (order_id_for_map and shipping_info_for_map):
                        print(f"DEBUG-MISSING-MAP-END: 订单ID='{order_id_for_map}', 收货信息={(shipping_info_for_map is not None)}")
                        self._log_info("订单ID或收货信息缺失，无法建立映射", "orange")
                        
                    # 特殊处理最后一个订单
                    if i == total_orders:  # 最后一个订单
                        self._log_info("处理最后一个订单，确保收货信息正确映射", "blue")
                        # 强制等待更长时间，确保剪贴板内容更新
                        time.sleep(3.0)
                        
                        # 再次尝试获取剪贴板内容
                        import pyperclip
                        final_clipboard = pyperclip.paste()
                        if final_clipboard and final_clipboard.strip():
                            if order_id_for_map:
                                clean_order_id = order_id_for_map.replace("订单编号：", "") if isinstance(order_id_for_map, str) else str(order_id_for_map)
                                self.order_clipboard_contents[clean_order_id] = final_clipboard
                                print(f"DEBUG-FINAL-ORDER: 订单ID '{clean_order_id}' -> 最终剪贴板内容长度: {len(final_clipboard)}")
                                self._log_info(f"最后一个订单特殊处理: 订单ID {clean_order_id} -> 收货信息长度: {len(final_clipboard)}", "green")
                    
                    if i > 1 and prev_order_id and order_id == prev_order_id:
                        self._log_info("检测到订单未切换，尝试加大滚动距离", "orange")
                        scroll_fail_count += 1
                        # 优先滚动容器，否则全局scrollBy
                        try:
                            if self.scroll_container_xpath:
                                script = f'''var el = document.evaluate("{self.scroll_container_xpath}", document, null, XPathResult.FIRST_ORDERED_NODE_TYPE, null).singleNodeValue; if (el) {{ el.scrollTop += {120*scroll_fail_count}; return el.scrollTop; }} else {{ window.scrollBy(0, {120*scroll_fail_count}); return window.scrollY; }}'''
                                self.driver.execute_script(script)
                                self._log_info(f"[智能滚动] 滚动容器 {self.scroll_container_xpath} +{120*scroll_fail_count}px", "blue")
                            else:
                                self.driver.execute_script(f'window.scrollBy(0, {120*scroll_fail_count});')
                                self._log_info(f"[全局滚动] window.scrollBy(0, {120*scroll_fail_count})", "blue")
                        except Exception as e:
                            self._log_info(f"智能滚动失败: {str(e)}", "red")
                        time.sleep(1)
                        continue
                    prev_order_id = order_id if 'order_id' in locals() else None
                    scroll_fail_count = 0
                    if order_data:
                        self.collected_data.append(order_data)
                    
                    # 设置当前订单索引用于翻页检查
                    self.current_order_index = i - 1  # 转换为0基索引
                    
                    # 检查是否需要翻页
                    if self.check_page_turn_needed():
                        self._log_info(f"已处理 {self.target_page_count} 个订单，开始翻页...", "blue")
                        if self.execute_page_turn():
                            self._log_info("翻页成功，继续处理订单", "green")
                            # 翻页后等待页面稳定
                            time.sleep(2)
                        else:
                            self._log_info("翻页失败，停止处理", "error")
                            break
                    
                    self.progress_bar["value"] = i
                    self.progress_label.config(text=f"{i}/{total_orders}")
                self._log_info(f"[智能循环] 已完成所有 {len(self.collected_data)} 个订单的处理", "green")
                
                # 循环结束后验证订单ID与收货信息映射
                self._log_info("循环结束后验证订单ID与收货信息映射:", "blue")
                
                # 确保order_clipboard_contents字典存在
                if not hasattr(self, 'order_clipboard_contents'):
                    self.order_clipboard_contents = {}
                    self._log_info("警告: 订单ID与收货信息映射字典不存在，已创建空字典", "red")
                    print("DEBUG-FINAL-ERROR: 订单ID与收货信息映射字典不存在，已创建空字典")
                
                print(f"DEBUG-FINAL-MAP-COUNT: 收集到 {len(self.collected_data)} 个订单，映射字典中有 {len(self.order_clipboard_contents)} 个映射")
                
                # 检查收集的订单数据与映射字典的一致性
                missing_mappings = []
                
                for i, order_data in enumerate(self.collected_data):
                    order_id = order_data.get('订单编号', '')
                    if isinstance(order_id, str) and order_id.startswith('订单编号：'):
                        clean_order_id = order_id.replace('订单编号：', '')
                    else:
                        clean_order_id = str(order_id)
                        
                    shipping_info = self.order_clipboard_contents.get(clean_order_id, "未找到")
                    if shipping_info == "未找到":
                        missing_mappings.append(clean_order_id)
                        
                    shipping_info_preview = shipping_info[:30] + "..." if shipping_info and shipping_info != "未找到" else shipping_info
                    print(f"DEBUG-FINAL-MAP-{i+1}: 订单ID '{clean_order_id}' -> 收货信息: '{shipping_info_preview}'")
                    self._log_info(f"订单{i+1} ID: {clean_order_id}, 收货信息: {shipping_info_preview}", "blue")
                
                # 报告缺失的映射
                if missing_mappings:
                    self._log_info(f"警告: 有 {len(missing_mappings)} 个订单缺少收货信息映射", "red")
                    print(f"DEBUG-FINAL-MISSING: 缺少映射的订单ID: {missing_mappings}")
                else:
                    self._log_info("所有订单都有对应的收货信息映射", "green")
                    print("DEBUG-FINAL-COMPLETE: 所有订单都有对应的收货信息映射")
                
                self._stop_collection()
                if len(self.collected_data) > 0:
                    self.excel_button.config(state=tk.NORMAL)
                    self.word_button.config(state=tk.NORMAL)
                return
            return self._process_orders_fallback()
        except Exception as e:
            self._log_info(f"订单处理过程中出错: {str(e)}", "red")
            import traceback
            self._log_info(traceback.format_exc(), "red")
            self._stop_collection()

    def _process_orders_fallback(self):
        """原有无参照点时的采集逻辑（兜底）"""
        try:
            total_orders = self._get_order_count()
            if not total_orders:
                self._log_info("无法获取订单数量，操作终止", "red")
                self._stop_collection()
                return
            self._log_info(f"检测到 {total_orders} 个待处理订单", "blue")
            self.progress_bar["maximum"] = total_orders
            self.progress_bar["value"] = 0
            self.progress_label.config(text=f"0/{total_orders}")
            self.collected_data = []
            once_executed_operations = set()
            prev_order_id = None
            scroll_fail_count = 0
            
            for i in range(total_orders):
                if not self.is_running or self.is_paused:
                    break
                self._log_info(f"正在处理第 {i + 1}/{total_orders} 个订单", "blue")
                order_data = {}
                if self.confirm_click.get() and i > 2:
                    self._log_info("已进入鼠标确认调试模式，仅循环前三个订单", "orange")
                    break
                if i + 1 == total_orders:
                    self._log_info("已进入最后一个订单，准备采集", "green")
                # 采集订单编号
                try:
                    # 优先使用智能提取方法获取订单ID
                    order_id = self._extract_current_order_id()
                    if order_id:
                        order_data['订单编号'] = f"订单编号：{order_id}"
                        self._log_info(f"智能提取到订单编号: {order_id}", "green")
                    else:
                        # 尝试使用旧方法
                        try:
                            element = self.driver.find_element(By.XPATH, f'//*[@id="order_id_{i+1}"]')
                            order_id = element.text.strip()
                            order_data['订单编号'] = order_id
                            self._log_info(f"采集到订单编号: {order_id}", "green")
                        except Exception as inner_e:
                            self._log_info(f"采集订单编号失败: {str(inner_e)}", "red")
                            order_id = f"未知_{len(self.collected_data)+1}"
                            order_data['订单编号'] = f"订单编号：{order_id}"
                            self._log_info(f"使用临时订单ID: {order_id}", "orange")
                except Exception as e:
                    self._log_info(f"订单ID提取失败: {str(e)}", "red")
                    order_id = f"未知_{len(self.collected_data)+1}"
                    order_data['订单编号'] = f"订单编号：{order_id}"
                    self._log_info(f"使用临时订单ID: {order_id}", "orange")
                # 去重校验
                if prev_order_id and order_id == prev_order_id:
                    self._log_info("检测到订单未切换，尝试加大滚动距离", "orange")
                    scroll_fail_count += 1
                    self._smart_scroll(step_mul=1+scroll_fail_count)
                    time.sleep(1)
                    continue
                prev_order_id = order_id
                scroll_fail_count = 0
                # 采集其他字段
                for op in self.operation_sequence:
                    if not self.is_running or self.is_paused:
                        break
                    if op.get('name') == '订单编号':
                        order_data[op['name']] = order_id
                        continue
                    # 动态生成同类元素XPath
                    op_xpath = op.get('xpath', '')
                    op_pattern = self._learn_xpath_pattern(self.ref1_xpath, op_xpath) if op_xpath else None
                    if op_pattern:
                        op_item_xpath = self._generate_xpath_for_item(op_xpath, i+1, op_pattern)
                    else:
                        op_item_xpath = op_xpath
                    op_copy = op.copy()
                    op_copy['xpath'] = op_item_xpath
                    result = self._execute_operation(op_copy)
                    # 修复：对于clickAndGetClipboard动作，确保使用最佳的剪贴板内容
                    if op_copy.get('action') == 'clickAndGetClipboard':
                        # 优先使用execute_operation返回的结果
                        if result and isinstance(result, str) and result.strip().lower() not in ["true", "false", ""]:
                            order_data[op['name']] = result
                            self._log_info(f"使用操作返回的剪贴板内容: '{result[:50]}...'", "blue")
                        # 其次尝试使用全局保存的最后剪贴板内容
                        elif hasattr(self, 'last_clipboard_content') and self.last_clipboard_content:
                            order_data[op['name']] = self.last_clipboard_content
                            self._log_info(f"使用全局保存的剪贴板内容: '{self.last_clipboard_content[:50]}...'", "blue")
                        # 最后再尝试直接读取剪贴板
                        else:
                            import pyperclip
                            clipboard_content = pyperclip.paste()
                            if clipboard_content and clipboard_content.strip().lower() not in ["true", "false", ""]:
                                order_data[op['name']] = clipboard_content
                                self._log_info(f"使用直接读取的剪贴板内容: '{clipboard_content[:50]}...'", "blue")
                            else:
                                # 如果都失败了，保存result
                                order_data[op['name']] = result
                                self._log_info(f"所有方法都失败，使用原始结果: {result}", "orange")
                    elif result is not None:
                        order_data[op['name']] = result
                    
                    # 如果"点击前确认"未勾选，使用自定义间隔；否则使用原有的短延迟
                    if not self.confirm_click.get():
                        time.sleep(self.auto_action_interval)
                        self._log_info(f"自动执行模式：操作间延迟 {self.auto_action_interval} 秒", "blue")
                    else:
                        time.sleep(0.2)
                if order_data:
                    self.collected_data.append(order_data)
                
                # 设置当前订单索引用于翻页检查
                self.current_order_index = i
                
                # 检查是否需要翻页
                if self.check_page_turn_needed():
                    self._log_info(f"已处理 {self.target_page_count} 个订单，开始翻页...", "blue")
                    if self.execute_page_turn():
                        self._log_info("翻页成功，继续处理订单", "green")
                        # 翻页后等待页面稳定
                        time.sleep(2)
                    else:
                        self._log_info("翻页失败，停止处理", "error")
                        break
                
                self.progress_bar["value"] = i + 1
                self.progress_label.config(text=f"{i + 1}/{total_orders}")
                # 滚动到下一个订单
                if i + 1 < total_orders:
                    self._scroll_to_next_order()
                    time.sleep(0.5)
            
                self._log_info(f"已完成所有 {len(self.collected_data)} 个订单的处理", "green")
                self._stop_collection()
                if len(self.collected_data) > 0:
                    self.excel_button.config(state=tk.NORMAL)
                    self.word_button.config(state=tk.NORMAL)
            return
        except Exception as e:
            self._log_info(f"订单处理过程中出错: {str(e)}", "red")
            import traceback
            self._log_info(traceback.format_exc(), "red")
            self._stop_collection()
    
    def _start_collection(self):
        # 启动验证码检测（如果已配置）
        if hasattr(self, 'template_images') and (self.template_images or self.use_mask_detection):
            if hasattr(self, 'target_window_handle') and self.target_window_handle:
                self._start_captcha_detection()
                self._log_info("验证码检测已自动启动", "green")
            else:
                self._log_info("提示：未选择验证码检测目标窗口，验证码检测未启动", "orange")
        
        mode = self.collection_mode.get()
        if mode == "正常模式":
            if not hasattr(self, 'operation_sequence') or not self.operation_sequence:
                self._configure_operations()
                return
                
            # 确保order_clipboard_contents字典已初始化
            if not hasattr(self, 'order_clipboard_contents'):
                self.order_clipboard_contents = {}
                self._log_info("初始化订单ID与收货信息映射字典", "blue")
                print("DEBUG-START-COLLECTION: 初始化订单ID与收货信息映射字典")
            
            # 程序启动时已清空所有历史数据，无需再次清理
                
            # 清空并记录初始剪贴板状态
            self.initial_clipboard_content = self._get_clipboard_content()
            self._log_info(f"记录初始剪贴板状态，长度: {len(self.initial_clipboard_content) if self.initial_clipboard_content else 0}", "blue")
            
            # 启动剪贴板监听
            self._start_clipboard_monitor()
            
            # 重置订单ID缓存
            self.current_order_id = None
            
            # 在主线程中预先获取订单数量（如果需要手动输入）
            order_count_elements = [op for op in self.operation_sequence if op.get("is_order_count", False)]
            manual_order_count = None
            
            if not order_count_elements:
                # 没有订单数量元素，需要手动输入订单数量
                self._log_info("未选择订单数量元素，将手动输入订单数量", "blue")
                try:
                    manual_order_count = self._get_order_count()
                    if manual_order_count == 0:
                        self._log_info('订单数量为0，无需执行。')
                        return
                    self._log_info(f"手动输入的订单数量: {manual_order_count}", "blue")
                except Exception as e:
                    self._log_info(f"获取订单数量失败: {str(e)}", "red")
                    return
            
            self._log_info("开始自动采集流程...", "green")
            self.is_running = True
            self.is_paused = False
            self.start_button.config(state=tk.DISABLED)
            self.pause_button.config(state=tk.NORMAL)
            self.continue_button.config(state=tk.DISABLED)
            self.stop_button.config(state=tk.NORMAL)
            self.configure_button.config(state=tk.DISABLED)
            threading.Thread(target=self.run_actions_loop, args=(manual_order_count,), daemon=True).start()
        elif mode == "采集模式":
            # 确保order_clipboard_contents字典已初始化
            if not hasattr(self, 'order_clipboard_contents'):
                self.order_clipboard_contents = {}
                self._log_info("初始化订单ID与收货信息映射字典", "blue")
                print("DEBUG-START-COLLECTION-MODE: 初始化订单ID与收货信息映射字典")
            
            # 启动剪贴板监听
            self._start_clipboard_monitor()
            
            # 重置订单ID缓存
            self.current_order_id = None
            
            self._log_info("进入采集准备状态，请按'.'键采集当前元素...", "green")
            self.is_running = True
            self.is_paused = False
            self.start_button.config(state=tk.DISABLED)
            self.pause_button.config(state=tk.NORMAL)
            self.continue_button.config(state=tk.DISABLED)
            self.stop_button.config(state=tk.NORMAL)
            self.configure_button.config(state=tk.DISABLED)
    
    def _pause_collection(self):
        """暂停采集流程"""
        self._log_info("已暂停采集", "orange")
        self.is_paused = True
        
        # 更新按钮状态
        self.pause_button.config(state=tk.DISABLED)
        self.continue_button.config(state=tk.NORMAL)
    
    def _continue_collection(self):
        """继续采集流程"""
        self._log_info("继续采集", "green")
        self.is_paused = False
        
        # 更新按钮状态
        self.pause_button.config(state=tk.NORMAL)
        self.continue_button.config(state=tk.DISABLED)
    
    def _stop_collection(self):
        if not self.is_running:
            return
        self._log_info("已终止采集", "red")
        self.is_running = False
        self.is_paused = False
        
        # 停止验证码检测
        if hasattr(self, 'captcha_running') and self.captcha_running:
            self._stop_captcha_detection()
            self._log_info("验证码检测已停止", "orange")
        
        # 停止剪贴板监听
        self._stop_clipboard_monitor()
        
        # 保存当前的剪贴板映射
        self._save_clipboard_mappings()
        
        # 删除辅助定位相关状态重置
        # 更新按钮状态
        self.start_button.config(state=tk.NORMAL)
        self.pause_button.config(state=tk.DISABLED)
        self.continue_button.config(state=tk.DISABLED)
        self.stop_button.config(state=tk.DISABLED)
        self.configure_button.config(state=tk.NORMAL if self.is_browser_connected else tk.DISABLED)
        # 如果有数据，根据模式启用不同导出按钮
        if hasattr(self, 'collected_data') and len(self.collected_data) > 0:
            if self.collection_mode.get() == "正常模式":
                self.excel_button.config(state=tk.NORMAL)
                self.word_button.config(state=tk.NORMAL)
            else:
                self.json_button.config(state=tk.NORMAL)
    
    def _collect_element(self):
        """重构：采集当前悬停元素的XPath，递归支持iframe，采集后清空lastHoveredElement，对齐代码逻辑.md。增强鲁棒性。修复置顶窗口导致焦点无法切换到浏览器的问题。"""
        if self.is_distance_learning:
            self._process_distance_learning()
            return
        if not self.driver:
            self._log_info("浏览器未连接，无法采集", "red")
            return
        # --- 修复：采集前临时取消置顶，采集后恢复 ---
        main_window = self.root.winfo_toplevel()
        was_topmost = main_window.attributes('-topmost')
        main_window.attributes('-topmost', False)
        try:
            # 每次采集前都注入监听脚本，防止页面刷新或iframe变化导致失效
            self._inject_hover_listener()
            # 切回主frame
            self.driver.switch_to.default_content()
            xpath = self._get_hovered_xpath_recursive()
            if not xpath:
                # 自动重试一次监听脚本注入+采集
                self._log_info("未检测到悬停元素，自动重试注入监听脚本...", "orange")
                self._inject_hover_listener()
                self.driver.switch_to.default_content()
                xpath = self._get_hovered_xpath_recursive()
            if not xpath:
                self._log_info("未检测到悬停元素，请确保鼠标已悬停在目标元素上，且页面未刷新。", "orange")
                return
            # 获取元素文本并高亮显示
            element = self.driver.find_element(By.XPATH, xpath)
            full_text = element.text.strip()
            chinese_text = ''.join(re.findall(r'[一-龥]+', full_text))
            
            # 高亮显示采集到的元素
            self._highlight_element(xpath)
            # 弹窗选择动作
            dialog = ActionSelectionDialog(self.root)
            action = dialog.result
            if action is None:
                self._log_info("已取消采集", "orange")
                return
            data_item = {
                "full_text": full_text,
                "chinese_only": chinese_text,
                "xpath": xpath,
                "action": action,
                "custom_name": dialog.custom_name,
                "element_type": element.tag_name,
                "captured_at": datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            }
            self.collected_data.append(data_item)
            if action == "getText":
                self._log_info(f"已采集(获取文本): {full_text} [名称: {dialog.custom_name}] [XPath: {xpath}]", "green")
            elif action == "click":
                self._log_info(f"已采集(点击): {full_text} [名称: {dialog.custom_name}] [XPath: {xpath}]", "blue")
            elif action == "clickAndGetClipboard":
                self._log_info(f"已采集(点击并获取剪贴板): {full_text} [名称: {dialog.custom_name}] [XPath: {xpath}]", "blue")
        except Exception as e:
            self._log_info(f"采集失败: {str(e)}", "red")
            logging.error(f"采集异常详情: {traceback.format_exc()}")
            return None
        finally:
            # 恢复置顶状态
            main_window.attributes('-topmost', was_topmost)
    
    def _get_hovered_xpath_recursive(self):
        """递归获取当前悬停元素的XPath，支持嵌套iframe"""
        if not self.driver:
            self._log_info('无法获取XPath：浏览器未连接', 'red')
            return None
        get_xpath_script = '''
        function getXPath(element) {
            if (element && element.id) return '//*[@id="' + element.id + '"]';
            var paths = [];
            for (; element && element.nodeType == 1; element = element.parentNode) {
                var index = 0;
                var hasSimilarSibling = false;
                for (var sibling = element.previousSibling; sibling; sibling = sibling.previousSibling) {
                    if (sibling.nodeType == 1 && sibling.tagName === element.tagName) {
                        index++;
                    }
                }
                for (var sibling = element.nextSibling; sibling; sibling = sibling.nextSibling) {
                    if (sibling.nodeType == 1 && sibling.tagName === element.tagName) {
                        hasSimilarSibling = true;
                        break;
                    }
                }
                var tagName = element.tagName.toLowerCase();
                var pathIndex = (index > 0 || hasSimilarSibling) ? `[${index + 1}]` : '';
                paths.splice(0, 0, tagName + pathIndex);
            }
            return paths.length ? '/' + paths.join('/') : null;
        }
        if (!window.lastHoveredElement) return null;
        var xpath = getXPath(window.lastHoveredElement);
        window.lastHoveredElement = null;
        return xpath;
        '''
        try:
            xpath = self.driver.execute_script(get_xpath_script)
            if xpath:
                return xpath
            iframes = self.driver.find_elements(By.TAG_NAME, 'iframe')
            for i in range(len(iframes)):
                try:
                    self.driver.switch_to.frame(i)
                    xpath_in_frame = self._get_hovered_xpath_recursive()
                    if xpath_in_frame:
                        return xpath_in_frame
                except Exception:
                    pass
                finally:
                    self.driver.switch_to.parent_frame()
        except Exception as e:
            self._log_info(f'通过JS获取XPath时出错: {e}', 'red')
        return None
        
    def _check_shipping_info_before_export(self):
        """检查并尝试修复收货信息字段，增加人工审核机制"""
        # 检查pyperclip依赖
        if pyperclip is None:
            self._log_info("pyperclip模块未正确导入，请安装: pip install pyperclip", "red")
            from tkinter import messagebox
            messagebox.showerror("错误", "pyperclip模块未安装，无法处理剪贴板内容")
            return False
            
        if not hasattr(self, 'collected_data') or not self.collected_data:
            self._log_info("没有收货信息可以导出", "red")
            from tkinter import messagebox
            messagebox.showwarning("导出失败", "没有收货信息可以导出，请先采集数据。")
            return False
            
        # 检查所有订单的收货信息
        orders_to_review = []
        
        # 1. 首先检查已标记为需要审核的订单
        if hasattr(self, 'orders_need_review') and self.orders_need_review:
            for order_id in self.orders_need_review:
                if order_id in self.order_clipboard_contents:
                    content = self.order_clipboard_contents[order_id]
                    is_valid, confidence, reason = self._is_valid_shipping_info(content)
                    orders_to_review.append((order_id, content, confidence, reason, "已标记为需要审核"))
        
        # 2. 再次检查所有订单的收货信息有效性
        for order_id, content in list(self.order_clipboard_contents.items()):
            if any(order_id == x[0] for x in orders_to_review):
                continue  # 跳过已添加的订单
            
            is_valid, confidence, reason = self._is_valid_shipping_info(content)
            
            # 如果可信度低于阈值，添加到需要审核的列表
            if not is_valid or confidence < 60:  # 使用更高的阈值进行二次检查
                orders_to_review.append((order_id, content, confidence, reason, "二次检查发现可能问题"))
        
        # 3. 如果存在需要审核的订单，提供批量审核界面
        if orders_to_review:
            self._log_info(f"检测到{len(orders_to_review)}个订单的收货信息需要人工审核", "orange")
            
            # 按可信度升序排序，先处理最可疑的
            orders_to_review.sort(key=lambda x: x[2])
            
            # 调用批量审核函数
            self._batch_review_shipping_info(orders_to_review)
            
        # 继续原有的导出流程
        field_name = '复制完整收货信息'  # 收货信息字段名
        fixed_count = 0
        
        # 确保order_clipboard_contents字典存在
        if not hasattr(self, 'order_clipboard_contents'):
            self.order_clipboard_contents = {}
            self._log_info("警告: 订单ID与收货信息映射字典不存在，已创建空字典", "red")
            print("DEBUG-EXPORT-ERROR: 订单ID与收货信息映射字典不存在，已创建空字典")
        else:
            # 导出前先确保所有映射都已保存到文件
            self._save_clipboard_mappings()
            self._log_info("已将当前所有映射保存到文件", "blue")
            
        # 输出当前订单ID与收货信息的映射关系
        self._log_info(f"导出前订单ID与收货信息映射关系：", "blue")
        print(f"DEBUG-EXPORT-MAP-COUNT: 映射字典中包含 {len(self.order_clipboard_contents)} 个订单")
        for order_id, content in self.order_clipboard_contents.items():
            content_preview = content[:30] + "..." if content else "空"
            print(f"DEBUG-EXPORT-MAP: 订单ID '{order_id}' -> 收货信息: '{content_preview}'")
            self._log_info(f"订单ID: {order_id}, 收货信息: {content_preview}", "blue")
        
        # 检查每个订单的收货信息是否唯一
        shipping_info_set = set()
        duplicate_count = 0
        
        for order_data in self.collected_data:
            # 获取当前订单ID
            order_id = order_data.get('订单编号', '')
            if isinstance(order_id, str) and order_id.startswith('订单编号：'):
                order_id = order_id.replace('订单编号：', '')
            
            self._log_info(f"处理订单: {order_id}", "blue")
            
            # 检查是否存在此字段
            if field_name in order_data:
                # 检查内容是否为布尔值或空
                if order_data[field_name] is True or order_data[field_name] is False or not order_data[field_name]:
                    self._log_info(f"检测到订单 {order_id} 的无效收货信息字段值: {order_data[field_name]}", "orange")
                    
                    # 优先使用订单专属的收货信息
                    if order_id and order_id in self.order_clipboard_contents:
                        order_data[field_name] = self.order_clipboard_contents[order_id]
                        fixed_count += 1
                        self._log_info(f"使用订单专属收货信息修复: '{self.order_clipboard_contents[order_id][:30]}...'", "green")
                        
                        # 检查是否是重复的收货信息
                        if order_data[field_name] in shipping_info_set:
                            duplicate_count += 1
                            self._log_info(f"警告: 订单 {order_id} 的收货信息与其他订单重复", "red")
                        else:
                            shipping_info_set.add(order_data[field_name])
                        
                        continue
                    
                    # 其次尝试使用全局变量
                    if hasattr(self, 'last_clipboard_content') and self.last_clipboard_content:
                        order_data[field_name] = self.last_clipboard_content
                        fixed_count += 1
                        self._log_info(f"使用全局变量修复了收货信息: '{self.last_clipboard_content[:30]}...'", "green")
                        continue
                    
                    # 最后尝试从剪贴板获取最新内容
                    import pyperclip
                    clipboard_content = pyperclip.paste()
                    if clipboard_content and clipboard_content.strip():
                        order_data[field_name] = clipboard_content
                        fixed_count += 1
                        self._log_info(f"使用剪贴板内容修复了收货信息: '{clipboard_content[:30]}...'", "green")
                        continue
                        
                    # 如果都失败了，记录一个明确的错误信息
                    order_data[field_name] = "【收货信息获取失败】"
                    self._log_info(f"无法修复订单 {order_id} 的收货信息字段，已标记为失败", "red")
            else:
                # 如果字段不存在，优先使用订单专属的收货信息
                if order_id and order_id in self.order_clipboard_contents:
                    order_data[field_name] = self.order_clipboard_contents[order_id]
                    fixed_count += 1
                    self._log_info(f"添加了订单 {order_id} 的专属收货信息: '{self.order_clipboard_contents[order_id][:30]}...'", "green")
                    continue
                
                # 其次尝试使用全局变量
                elif hasattr(self, 'last_clipboard_content') and self.last_clipboard_content:
                    order_data[field_name] = self.last_clipboard_content
                    fixed_count += 1
                    self._log_info(f"添加了缺失的收货信息字段: '{self.last_clipboard_content[:30]}...'", "green")
                else:
                    # 最后尝试从剪贴板获取
                    import pyperclip
                    clipboard_content = pyperclip.paste()
                    if clipboard_content and clipboard_content.strip():
                        order_data[field_name] = clipboard_content
                        fixed_count += 1
                        self._log_info(f"添加了缺失的收货信息字段: '{clipboard_content[:30]}...'", "green")
        
        if fixed_count > 0:
            self._log_info(f"导出前共修复了 {fixed_count} 条收货信息记录", "green")
        
        if duplicate_count > 0:
            self._log_info(f"警告: 检测到 {duplicate_count} 个订单的收货信息重复", "red")
            
        # 1. 首先检查已标记为需要审核的订单
        if hasattr(self, 'orders_need_review') and self.orders_need_review:
            for order_id in self.orders_need_review:
                if order_id in self.order_clipboard_contents:
                    content = self.order_clipboard_contents[order_id]
                    is_valid, confidence, reason = self._is_valid_shipping_info(content)
                    orders_to_review.append((order_id, content, confidence, reason, "已标记为需要审核"))
        
        # 2. 再次检查所有订单的收货信息有效性
        for order_id, content in list(self.order_clipboard_contents.items()):
            if any(order_id == x[0] for x in orders_to_review):
                continue  # 跳过已添加的订单
            
            is_valid, confidence, reason = self._is_valid_shipping_info(content)
            
            # 如果可信度低于阈值，添加到需要审核的列表
            if not is_valid or confidence < 60:  # 使用更高的阈值进行二次检查
                orders_to_review.append((order_id, content, confidence, reason, "二次检查发现可能问题"))
        
        # 3. 如果存在需要审核的订单，提供批量审核界面
        if orders_to_review:
            self._log_info(f"检测到{len(orders_to_review)}个订单的收货信息需要人工审核", "orange")
            
            # 按可信度升序排序，先处理最可疑的
            orders_to_review.sort(key=lambda x: x[2])
            
            # 强制主线程处理UI事件
            self.root.update_idletasks()
            self.root.update()
            
            # 调用批量审核函数
            self._batch_review_shipping_info(orders_to_review)
            
            # 确保界面显示
            self.root.update_idletasks()
            self.root.update()
        
        # 导出前检查收货信息字段内容
        self._log_info("导出前检查收货信息字段内容:", "blue")
        for i, order_data in enumerate(self.collected_data):
            order_id = order_data.get('订单编号', '')
            if isinstance(order_id, str) and order_id.startswith('订单编号：'):
                order_id = order_id.replace('订单编号：', '')
            else:
                order_id = str(order_id)
                
            shipping_info = order_data.get(field_name, "未设置")
            shipping_info_type = type(shipping_info).__name__
            shipping_info_preview = shipping_info[:30] + "..." if isinstance(shipping_info, str) and shipping_info != "未设置" else str(shipping_info)
            
            print(f"DEBUG-EXPORT-CHECK-{i+1}: 订单ID '{order_id}', {field_name} = {shipping_info_preview} (类型: {shipping_info_type})")
            self._log_info(f"记录 {i+1}: 订单ID={order_id}, {field_name} = {shipping_info_preview} (类型: {shipping_info_type})", "blue")
            
        return fixed_count > 0
    
    def _export_excel(self):
        """导出数据到Excel（正常模式专用）"""
        # 检查pandas依赖
        if pd is None:
            self._log_info("pandas模块未正确导入，请安装: pip install pandas", "red")
            messagebox.showerror("错误", "pandas模块未安装，无法导出Excel文件")
            return
            
        # 确保是在正常模式下使用此功能
        if self.collection_mode.get() != "正常模式":
            messagebox.showinfo("提示", "Excel导出功能仅在正常模式下可用")
            return
            
        if len(self.collected_data) == 0:
            messagebox.showinfo("提示", "没有可导出的数据")
            return
        
        # 确保order_clipboard_contents字典存在
        if not hasattr(self, 'order_clipboard_contents'):
            self.order_clipboard_contents = {}
            self._log_info("警告: 导出前发现订单ID与收货信息映射字典不存在，已创建空字典", "red")
            print("DEBUG-EXPORT-EXCEL-ERROR: 订单ID与收货信息映射字典不存在，已创建空字典")
            
        # 检查并尝试修复收货信息
        self._check_shipping_info_before_export()
        
        # 添加调试日志，显示每条记录的收货信息字段内容
        self._log_info("导出前检查收货信息字段内容:", "blue")
        for i, order_data in enumerate(self.collected_data):
            field_name = '复制完整收货信息'
            order_id = order_data.get('订单编号', '')
            if isinstance(order_id, str) and order_id.startswith('订单编号：'):
                order_id = order_id.replace('订单编号：', '')
                
            if field_name in order_data:
                value = order_data[field_name]
                value_type = type(value).__name__
                value_preview = str(value)[:30] + "..." if isinstance(value, str) and len(str(value)) > 30 else str(value)
                
                # 检查是否使用了订单专属收货信息
                if order_id and order_id in self.order_clipboard_contents:
                    order_specific_info = self.order_clipboard_contents[order_id]
                    is_using_specific = (value == order_specific_info)
                    source_info = "（使用订单专属收货信息）" if is_using_specific else "（未使用订单专属收货信息）"
                    self._log_info(f"记录 {i+1}: 订单ID={order_id}, {field_name} = {value_preview} {source_info} (类型: {value_type})", "blue")
                else:
                    self._log_info(f"记录 {i+1}: 订单ID={order_id}, {field_name} = {value_preview} (类型: {value_type})", "blue")
            else:
                self._log_info(f"记录 {i+1}: 订单ID={order_id}, {field_name} 字段不存在", "orange")
                
        try:
            # 检查是否安装了openpyxl
            try:
                import openpyxl
            except ImportError:
                result = messagebox.askquestion("提示", 
                    "缺少Excel导出所需的openpyxl库。是否安装？\n(需要联网，可能需要几分钟)")
                if result == 'yes':
                    self._log_info("正在安装openpyxl库...", "blue")
                    import subprocess
                    import sys
                    try:
                        subprocess.check_call([sys.executable, "-m", "pip", "install", "openpyxl"])
                        self._log_info("openpyxl库安装成功", "green")
                        import openpyxl
                    except Exception as e:
                        self._log_info(f"安装失败: {str(e)}", "red")
                        messagebox.showerror("错误", f"安装openpyxl失败，请手动安装:\npip install openpyxl\n\n错误信息: {str(e)}")
                        return
                else:
                    return
            
            file_path = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel files", "*.xlsx")]
            )
            if not file_path:
                return
            
            # 创建数据框架，支持新的字典格式
            # 从collected_data创建DataFrame，每个订单一行
            df = pd.DataFrame(self.collected_data)
            
            # 检测并去重重复的列（如'复制完整收货信息'和'复制完整的收货信息'）
            duplicate_columns = []
            columns_to_merge = {
                '复制完整收货信息': ['复制完整收货信息', '复制完整的收货信息'],
                # 可以在这里添加其他需要合并的列
            }
            
            for target_col, source_cols in columns_to_merge.items():
                existing_cols = [col for col in source_cols if col in df.columns]
                if len(existing_cols) > 1:
                    self._log_info(f"检测到重复列: {existing_cols}，将合并为: {target_col}", "blue")
                    # 合并列数据，优先使用非空值
                    merged_data = []
                    for index, row in df.iterrows():
                        merged_value = ""
                        for col in existing_cols:
                            value = row.get(col, "")
                            if value and str(value).strip() and str(value).strip() != "nan":
                                merged_value = value
                                break
                        merged_data.append(merged_value)
                    
                    # 添加合并后的列
                    df[target_col] = merged_data
                    
                    # 标记要删除的重复列
                    duplicate_columns.extend(existing_cols)
            
            # 删除重复列
            if duplicate_columns:
                # 保留目标列，删除源列
                cols_to_drop = [col for col in duplicate_columns if col in columns_to_merge.keys() and col in df.columns]
                remaining_cols_to_drop = [col for col in duplicate_columns if col not in columns_to_merge.keys()]
                df = df.drop(columns=remaining_cols_to_drop)
                self._log_info(f"已删除重复列: {remaining_cols_to_drop}", "green")
            
            # 导出到Excel
            df.to_excel(file_path, index=False)
            self._log_info(f"Excel导出成功: {file_path}", "green")
        except Exception as e:
            self._log_info(f"Excel导出失败: {str(e)}", "red")
            import traceback
            self._log_info(traceback.format_exc(), "red")
    
    def _export_word(self):
        """导出数据到Word（正常模式专用）- Markdown格式，按商品名称分组"""
        # 检查docx依赖
        if Document is None:
            self._log_info("python-docx模块未正确导入，请安装: pip install python-docx", "red")
            messagebox.showerror("错误", "python-docx模块未安装，无法导出Word文件")
            return
            
        # 确保是在正常模式下使用此功能
        if self.collection_mode.get() != "正常模式":
            messagebox.showinfo("提示", "Word导出功能仅在正常模式下可用")
            return
            
        if len(self.collected_data) == 0:
            messagebox.showinfo("提示", "没有可导出的数据")
            return
            
        # 确保order_clipboard_contents字典存在
        if not hasattr(self, 'order_clipboard_contents'):
            self.order_clipboard_contents = {}
            self._log_info("警告: 导出前发现订单ID与收货信息映射字典不存在，已创建空字典", "red")
            print("DEBUG-EXPORT-WORD-ERROR: 订单ID与收货信息映射字典不存在，已创建空字典")
            
        # 检查并尝试修复收货信息
        self._check_shipping_info_before_export()
        
        # 添加调试日志，显示每条记录的收货信息字段内容
        self._log_info("导出前检查收货信息字段内容:", "blue")
        for i, order_data in enumerate(self.collected_data):
            field_name = '复制完整收货信息'
            order_id = order_data.get('订单编号', '')
            if isinstance(order_id, str) and order_id.startswith('订单编号：'):
                order_id = order_id.replace('订单编号：', '')
                
            if field_name in order_data:
                value = order_data[field_name]
                value_type = type(value).__name__
                value_preview = str(value)[:30] + "..." if isinstance(value, str) and len(str(value)) > 30 else str(value)
                
                # 检查是否使用了订单专属收货信息
                if order_id and order_id in self.order_clipboard_contents:
                    order_specific_info = self.order_clipboard_contents[order_id]
                    is_using_specific = (value == order_specific_info)
                    source_info = "（使用订单专属收货信息）" if is_using_specific else "（未使用订单专属收货信息）"
                    self._log_info(f"记录 {i+1}: 订单ID={order_id}, {field_name} = {value_preview} {source_info} (类型: {value_type})", "blue")
                else:
                    self._log_info(f"记录 {i+1}: 订单ID={order_id}, {field_name} = {value_preview} (类型: {value_type})", "blue")
            else:
                self._log_info(f"记录 {i+1}: 订单ID={order_id}, {field_name} 字段不存在", "orange")
                
        try:
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word files", "*.docx")]
            )
            if not file_path:
                return
                
            doc = Document()
            
            # 添加文档标题
            title = doc.add_heading("收货信息采集报告", level=1)
            title.alignment = 1  # 居中对齐
            
            # 添加时间戳
            timestamp = doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            timestamp.alignment = 1  # 居中对齐
            
            # 处理重复字段的去重逻辑
            columns_to_merge = {
                '复制完整收货信息': ['复制完整收货信息', '复制完整的收货信息'],
                # 可以在这里添加其他需要合并的列
            }
            
            # 创建处理后的数据副本
            processed_data = []
            for order_data in self.collected_data:
                processed_order = order_data.copy()
                
                # 处理每个需要合并的字段组
                for target_col, source_cols in columns_to_merge.items():
                    existing_cols = [col for col in source_cols if col in processed_order]
                    if len(existing_cols) > 1:
                        # 合并字段数据，优先使用非空值
                        merged_value = ""
                        for col in existing_cols:
                            value = processed_order.get(col, "")
                            if value and str(value).strip() and str(value).strip() != "nan":
                                merged_value = value
                                break
                        
                        # 设置合并后的值
                        processed_order[target_col] = merged_value
                        
                        # 删除重复的源字段（保留目标字段）
                        for col in existing_cols:
                            if col != target_col:
                                processed_order.pop(col, None)
                
                processed_data.append(processed_order)
            
            # 记录去重信息
            original_fields = set()
            processed_fields = set()
            for order_data in self.collected_data:
                original_fields.update(order_data.keys())
            for order_data in processed_data:
                processed_fields.update(order_data.keys())
            
            removed_fields = original_fields - processed_fields
            if removed_fields:
                self._log_info(f"Word导出时已去重字段: {list(removed_fields)}", "green")
            
            # 按商品名称分组数据
            grouped_data = {}
            for order_data in processed_data:
                # 尝试从多个可能的字段中获取商品名称
                product_name = None
                for field_name in ['商品名称', '商品', '产品名称', '产品', '商品信息']:
                    if field_name in order_data and order_data[field_name]:
                        product_name = str(order_data[field_name]).strip()
                        break
                
                # 如果没有找到商品名称，使用默认分组
                if not product_name or product_name == "nan":
                    product_name = "未知商品"
                
                if product_name not in grouped_data:
                    grouped_data[product_name] = []
                grouped_data[product_name].append(order_data)
            
            # 按商品名称排序
            sorted_products = sorted(grouped_data.keys())
            
            # 为每个商品组生成内容
            for product_name in sorted_products:
                orders = grouped_data[product_name]
                
                # 添加商品名称小标题
                product_heading = doc.add_heading(f"商品名称：{product_name}", level=2)
                
                # 为每个订单添加详细信息
                for i, order_data in enumerate(orders):
                    # 添加订单序号（如果同一商品有多个订单）
                    if len(orders) > 1:
                        order_heading = doc.add_heading(f"订单 {i+1}", level=3)
                    
                    # 以markdown格式添加订单详细信息
                    for field_name, field_value in order_data.items():
                        if field_value and str(field_value).strip() and str(field_value).strip() != "nan":
                            # 创建字段信息段落
                            p = doc.add_paragraph()
                            # 添加字段名（加粗）
                            field_run = p.add_run(f"{field_name}：")
                            field_run.bold = True
                            # 添加字段值
                            value_run = p.add_run(str(field_value))
                    
                    # 在订单之间添加分隔线（除了最后一个订单）
                    if i < len(orders) - 1:
                        doc.add_paragraph("" + "-" * 50)
                
                # 在商品组之间添加空行
                if product_name != sorted_products[-1]:  # 不是最后一个商品组
                    doc.add_paragraph("")
                    doc.add_paragraph("" + "=" * 80)
                    doc.add_paragraph("")
            
            doc.save(file_path)
            self._log_info(f"Word导出成功: {file_path} (按商品名称分组，共{len(sorted_products)}个商品组)", "green")
        except Exception as e:
            self._log_info(f"Word导出失败: {str(e)}", "red")
            import traceback
            self._log_info(f"详细错误信息: {traceback.format_exc()}", "red")
    
    def _export_json(self):
        """导出数据到JSON（采集模式专用）"""
        # 确保是在采集模式下使用此功能
        if self.collection_mode.get() != "采集模式":
            messagebox.showinfo("提示", "JSON导出功能仅在采集模式下可用")
            return
            
        if not hasattr(self, 'collected_data') or len(self.collected_data) == 0:
            messagebox.showinfo("提示", "没有可导出的数据")
            return
            
        # 确保order_clipboard_contents字典存在
        if not hasattr(self, 'order_clipboard_contents'):
            self.order_clipboard_contents = {}
            self._log_info("警告: 导出前发现订单ID与收货信息映射字典不存在，已创建空字典", "red")
            print("DEBUG-EXPORT-JSON-ERROR: 订单ID与收货信息映射字典不存在，已创建空字典")
            
        # 检查并尝试修复收货信息
        self._check_shipping_info_before_export()
            
        try:
            file_path = filedialog.asksaveasfilename(
                defaultextension=".json",
                filetypes=[("JSON files", "*.json")]
            )
            if not file_path:
                return
                
            # 导出数据，确保JSON友好
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(self.collected_data, f, ensure_ascii=False, indent=2)
                
            self._log_info(f"JSON导出成功: {file_path}", "green")
        except Exception as e:
            self._log_info(f"JSON导出失败: {str(e)}", "red")
            logging.error(f"JSON导出异常: {traceback.format_exc()}")
    
    def _start_distance_learning(self):
        """开始辅助定位（学习连续订单之间的滚动距离）"""
        if not self.is_browser_connected:
            self._log_info("浏览器未连接，无法开始辅助定位", "red")
            return
            
        # 重置相关状态变量
        self.is_distance_learning = True
        self.distance_learning_step = 0
        self.first_element_position = None
        self.second_element_position = None
        self.scroll_distance = None
        
        # 禁用其他按钮，只保留终止按钮可用
        self.start_button.config(state=tk.DISABLED)
        self.pause_button.config(state=tk.DISABLED)
        self.continue_button.config(state=tk.DISABLED)
        self.stop_button.config(state=tk.NORMAL)
        
        self._log_info("已进入辅助定位模式", "blue")
        self._log_info("请将鼠标放在第1个订单的目标元素上（如订单号），按'.'键采集", "blue")
    
    def _get_element_at_cursor(self):
        """获取当前鼠标位置下的元素及其位置信息"""
        if not self.driver:
            self._log_info("浏览器未连接，无法获取元素", "red")
            return None
            
        try:
            # 获取当前鼠标屏幕位置
            pt = POINT()
            windll.user32.GetCursorPos(ctypes.byref(pt))
            mouse_x, mouse_y = pt.x, pt.y
            
            # 获取浏览器窗口位置和大小
            window_pos = self.driver.get_window_position()
            window_size = self.driver.get_window_size()
            browser_left = window_pos['x']
            browser_top = window_pos['y']
            browser_right = browser_left + window_size['width']
            browser_bottom = browser_top + window_size['height']
            
            if not (browser_left <= mouse_x <= browser_right and browser_top <= mouse_y <= browser_bottom):
                self._log_info("鼠标不在浏览器窗口内", "red")
                return None
            
            # 动态获取内容区顶部偏移
            js = "return {screenY: window.screenY, outerHeight: window.outerHeight, innerHeight: window.innerHeight};"
            win_metrics = self.driver.execute_script(js)
            content_top = win_metrics['screenY'] + (win_metrics['outerHeight'] - win_metrics['innerHeight'])
            viewport_x = mouse_x - browser_left
            viewport_y = mouse_y - content_top
            
            # JavaScript脚本获取元素信息
            script = '''
            function getElementInfoWithPosition(x, y) {
                let el = document.elementFromPoint(x, y);
                
                if (!el) {
                    return { found: false, reason: "未找到元素" };
                }
                
                // 检查元素是否可见
                const rect = el.getBoundingClientRect();
                const isVisible = !!(rect.width && rect.height && 
                    window.getComputedStyle(el).visibility !== 'hidden' && 
                    window.getComputedStyle(el).display !== 'none');
                
                if (!isVisible) {
                    return { found: true, visible: false, reason: "元素存在但不可见" };
                }
                
                // 获取文本内容
                const text = el.innerText || el.textContent || "";
                
                // 生成XPath
                function getElementXPath(element) {
                    if (!element) return '';
                    
                    // 处理document节点
                    if (element.nodeType === 9) return '/';
                    
                    // 处理普通节点
                    var path = [];
                    var current = element;
                    
                    while (current && current.nodeType === 1) {
                        var index = 0;
                        var found = false;
                        
                        for (var i = 0; i < current.parentNode.childNodes.length; i++) {
                            var node = current.parentNode.childNodes[i];
                            if (node.nodeType === 1 && node.tagName === current.tagName) {
                                index++;
                            }
                            if (node === current) {
                                found = true;
                                break;
                            }
                        }
                        
                        var tagName = current.tagName.toLowerCase();
                        var pathIndex = (index > 1 || !found) ? '[' + index + ']' : '';
                        
                        if (current.id) {
                            path.unshift(tagName + '[@id="' + current.id + '"]');
                            break;
                        } else {
                            path.unshift(tagName + pathIndex);
                        }
                        
                        current = current.parentNode;
                    }
                    
                    return '/' + path.join('/');
                }
                
                // 获取元素位置和属性信息
                return {
                    found: true,
                    visible: true,
                    text: text,
                    xpath: getElementXPath(el),
                    y_position: rect.top,
                    x_position: rect.left,
                    height: rect.height,
                    width: rect.width,
                    element: {
                        tagName: el.tagName,
                        id: el.id || "",
                        className: el.className || "",
                        type: el.getAttribute('type') || "",
                        name: el.getAttribute('name') || ""
                    }
                };
            }
            
            return getElementInfoWithPosition(arguments[0], arguments[1]);
            '''
            
            # 执行脚本获取元素信息
            element_info = self.driver.execute_script(script, viewport_x, viewport_y)
            
            # 验证返回结果
            if not element_info:
                self._log_info("无法获取元素信息", "red")
                return None
                
            if not element_info.get('found', False):
                reason = element_info.get('reason', '未知原因')
                self._log_info(f"未找到元素: {reason}", "red")
                return None
                
            if not element_info.get('visible', False):
                self._log_info("元素存在但不可见", "red")
                return None
                
            return element_info
            
        except Exception as e:
            self._log_info(f"获取元素信息失败: {str(e)}", "red")
            logging.error(f"获取元素信息异常详情: {traceback.format_exc()}")
            return None
    
    def _process_distance_learning(self):
        """处理辅助定位模式下的元素采集"""
        try:
            # 获取当前鼠标位置下的元素信息
            element_info = self._get_element_at_cursor()
            if not element_info:
                return
                
            # 提取所需信息
            xpath = element_info.get('xpath', '')
            y_position = element_info.get('y_position', 0)
            text = element_info.get('text', '')
            tag_name = element_info.get('element', {}).get('tagName', '')
            class_name = element_info.get('element', {}).get('className', '')
            
            # 根据当前学习步骤处理
            if self.distance_learning_step == 0:
                # 第一次采集
                self.first_element_position = {
                    'xpath': xpath,
                    'y': y_position,
                    'text': text,
                    'tag': tag_name,
                    'class': class_name
                }
                self.distance_learning_step = 1
                self._log_info(f"已采集第1个订单的元素 [XPath: {xpath}]", "green")
                self._log_info("请滚动页面，将鼠标放在第2个订单的相同类型元素上（位置相同但内容可能不同的元素），然后按'.'键采集", "blue")
                
            elif self.distance_learning_step == 1:
                # 第二次采集
                self.second_element_position = {
                    'xpath': xpath,
                    'y': y_position,
                    'text': text,
                    'tag': tag_name,
                    'class': class_name
                }
                
                # 验证元素类型一致性
                if self.first_element_position and self.second_element_position and self.first_element_position.get('tag') != self.second_element_position.get('tag'):
                    self._log_info(f"警告：两次采集的元素类型不同 (第一次: {self.first_element_position.get('tag')}, 第二次: {self.second_element_position.get('tag')})", "orange")
                    self._log_info("建议重新开始辅助定位，确保选择相同类型的元素", "orange")
                
                # 确保两个元素位置都有效
                if self.first_element_position and self.second_element_position and 'y' in self.first_element_position and 'y' in self.second_element_position:
                    # 计算距离
                    distance = self.second_element_position['y'] - self.first_element_position['y']
                    
                    # 验证距离合理性
                    if distance <= 0:
                        self._log_info(f"警告：计算的滚动距离为负值或零 ({distance}px)，这可能不正确", "orange")
                        self._log_info("请确保第二个元素在第一个元素下方，建议重新开始辅助定位", "orange")
                        # 不保存这个不合理的距离值
                    elif distance > 1000:
                        self._log_info(f"警告：计算的滚动距离过大 ({distance}px)，这可能不正确", "orange")
                        self._log_info("请确保选择了相邻两个订单的元素，建议重新开始辅助定位", "orange")
                        # 不保存这个不合理的距离值
                    else:
                        self.scroll_distance = distance
                        self._log_info(f"辅助定位成功完成！两个订单之间的滚动距离为: {distance}px", "green")
                        self._log_info("现在您可以开始正常采集流程，系统将使用这个距离自动滚动到下一个订单", "green")
                else:
                    self._log_info("错误：无法计算滚动距离，元素位置信息不完整", "red")
                    self._log_info("请重新开始辅助定位", "red")
                
                # 完成学习并恢复状态
                self.is_distance_learning = False
                self.distance_learning_step = 2
                
                # 恢复按钮状态
                self.start_button.config(state=tk.NORMAL)
                self.stop_button.config(state=tk.NORMAL)
                
        except Exception as e:
            self._log_info(f"辅助定位过程中出错: {str(e)}", "red")
            logging.error(f"辅助定位异常: {traceback.format_exc()}")
            
            # 出错时重置状态
            self.is_distance_learning = False
            self.distance_learning_step = 0
            self.first_element_position = None
            self.second_element_position = None
            
            # 恢复按钮状态
            self.start_button.config(state=tk.NORMAL)
            self.stop_button.config(state=tk.NORMAL)
    
    def _scroll_to_next_order(self):
        """使用多种滚动策略尝试滚动到下一个订单"""
        # 检查是否已终止操作
        if not self.is_running:
            self._log_info("操作已终止，停止滚动操作", "orange")
            return False
            
        if not self.driver:
            self._log_info("无法滚动：浏览器未连接", "red")
            return False
            
        # 获取当前订单ID，用于后续比较
        current_order_id = self._extract_current_order_id()
        self._log_info(f"滚动前订单ID: {current_order_id}", "blue")
        
        # 保存滚动前的截图（如果启用了调试模式）
        if hasattr(self, 'confirm_click') and self.confirm_click.get():
            self._save_screenshot("before_scroll")
        
        # 根据是否是调试模式选择不同的滚动策略
        if hasattr(self, 'confirm_click') and self.confirm_click.get():
            # 调试模式下，只使用JavaScript滚动和键盘滚动，不使用鼠标拖动
            scroll_methods = [
                self._scroll_with_javascript,
                self._scroll_with_keys,
                self._click_next_page
            ]
            self._log_info("调试模式下使用精确滚动", "blue")
            
            # 调试模式下增加重试次数和等待时间
            max_retries = 5  # 增加到5次尝试
            wait_time = 3.0  # 调试模式下等待更长时间
        else:
            # 正常模式下，使用JavaScript滚动和键盘滚动，不使用鼠标拖动
            scroll_methods = [
                self._scroll_with_javascript,
                self._scroll_with_keys,
                self._click_next_page
            ]
            max_retries = 1
            wait_time = 1.5
        
        # 尝试多次滚动
        for retry in range(max_retries):
            if retry > 0:
                self._log_info(f"第 {retry+1}/{max_retries} 次尝试滚动", "blue")
                
                # 重试时增加滚动距离
                multiplier = (retry + 1) * 0.8  # 增加倍数，使滚动距离更接近正确位置
                self._log_info(f"增加滚动距离倍数: {multiplier}x", "blue")
                
                # 在调试模式下，每次重试前等待更长时间让页面稳定
                if hasattr(self, 'confirm_click') and self.confirm_click.get():
                    time.sleep(1.0)
            else:
                multiplier = 1.0
            
            # 在调试模式下，先尝试使用JavaScript滚动
            if hasattr(self, 'confirm_click') and self.confirm_click.get():
                try:
                    self._log_info(f"调试模式下先尝试JavaScript滚动 (倍数: {multiplier}x)", "blue")
                    success = self._scroll_with_javascript(multiplier=multiplier)
                    if not success:
                        self._log_info("JavaScript滚动返回失败状态", "orange")
                    
                    # 增加滚动后等待时间，让页面有足够时间更新
                    time.sleep(wait_time * 1.5)
                    
                    # 检查是否滚动到了新订单
                    new_order_id = self._extract_current_order_id()
                    self._log_info(f"滚动后订单ID: {new_order_id}", "blue")
                    
                    if new_order_id and new_order_id != current_order_id:
                        self._log_info(f"成功滚动到新订单: {new_order_id}", "green")
                        return True
                    
                    # 参考代码的策略：即使ID没变，也尝试查找下一个订单的元素
                    # 这是关键改进：先滚动，然后尝试查找下一个订单的元素
                    try:
                        # 尝试定位第二个订单的元素
                        if hasattr(self, 'ref1_xpath') and hasattr(self, 'ref2_xpath') and self.ref1_xpath and self.ref2_xpath:
                            self._log_info("尝试直接查找第二个订单的参照元素", "blue")
                            element = self.driver.find_element(By.XPATH, self.ref2_xpath)
                            if element:
                                self._log_info("成功找到第二个订单的参照元素，滚动成功", "green")
                                return True
                    except Exception as e:
                        self._log_info(f"查找第二个订单参照元素失败: {str(e)}", "orange")
                except Exception as e:
                    self._log_info(f"JavaScript滚动失败: {str(e)}", "red")
            
            # 尝试所有滚动方法
            for method in scroll_methods:
                try:
                    method_name = method.__name__
                    self._log_info(f"尝试使用 {method_name} 滚动到下一个订单", "blue")
                    
                    # 对JavaScript滚动方法传递倍数参数
                    if method_name == '_scroll_with_javascript':
                        success = method(multiplier=multiplier)
                        if not success:
                            self._log_info(f"{method_name} 返回失败状态", "orange")
                            continue
                    else:
                        method()
                    
                    # 增加滚动后等待时间，让页面有足够时间更新
                    time.sleep(wait_time * 1.5)
                    
                    # 检查是否滚动到了新订单
                    new_order_id = self._extract_current_order_id()
                    self._log_info(f"滚动后订单ID: {new_order_id}", "blue")
                    
                    if new_order_id and new_order_id != current_order_id:
                        self._log_info(f"成功滚动到新订单: {new_order_id}", "green")
                        return True
                    
                    # 参考代码的策略：即使ID没变，也尝试查找下一个订单的元素
                    try:
                        # 尝试定位第二个订单的元素
                        if hasattr(self, 'ref1_xpath') and hasattr(self, 'ref2_xpath') and self.ref1_xpath and self.ref2_xpath:
                            self._log_info("尝试直接查找第二个订单的参照元素", "blue")
                            element = self.driver.find_element(By.XPATH, self.ref2_xpath)
                            if element:
                                self._log_info("成功找到第二个订单的参照元素，滚动成功", "green")
                                return True
                    except Exception as e:
                        self._log_info(f"查找第二个订单参照元素失败: {str(e)}", "orange")
                    
                    self._log_info(f"{method_name} 滚动后订单未变化，尝试下一种方法", "orange")
                    
                    # 验证页面是否有变化（通过检查页面高度或其他元素）
                    try:
                        scroll_position_before = self.driver.execute_script("return window.pageYOffset;")
                        page_height_before = self.driver.execute_script("return document.body.scrollHeight;")
                        
                        # 尝试再次小幅滚动
                        self.driver.execute_script("window.scrollBy(0, 50);")
                        time.sleep(0.5)
                        
                        scroll_position_after = self.driver.execute_script("return window.pageYOffset;")
                        page_height_after = self.driver.execute_script("return document.body.scrollHeight;")
                        
                        if scroll_position_after > scroll_position_before or page_height_after != page_height_before:
                            self._log_info("检测到页面有变化，可能正在加载新内容", "blue")
                            time.sleep(1.0)  # 等待更长时间让内容加载
                            
                            # 再次检查订单ID
                            new_order_id = self._extract_current_order_id()
                            if new_order_id and new_order_id != current_order_id:
                                self._log_info(f"延迟检测到新订单: {new_order_id}", "green")
                                return True
                            # 再次尝试查找第二个订单的元素
                            try:
                                if hasattr(self, 'ref1_xpath') and hasattr(self, 'ref2_xpath') and self.ref1_xpath and self.ref2_xpath:
                                    element = self.driver.find_element(By.XPATH, self.ref2_xpath)
                                    if element:
                                        self._log_info("延迟成功找到第二个订单的参照元素，滚动成功", "green")
                                        return True
                            except Exception:
                                pass
                    except Exception as e:
                        self._log_info(f"页面变化检测失败: {str(e)}", "orange")
                except Exception as e:
                    self._log_info(f"{method_name} 滚动失败: {str(e)}", "red")
                    import traceback
                    logging.error(f"滚动异常: {traceback.format_exc()}")
            
            # 如果所有方法都失败，尝试点击页面中间然后按Page Down键
            if retry == max_retries - 1:  # 最后一次重试
                try:
                    self._log_info("尝试点击页面中间并按Page Down键", "blue")
                    
                    # 获取浏览器窗口位置和大小
                    window_rect = self.driver.execute_script("""
                        return {
                            width: window.innerWidth,
                            height: window.innerHeight
                        };
                    """)
                    
                    # 点击页面中间
                    actions = ActionChains(self.driver)
                    actions.move_by_offset(window_rect['width'] // 2, window_rect['height'] // 2)
                    actions.click()
                    actions.perform()
                    
                    # 等待一下
                    time.sleep(0.5)
                    
                    # 发送Page Down键
                    actions = ActionChains(self.driver)
                    actions.send_keys(Keys.PAGE_DOWN)
                    actions.perform()
                    
                    # 增加滚动后等待时间
                    time.sleep(wait_time * 1.5)
                    
                    # 检查是否滚动到了新订单
                    new_order_id = self._extract_current_order_id()
                    if new_order_id and new_order_id != current_order_id:
                        self._log_info(f"点击页面中间并按Page Down键成功滚动到新订单: {new_order_id}", "green")
                        return True
                    
                    # 参考代码的策略：即使ID没变，也尝试查找下一个订单的元素
                    try:
                        if hasattr(self, 'ref1_xpath') and hasattr(self, 'ref2_xpath') and self.ref1_xpath and self.ref2_xpath:
                            element = self.driver.find_element(By.XPATH, self.ref2_xpath)
                            if element:
                                self._log_info("通过Page Down成功找到第二个订单的参照元素，滚动成功", "green")
                                return True
                    except Exception:
                        pass
                except Exception as e:
                    self._log_info(f"点击页面中间并按Page Down键失败: {str(e)}", "red")
        
        # 如果所有方法都失败了，尝试增加滚动距离并重试
        self._log_info("所有滚动方法都失败，尝试增加滚动距离", "orange")
        try:
            # 增加滚动距离，再次尝试JavaScript滚动
            multiplier = 2.0 if hasattr(self, 'confirm_click') and self.confirm_click.get() else 1.5
            self._log_info(f"使用 {multiplier}x 倍滚动距离", "blue")
            success = self._scroll_with_javascript(multiplier=multiplier)
            if not success:
                self._log_info("最终JavaScript滚动返回失败状态", "orange")
                
            # 增加滚动后等待时间
            time.sleep(wait_time * 2)
            
            # 再次检查是否滚动到了新订单
            new_order_id = self._extract_current_order_id()
            if new_order_id and new_order_id != current_order_id:
                self._log_info(f"增加滚动距离后成功到达新订单: {new_order_id}", "green")
                return True
            
            # 参考代码的策略：即使ID没变，也尝试查找下一个订单的元素
            try:
                if hasattr(self, 'ref1_xpath') and hasattr(self, 'ref2_xpath') and self.ref1_xpath and self.ref2_xpath:
                    element = self.driver.find_element(By.XPATH, self.ref2_xpath)
                    if element:
                        self._log_info("最终成功找到第二个订单的参照元素，滚动成功", "green")
                        return True
            except Exception:
                pass
        except Exception as e:
            self._log_info(f"增加滚动距离失败: {str(e)}", "red")
        
            self._log_info("所有滚动方法都失败，无法到达下一个订单", "red")
            return False
        
    def _scroll_with_javascript(self, multiplier=1.0):
        """使用JavaScript滚动页面，确保第二个容器滚动到第一个容器的位置"""
        # 默认滚动距离，可以根据实际情况调整
        scroll_distance = 150 * multiplier  # 增加基础滚动距离
        
        # 调试模式下使用更精确的滚动
        if hasattr(self, 'confirm_click') and self.confirm_click.get():
            scroll_distance = 200 * multiplier  # 调试模式下使用更大的滚动距离
            self._log_info(f"调试模式下使用精确滚动距离: {scroll_distance}px", "blue")
            
            # 在调试模式下，保存滚动前的截图
            self._save_screenshot("before_js_scroll")
        
        # 尝试通过参照点计算精确滚动距离
        try:
            # 如果有两个参照点，尝试计算它们之间的距离
            if hasattr(self, 'ref1_xpath') and hasattr(self, 'ref2_xpath') and self.ref1_xpath and self.ref2_xpath:
                self._log_info("尝试使用参照点计算滚动距离", "blue")
                
                # 获取第一个参照点的位置
                ref1_position = self.driver.execute_script(f"""
                    const element = document.evaluate('{self.ref1_xpath}', document, null, XPathResult.FIRST_ORDERED_NODE_TYPE, null).singleNodeValue;
                    if (element) {{
                        const rect = element.getBoundingClientRect();
                        return {{
                            top: rect.top,
                            left: rect.left,
                            height: rect.height,
                            exists: true
                        }};
                    }}
                    return {{ exists: false }};
                """)
                
                # 获取第二个参照点的位置
                ref2_position = self.driver.execute_script(f"""
                    const element = document.evaluate('{self.ref2_xpath}', document, null, XPathResult.FIRST_ORDERED_NODE_TYPE, null).singleNodeValue;
                    if (element) {{
                        const rect = element.getBoundingClientRect();
                        return {{
                            top: rect.top,
                            left: rect.left,
                            height: rect.height,
                            exists: true
                        }};
                    }}
                    return {{ exists: false }};
                """)
                
                # 如果两个参照点都存在，计算它们之间的距离
                if ref1_position.get('exists', False) and ref2_position.get('exists', False):
                    # 计算两个参照点之间的垂直距离
                    distance = ref2_position['top'] - ref1_position['top']
                    
                    if distance > 0:
                        self._log_info(f"参照点之间的垂直距离: {distance}px", "blue")
                        
                        # 获取页面总高度和可视区域高度，以便计算滚动比例
                        page_metrics = self.driver.execute_script("""
                            return {
                                scrollHeight: document.documentElement.scrollHeight,
                                clientHeight: document.documentElement.clientHeight,
                                scrollTop: document.documentElement.scrollTop
                            };
                        """)
                        
                        # 计算滚动比例
                        scroll_height = page_metrics.get('scrollHeight', 3000)
                        client_height = page_metrics.get('clientHeight', 800)
                        
                        # 计算滚动条移动距离与实际页面移动距离的比例
                        if scroll_height > client_height:
                            scroll_ratio = (scroll_height - client_height) / scroll_height
                            # 使用参照点之间的距离作为滚动距离，考虑滚动条比例并增加一点额外距离
                            scroll_distance = distance * scroll_ratio * 1.2  # 增加20%的距离
                            self._log_info(f"页面高度: {scroll_height}px, 可视区域: {client_height}px, 滚动比例: {scroll_ratio:.2f}", "blue")
                        else:
                            # 如果页面不需要滚动，使用默认值
                            scroll_distance = distance * 0.8  # 增加滚动距离
                        
                        self._log_info(f"使用调整后的滚动距离: {scroll_distance}px", "blue")
                    else:
                        self._log_info(f"参照点之间的垂直距离为负或零: {distance}px，使用默认距离", "orange")
                else:
                    self._log_info("无法获取参照点位置，使用默认滚动距离", "orange")
        except Exception as e:
            self._log_info(f"计算参照点距离失败: {str(e)}", "orange")
        
        # 直接滚动页面，不尝试滚动容器
        self._log_info(f"直接滚动整个页面 {scroll_distance}px", "blue")
        
        # 使用更精确的滚动方法，考虑滚动条和页面的比例关系
        try:
            # 获取页面总高度和可视区域高度，以便计算滚动比例
            page_metrics = self.driver.execute_script("""
                return {
                    scrollHeight: document.documentElement.scrollHeight,
                    clientHeight: document.documentElement.clientHeight,
                    scrollTop: document.documentElement.scrollTop,
                    pageYOffset: window.pageYOffset
                };
            """)
            
            self._log_info(f"页面高度: {page_metrics.get('scrollHeight', 0)}px, 可视区域高度: {page_metrics.get('clientHeight', 0)}px, 当前滚动位置: {page_metrics.get('scrollTop', 0)}px", "blue")
            
            # 计算滚动步长，分多次小幅度滚动以确保页面内容能够正确加载
            total_scroll = scroll_distance
            steps = 2  # 减少滚动步骤，每步滚动更多
            step_distance = total_scroll / steps
            
            for i in range(steps):
                # 执行滚动
                result = self.driver.execute_script(f"""
                    const originalScrollTop = window.pageYOffset;
                    window.scrollBy(0, {step_distance});
                    return {{
                        before: originalScrollTop,
                        after: window.pageYOffset,
                        scrolled: window.pageYOffset - originalScrollTop
                    }};
                """)
                
                before_position = result.get('before', 0)
                after_position = result.get('after', 0)
                scrolled = result.get('scrolled', 0)
                
                self._log_info(f"滚动步骤 {i+1}/{steps}: 从 {before_position} 到 {after_position} (滚动: {scrolled}px)", "blue")
                
                # 短暂等待，让页面内容加载
                time.sleep(0.3)
            
            # 在调试模式下，保存滚动后的截图
            if hasattr(self, 'confirm_click') and self.confirm_click.get():
                self._save_screenshot("after_window_scroll")
                
            # 如果滚动距离太小，尝试一次更大的滚动
            if total_scroll < 100:
                self._log_info("滚动距离太小，尝试更大的滚动", "blue")
                result = self.driver.execute_script(f"""
                    const originalScrollTop = window.pageYOffset;
                    window.scrollBy(0, 200);
                    return {{
                        before: originalScrollTop,
                        after: window.pageYOffset,
                        scrolled: window.pageYOffset - originalScrollTop
                    }};
                """)
                
                before_position = result.get('before', 0)
                after_position = result.get('after', 0)
                scrolled = result.get('scrolled', 0)
                
                self._log_info(f"额外滚动: 从 {before_position} 到 {after_position} (滚动: {scrolled}px)", "blue")
            
            return True
        except Exception as e:
            self._log_info(f"精确滚动失败: {str(e)}", "orange")
        
        # 如果精确滚动失败，尝试常规滚动方法
        scroll_methods = [
            f"window.scrollBy(0, {scroll_distance}); return {{before: window.pageYOffset, after: window.pageYOffset + {scroll_distance}}};",
            f"const before = window.pageYOffset; window.scrollTo(0, before + {scroll_distance}); return {{before: before, after: window.pageYOffset}};",
            f"const before = document.documentElement.scrollTop; document.documentElement.scrollTop += {scroll_distance}; return {{before: before, after: document.documentElement.scrollTop}};",
            f"const before = document.body.scrollTop; document.body.scrollTop += {scroll_distance}; return {{before: before, after: document.body.scrollTop}};"
        ]
        
        for method in scroll_methods:
            try:
                result = self.driver.execute_script(method)
                before_position = result.get('before', 0)
                after_position = result.get('after', 0)
                
                if after_position > before_position:
                    self._log_info(f"已滚动窗口从 {before_position} 到 {after_position} (距离: {after_position - before_position}px)", "blue")
                    
                    # 在调试模式下，保存滚动后的截图
                    if hasattr(self, 'confirm_click') and self.confirm_click.get():
                        self._save_screenshot("after_window_scroll")
                        
                    return True
            except Exception as e:
                continue
        
        # 最后尝试使用键盘滚动
        try:
            actions = ActionChains(self.driver)
            actions.send_keys(Keys.PAGE_DOWN)
            actions.perform()
            self._log_info("使用键盘PAGE_DOWN滚动", "blue")
            
            # 在调试模式下，保存滚动后的截图
            if hasattr(self, 'confirm_click') and self.confirm_click.get():
                self._save_screenshot("after_pagedown_scroll")
                
            return True
        except Exception as e:
            self._log_info(f"键盘滚动失败: {str(e)}", "orange")
        
        self._log_info("所有窗口滚动方法都失败", "orange")
        return False
    
    def _swipe_with_pyautogui(self):
        """使用PyAutoGUI模拟滑动手势"""
        try:
            # 获取浏览器窗口位置和大小
            window_rect = self.driver.execute_script("""
                return {
                    left: window.screenX || window.screenLeft,
                    top: window.screenY || window.screenTop,
                    width: window.outerWidth,
                    height: window.innerHeight
                };
            """)
            
            # 计算滑动的起点和终点（从窗口中心开始，向上滑动）
            start_x = window_rect['left'] + window_rect['width'] // 2
            start_y = window_rect['top'] + window_rect['height'] // 2
            end_y = start_y - 300  # 向上滑动300像素
            
            # 执行滑动操作
            pyautogui.moveTo(start_x, start_y)
            pyautogui.mouseDown()
            pyautogui.moveTo(start_x, end_y, duration=0.5)
            pyautogui.mouseUp()
            
            self._log_info(f"已执行滑动: 从 ({start_x}, {start_y}) 到 ({start_x}, {end_y})", "blue")
        except Exception as e:
            self._log_info(f"PyAutoGUI滑动失败: {str(e)}", "red")
            raise
    
    def _scroll_with_keys(self):
        """使用键盘按键滚动页面"""
        try:
            # 确保页面有焦点
            self.driver.execute_script("window.focus();")
            
            # 发送Page Down键
            actions = ActionChains(self.driver)
            actions.send_keys(Keys.PAGE_DOWN)
            actions.perform()
            
            self._log_info("已发送Page Down键", "blue")
        except Exception as e:
            self._log_info(f"键盘滚动失败: {str(e)}", "red")
            raise
    
    def _click_next_page(self):
        """尝试点击"下一页"按钮"""
        try:
            # 尝试多种可能的"下一页"按钮选择器
            next_page_selectors = [
                "//a[contains(text(), '下一页')]",
                "//button[contains(text(), '下一页')]",
                "//span[contains(text(), '下一页')]",
                "//a[contains(@class, 'next')]",
                "//button[contains(@class, 'next')]",
                "//li[contains(@class, 'next')]",
                "//i[contains(@class, 'next')]"
            ]
            
            for selector in next_page_selectors:
                try:
                    elements = self.driver.find_elements(By.XPATH, selector)
                    if elements:
                        elements[0].click()
                        self._log_info(f"已点击下一页按钮: {selector}", "blue")
                        return
                except:
                    continue
            
            self._log_info("未找到下一页按钮", "orange")
            raise Exception("未找到下一页按钮")
        except Exception as e:
            self._log_info(f"点击下一页失败: {str(e)}", "red")
            raise
            
    def _extract_current_order_id(self):
        """提取当前页面上的订单ID"""
        # 检查是否已终止操作
        if not self.is_running:
            self._log_info("操作已终止，停止提取订单ID", "orange")
            return None
            
        import re  # 导入re模块解决"cannot access local variable 're'"问题
        try:
            # 策略1: 尝试从URL中提取
            url = self.driver.current_url
            url_match = re.search(r'order[_=-]?id=([A-Za-z0-9-]+)', url, re.IGNORECASE)
            if url_match:
                order_id = url_match.group(1)
                self._log_info(f"从URL成功提取订单ID: {order_id}", "green")
                print(f"DEBUG-URL-ID-EXTRACTED: '{order_id}'")
                
                # 保存到最近使用的ID列表
                if order_id not in self.last_order_ids:
                    self.last_order_ids.append(order_id)
                    if len(self.last_order_ids) > 5:
                        self.last_order_ids.pop(0)
                
                # 更新当前订单ID
                self.current_order_id = order_id
                
                return order_id
            
            # 策略2: 尝试多种可能的订单ID选择器
            order_id_selectors = [
                "//span[contains(text(), '订单编号')]/following-sibling::span",
                "//span[contains(text(), '订单编号')]",
                "//div[contains(text(), '订单编号')]",
                "//td[contains(text(), '订单编号')]",
                "//th[contains(text(), '订单编号')]/following-sibling::td",
                "//label[contains(text(), '订单编号')]/following-sibling::*"
            ]
            
            self._log_info("开始提取当前可见订单ID...", "blue")
            print("DEBUG-EXTRACT-ID: 开始提取当前可见订单ID")
            
            # 添加重试机制
            max_retries = 3
            for retry in range(max_retries):
                if retry > 0:
                    self._log_info(f"第{retry+1}次尝试提取订单ID...", "orange")
                    print(f"DEBUG-EXTRACT-RETRY: 第{retry+1}次尝试提取订单ID")
                    # 短暂等待页面可能的更新
                    time.sleep(0.5)
                
                # 获取当前视口内可见元素的信息
                visible_elements_js = """
                return Array.from(document.querySelectorAll('*')).filter(el => {
                    const rect = el.getBoundingClientRect();
                    return rect.top >= 0 && 
                           rect.left >= 0 && 
                           rect.bottom <= window.innerHeight && 
                           rect.right <= window.innerWidth &&
                           el.textContent.includes('订单编号');
                }).map(el => ({
                    text: el.textContent,
                    top: el.getBoundingClientRect().top,
                    visible: true
                }));
                """
                
                try:
                    visible_elements = self.driver.execute_script(visible_elements_js)
                    if visible_elements:
                        print(f"DEBUG-VISIBLE-ELEMENTS: 找到{len(visible_elements)}个可见元素包含'订单编号'")
                        for i, elem in enumerate(visible_elements):
                            print(f"DEBUG-VISIBLE-ELEMENT-{i+1}: {elem['text'][:50]}, top={elem['top']}")
                        
                        # 优先使用可见元素中最靠近视口中心的元素
                        if len(visible_elements) > 0:
                            # 按照元素距离视口顶部的距离排序
                            sorted_elements = sorted(visible_elements, key=lambda e: abs(e['top'] - (self.driver.execute_script("return window.innerHeight") / 2)))
                            center_element_text = sorted_elements[0]['text']
                            print(f"DEBUG-CENTER-ELEMENT: 选择最靠近视口中心的元素: '{center_element_text[:50]}'")
                            
                            # 从文本中提取订单ID
                            match = re.search(r'订单编号[：:]\s*([0-9a-zA-Z-]+)', center_element_text)
                            if match:
                                extracted_id = match.group(1)
                                self._log_info(f"从可见元素中提取到订单ID: '{extracted_id}'", "green")
                                print(f"DEBUG-VISIBLE-ID-EXTRACTED: '{extracted_id}'")
                                
                                # 保存到最近使用的ID列表
                                if extracted_id not in self.last_order_ids:
                                    self.last_order_ids.append(extracted_id)
                                    if len(self.last_order_ids) > 5:
                                        self.last_order_ids.pop(0)
                                
                                # 更新当前订单ID
                                self.current_order_id = extracted_id
                                
                                # 验证提取的ID格式是否合理
                                if len(extracted_id) >= 5:  # 订单ID通常较长
                                    return extracted_id
                                else:
                                    self._log_info(f"提取的订单ID '{extracted_id}' 格式可疑，尝试其他方法", "orange")
                except Exception as js_e:
                    print(f"DEBUG-JS-ERROR: 获取可见元素失败: {str(js_e)}")
                
                for selector in order_id_selectors:
                    try:
                        elements = self.driver.find_elements(By.XPATH, selector)
                        if elements:
                            order_id = elements[0].text
                            self._log_info(f"找到订单ID元素，原始文本: '{order_id}'", "blue")
                            print(f"DEBUG-ORDER-ID-ELEMENT: '{order_id}'")
                            
                            # 提取数字和字母部分作为订单ID
                            import re
                            match = re.search(r'[0-9a-zA-Z-]+', order_id)
                            if match:
                                extracted_id = match.group(0)
                                self._log_info(f"提取到订单ID: '{extracted_id}'", "green")
                                print(f"DEBUG-ORDER-ID-EXTRACTED: '{extracted_id}'")
                                
                                # 保存到最近使用的ID列表
                                if extracted_id not in self.last_order_ids:
                                    self.last_order_ids.append(extracted_id)
                                    if len(self.last_order_ids) > 5:
                                        self.last_order_ids.pop(0)
                                
                                # 更新当前订单ID
                                self.current_order_id = extracted_id
                                
                                # 验证提取的ID格式是否合理
                                if len(extracted_id) >= 5:  # 订单ID通常较长
                                    return extracted_id
                                else:
                                    self._log_info(f"提取的订单ID '{extracted_id}' 格式可疑，尝试下一个选择器", "orange")
                                    continue
                            
                            cleaned_id = order_id.strip()
                            self._log_info(f"清理后的订单ID: '{cleaned_id}'", "green")
                            print(f"DEBUG-ORDER-ID-CLEANED: '{cleaned_id}'")
                            
                            # 保存到最近使用的ID列表
                            if cleaned_id not in self.last_order_ids:
                                self.last_order_ids.append(cleaned_id)
                                if len(self.last_order_ids) > 5:
                                    self.last_order_ids.pop(0)
                            
                            # 更新当前订单ID
                            self.current_order_id = cleaned_id
                            
                            # 验证清理后的ID格式是否合理
                            if len(cleaned_id) >= 5:  # 订单ID通常较长
                                return cleaned_id
                            else:
                                self._log_info(f"清理后的订单ID '{cleaned_id}' 太短，尝试下一个选择器", "orange")
                    except Exception as e:
                        print(f"DEBUG-SELECTOR-ERROR: 选择器 '{selector}' 失败: {str(e)}")
            
            # 策略3: 尝试从页面源代码中查找订单ID模式
            page_source = self.driver.page_source
            patterns = [
                r'订单编号[：:]\s*([0-9a-zA-Z-]{5,})',
                r'订单号[：:]\s*([0-9a-zA-Z-]{5,})',
                r'订单[：:]\s*([0-9a-zA-Z-]{5,})',
                r'单号[：:]\s*([0-9a-zA-Z-]{5,})',
                r'order[_\s]?id[：:=]\s*([0-9a-zA-Z-]{5,})',
                r'order[_\s]?number[：:=]\s*([0-9a-zA-Z-]{5,})'
            ]
            
            for pattern in patterns:
                match = re.search(pattern, page_source, re.IGNORECASE)
                if match:
                    extracted_id = match.group(1)
                    self._log_info(f"从页面源码提取订单ID: '{extracted_id}'", "green")
                    print(f"DEBUG-SOURCE-ID-EXTRACTED: '{extracted_id}'")
                    
                    # 保存到最近使用的ID列表
                    if extracted_id not in self.last_order_ids:
                        self.last_order_ids.append(extracted_id)
                        if len(self.last_order_ids) > 5:
                            self.last_order_ids.pop(0)
                    
                    # 更新当前订单ID
                    self.current_order_id = extracted_id
                    
                    return extracted_id
            
            # 策略4: 如果有当前缓存的订单ID，使用它
            if hasattr(self, 'current_order_id') and self.current_order_id:
                self._log_info(f"使用当前缓存的订单ID: '{self.current_order_id}'", "orange")
                print(f"DEBUG-CACHED-ID-USED: '{self.current_order_id}'")
                return self.current_order_id
            
            self._log_info("无法提取订单ID", "red")
            print("DEBUG-NO-ID-EXTRACTED")
            return None
        except Exception as e:
            self._log_info(f"提取订单ID时发生异常: {str(e)}", "red")
            print(f"DEBUG-EXTRACT-ID-ERROR: {str(e)}")
            
            # 在异常情况下，如果有当前缓存的订单ID，使用它
            if hasattr(self, 'current_order_id') and self.current_order_id:
                self._log_info(f"异常情况下使用当前缓存的订单ID: '{self.current_order_id}'", "orange")
                print(f"DEBUG-ERROR-CACHED-ID-USED: '{self.current_order_id}'")
                return self.current_order_id
                
            return None
    
    def _save_screenshot(self, tag):
        """保存屏幕截图用于调试 - 已禁用"""
        # 不再保存截图
        pass
    
    def _load_elements_from_json(self, file_path=None):
        """从JSON文件加载元素数据"""
        if not file_path:
            file_path = filedialog.askopenfilename(
                title="选择元素配置文件",
                filetypes=[("JSON文件", "*.json")],
                initialdir=os.path.dirname(os.path.abspath(__file__))
            )
            
        if not file_path:
            return None
            
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                elements_data = json.load(f)
                
            # 处理并验证JSON数据
            processed_data = []
            for i, elem in enumerate(elements_data):
                # 转换为内部数据结构，添加默认顺序和启用状态
                processed_data.append({
                    "element_id": i,
                    "name": elem.get("custom_name", f"元素{i+1}"),
                    "action": elem.get("action", "getText"),
                    "xpath": elem.get("xpath", ""),
                    "order": i+1,  # 默认顺序为加载顺序
                    "enabled": True  # 默认启用
                })
            return processed_data
        except Exception as e:
            self._log_info(f"加载元素数据失败: {str(e)}", "red")
            return None
    
    def _configure_operations(self):
        """显示操作选择与排序对话框"""
        if not self.is_browser_connected:
            self._log_info("浏览器未连接，无法配置操作", "red")
            return
            
        elements_data = self._load_elements_from_json()
        if not elements_data:
            self._log_info("未找到有效的元素数据，请先使用采集模式收集元素", "red")
            return
            
        dialog = OperationSequenceDialog(self.root, elements_data)
        if dialog.result:
            self.operation_sequence = dialog.result
            # 按顺序排序
            self.operation_sequence.sort(key=lambda x: x["order"])
            # 仅保留启用的操作
            self.operation_sequence = [op for op in self.operation_sequence if op["enabled"]]
            self._log_info(f"已配置{len(self.operation_sequence)}个操作，准备就绪", "green")
            # 启用开始按钮
            self.start_button.config(state=tk.NORMAL)
            
    def _preview_element(self, xpath):
        """高亮显示指定XPath的元素，用于预览"""
        if not self.driver:
            self._log_info("浏览器未连接，无法预览元素", "red")
            return False
            
        try:
            # 查找元素
            element = self.driver.find_element(By.XPATH, xpath)
            
            # 使用JavaScript高亮元素
            script = """
            function highlightElement(element) {
                var originalStyle = element.getAttribute('style');
                element.setAttribute('style', originalStyle + '; border: 2px solid red !important; background-color: yellow !important;');
                setTimeout(function() {
                    element.setAttribute('style', originalStyle);
                }, 2000);
            }
            arguments[0].scrollIntoView({behavior: "smooth", block: "center"});
            highlightElement(arguments[0]);
            """
            self.driver.execute_script(script, element)
            return True
        except Exception as e:
            self._log_info(f"预览元素失败: {str(e)}", "red")
            return False

    def _get_order_count(self):
        """从页面获取待处理订单数量"""
        if not self.driver:
            return 0
            
        # 如果操作序列中包含了订单数量元素，使用它获取数量
        order_count_elements = [op for op in self.operation_sequence if op.get("is_order_count", False)]
        
        if order_count_elements:
            try:
                # 使用指定的XPath获取订单数量元素
                op = order_count_elements[0]
                xpath = op["xpath"]
                element = self.driver.find_element(By.XPATH, xpath)
                text = element.text
                
                # 提取数字
                import re
                numbers = re.findall(r'\d+', text)
                if numbers:
                    count = int(numbers[0])
                    self._log_info(f"自动获取订单数量: {count}", "green")
                    return count
            except Exception as e:
                self._log_info(f"获取订单数量失败: {str(e)}", "red")
        
        # 如果没有指定订单数量元素或获取失败，询问用户
        count = simpledialog.askinteger(
            "输入订单数量", 
            "请手动输入本次要处理的订单数量:", 
            minvalue=1, 
            maxvalue=1000
        )
        if count:
            self._log_info(f"手动输入订单数量: {count}", "blue")
            return count
        
        # 默认值
        default_count = 5
        self._log_info(f"使用默认订单数量: {default_count}", "orange")
        return default_count

    def _update_always_on_top(self):
        """更新窗口置顶状态"""
        if self.always_on_top.get():
            self.root.attributes('-topmost', True)
        else:
            self.root.attributes('-topmost', False)
    
    def _reset_offsets(self):
        """重置所有元素的偏移量"""
        # 重置为空字典或默认值
        self.element_offsets = {elem: {"x": 0, "y": 0} for elem in ["订单编号", "商品名称", "成交金额", "查看1", "查看2", "复制完整收货信息"]}
        
        # 更新显示
        if hasattr(self, 'offset_label') and self.offset_label:
            self.offset_label.config(text="元素偏移量: 已重置")
        
        # 保存到配置文件
        self._save_offset_config()
        self._log_info("已重置所有元素的偏移量", "blue")
    
    def _load_offset_config(self):
        """加载元素偏移量配置"""
        try:
            config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "offset_config.json")
            if os.path.exists(config_path):
                with open(config_path, 'r', encoding='utf-8') as f:
                    config = json.load(f)
                    
                # 兼容旧版配置
                if "offset" in config and "x" in config["offset"] and "y" in config["offset"]:
                    # 旧格式，创建默认字典，并填入旧的偏移量
                    default_x = config["offset"]["x"]
                    default_y = config["offset"]["y"]
                    self._log_info(f"已从配置文件加载偏移量: X={default_x}, Y={default_y}", "blue")
                    # 初始化元素偏移量字典
                    self.element_offsets = {elem: {"x": default_x, "y": default_y} for elem in ["订单编号", "商品名称", "成交金额", "查看1", "查看2", "复制完整收货信息"]}
                elif "element_offsets" in config:
                    # 新格式，直接加载
                    self.element_offsets = config["element_offsets"]
                    elements_loaded = len(self.element_offsets)
                    self._log_info(f"已加载{elements_loaded}个元素的偏移量配置", "blue")
                    # 显示部分元素的偏移量
                    for name, offset in list(self.element_offsets.items())[:3]:
                        self._log_info(f"元素'{name}'偏移量: X={offset['x']}, Y={offset['y']}", "blue")
                        
            else:
                # 文件不存在，创建默认字典
                self.element_offsets = {elem: {"x": 0, "y": 0} for elem in ["订单编号", "商品名称", "成交金额", "查看1", "查看2", "复制完整收货信息"]}
                self._log_info("偏移量配置文件不存在，已创建默认配置", "blue")
                # 保存默认配置
                self._save_offset_config()
        except Exception as e:
            self.element_offsets = {elem: {"x": 0, "y": 0} for elem in ["订单编号", "商品名称", "成交金额", "查看1", "查看2", "复制完整收货信息"]}
            self._log_info(f"加载偏移量配置失败: {str(e)}，使用默认值", "orange")
    
    def _save_offset_config(self):
        """保存元素偏移量配置"""
        try:
            config = {
                "element_offsets": self.element_offsets,
                "updated_at": time.strftime("%Y-%m-%d %H:%M:%S")
            }
            
            config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "offset_config.json")
            with open(config_path, 'w', encoding='utf-8') as f:
                json.dump(config, f, ensure_ascii=False, indent=2)
                
            self._log_info(f"已保存{len(self.element_offsets)}个元素的偏移量配置", "green")
            
        except Exception as e:
            self._log_info(f"保存偏移量配置失败: {str(e)}", "red")

    def _show_offset_manager(self):
        """显示元素偏移量管理界面"""
        # 创建偏移量管理窗口
        offset_win = tk.Toplevel(self.root)
        offset_win.title("元素偏移量管理")
        offset_win.geometry("500x400")
        offset_win.transient(self.root)
        offset_win.grab_set()
        
        # 创建表格显示所有元素的偏移量
        columns = ("元素名称", "X偏移", "Y偏移", "操作")
        tree = ttk.Treeview(offset_win, columns=columns, show="headings", selectmode="browse")
        
        # 设置列标题
        for col in columns:
            tree.heading(col, text=col)
            if col == "元素名称":
                tree.column(col, width=200, anchor="w")
            elif col in ("X偏移", "Y偏移"):
                tree.column(col, width=80, anchor="center")
            else:
                tree.column(col, width=100, anchor="center")
        
        # 添加滚动条
        scrollbar = ttk.Scrollbar(offset_win, orient="vertical", command=tree.yview)
        tree.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side="right", fill="y")
        tree.pack(fill="both", expand=True, padx=10, pady=10)
        
        # 填充数据
        def refresh_tree():
            # 清空现有数据
            tree.delete(*tree.get_children())
            
            # 添加元素数据
            for name, offset in self.element_offsets.items():
                tree.insert("", "end", values=(name, offset.get("x", 0), offset.get("y", 0), "重置"))
        
        refresh_tree()
        
        # 创建按钮框架
        btn_frame = ttk.Frame(offset_win)
        btn_frame.pack(fill="x", padx=10, pady=10)
        
        # 添加重置所有按钮
        reset_all_btn = ttk.Button(
            btn_frame, 
            text="重置所有偏移量", 
            command=lambda: [self._reset_offsets(), refresh_tree()]
        )
        reset_all_btn.pack(side="left", padx=5)
        
        # 添加导出按钮
        export_btn = ttk.Button(
            btn_frame, 
            text="导出配置", 
            command=self._export_offset_config
        )
        export_btn.pack(side="left", padx=5)
        
        # 添加导入按钮
        import_btn = ttk.Button(
            btn_frame, 
            text="导入配置", 
            command=lambda: [self._import_offset_config(), refresh_tree()]
        )
        import_btn.pack(side="left", padx=5)
        
        # 添加关闭按钮
        close_btn = ttk.Button(
            btn_frame, 
            text="关闭", 
            command=offset_win.destroy
        )
        close_btn.pack(side="right", padx=5)
        
        # 处理点击"重置"操作
        def on_tree_click(event):
            # 获取点击的行和列
            item = tree.identify_row(event.y)
            column = tree.identify_column(event.x)
            
            if item and column == "#4":  # 操作列
                # 获取元素名称
                values = tree.item(item, "values")
                elem_name = values[0]
                
                # 重置该元素的偏移量
                if elem_name in self.element_offsets:
                    self.element_offsets[elem_name] = {"x": 0, "y": 0}
                    self._save_offset_config()
                    self._log_info(f"已重置元素'{elem_name}'的偏移量", "blue")
                    refresh_tree()
        
        tree.bind("<Button-1>", on_tree_click)
    
    def _export_offset_config(self):
        """导出元素偏移量配置到文件"""
        try:
            file_path = filedialog.asksaveasfilename(
                title="导出偏移量配置",
                defaultextension=".json",
                filetypes=[("JSON文件", "*.json")],
                initialfile="element_offsets_config.json"
            )
            
            if not file_path:
                return
                
            config = {
                "element_offsets": self.element_offsets,
                "exported_at": time.strftime("%Y-%m-%d %H:%M:%S"),
                "description": "元素偏移量配置文件"
            }
            
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(config, f, ensure_ascii=False, indent=2)
                
            self._log_info(f"已导出偏移量配置到: {file_path}", "green")
            
        except Exception as e:
            self._log_info(f"导出偏移量配置失败: {str(e)}", "red")
            
    def _import_offset_config(self):
        """从文件导入元素偏移量配置"""
        try:
            file_path = filedialog.askopenfilename(
                title="导入偏移量配置",
                filetypes=[("JSON文件", "*.json")],
                initialdir=os.path.dirname(os.path.abspath(__file__))
            )
            
            if not file_path:
                return
                
            with open(file_path, 'r', encoding='utf-8') as f:
                config = json.load(f)
                
            if "element_offsets" in config:
                # 保存原始配置以防导入失败
                original_offsets = copy.deepcopy(self.element_offsets)
                
                try:
                    # 验证导入的配置格式
                    for name, offset in config["element_offsets"].items():
                        if not isinstance(name, str) or not isinstance(offset, dict):
                            raise ValueError("配置格式错误")
                        if "x" not in offset or "y" not in offset:
                            raise ValueError(f"元素'{name}'的偏移量格式错误")
                            
                    # 更新配置
                    self.element_offsets = config["element_offsets"]
                    self._save_offset_config()
                    self._log_info(f"已导入偏移量配置，包含{len(self.element_offsets)}个元素", "green")
                    
                except Exception as e:
                    # 恢复原始配置
                    self.element_offsets = original_offsets
                    self._log_info(f"导入偏移量配置验证失败: {str(e)}", "red")
            else:
                self._log_info("无效的偏移量配置文件", "red")
                
        except Exception as e:
            self._log_info(f"导入偏移量配置失败: {str(e)}", "red")

    def _collect_ref1_xpath(self):
        self._log_info("请将鼠标悬停在第1个订单的参照元素上，然后按'.'键采集", "blue")
        self._pending_collect = 'ref1'

    def _collect_ref2_xpath(self):
        self._log_info("请将鼠标悬停在第2个订单的参照元素上，然后按'.'键采集", "blue")
        self._pending_collect = 'ref2'

    def _collect_scroll_container_xpath(self):
        self._log_info("请将鼠标悬停在滚动容器上，然后按'.'键采集", "blue")
        self._pending_collect = 'scroll_container'

    def _learn_xpath_pattern(self, xpath1, xpath2):
        """参考脚本的循环XPath推算：找到唯一递增索引位置，允许中间层级差异"""
        import re
        
        # 完全相同的XPath，无法学习
        if xpath1 == xpath2:
            self._log_info('两个参照XPath完全相同，无法学习循环模式。', 'red')
            return None
            
        parts1 = xpath1.split('/')
        parts2 = xpath2.split('/')
        
        # 处理不同长度的XPath
        if len(parts1) != len(parts2):
            self._log_info(f'参照XPath层级不同: 第1个有{len(parts1)}层，第2个有{len(parts2)}层', 'orange')
            # 尝试从后向前匹配，找到共同部分
            min_len = min(len(parts1), len(parts2))
            parts1 = parts1[-min_len:]
            parts2 = parts2[-min_len:]
            self._log_info(f'尝试使用后{min_len}层进行匹配', 'blue')
            
        # 查找所有不同的部分
        diff_segments = []
        for i in range(len(parts1)):
            if parts1[i] != parts2[i]:
                diff_segments.append((i, parts1[i], parts2[i]))
                
        # 没有差异，不应该发生，因为之前已经检查了完全相同的情况
        if not diff_segments:
            self._log_info('截取后的XPath没有差异，无法学习。', 'red')
            return None
            
        # 筛选出包含数字索引的差异部分
        index_diffs = []
        for idx, part1, part2 in diff_segments:
            match1 = re.search(r'\[(\d+)\]', part1)
            match2 = re.search(r'\[(\d+)\]', part2)
            if match1 and match2:
                index1 = int(match1.group(1))
                index2 = int(match2.group(1))
                if index2 > index1:  # 确保索引是递增的
                    index_diffs.append((idx, index1, index2, part1, part2))
                    
        # 如果没有找到递增的数字索引差异
        if not index_diffs:
            # 尝试找出任何包含数字索引的部分
            for idx, part1, part2 in diff_segments:
                match1 = re.search(r'\[(\d+)\]', part1)
                match2 = re.search(r'\[(\d+)\]', part2)
                if match1 or match2:  # 只要有一个包含索引
                    if match1 and not match2:
                        self._log_info(f'XPath差异: 第1个有索引[{match1.group(1)}]，第2个没有索引', 'orange')
                    elif match2 and not match1:
                        self._log_info(f'XPath差异: 第1个没有索引，第2个有索引[{match2.group(1)}]', 'orange')
                    else:
                        self._log_info(f'XPath索引不是递增的: {match1.group(1)} -> {match2.group(1)}', 'orange')
            
            self._log_info('未找到有效的递增索引模式，尝试使用最后一个数字索引作为模式。', 'orange')
            
            # 查找最后一个包含数字索引的部分
            last_index_part = None
            last_index_value = None
            last_index_idx = -1
            
            for i, part in enumerate(parts1):
                match = re.search(r'\[(\d+)\]', part)
                if match:
                    last_index_part = part
                    last_index_value = int(match.group(1))
                    last_index_idx = i
                    
            if last_index_part:
                self._log_info(f'使用最后一个索引 [{last_index_value}] 作为基准', 'blue')
                template = re.sub(r'\[\d+\]', '[{}]', last_index_part, count=1)
                return {'diff_segment_index': last_index_idx, 'template': template, 'start_index': last_index_value}
            else:
                self._log_info('XPath中没有找到任何数字索引，无法学习。', 'red')
                return None
                
        # 如果找到了多个递增索引差异，使用最可能的那个（通常是第一个）
        if len(index_diffs) > 1:
            self._log_info(f'发现{len(index_diffs)}处递增索引差异，使用第一处作为模式。', 'orange')
            
        # 使用找到的差异生成模式
        diff_idx, start_index, next_index, part1, part2 = index_diffs[0]
        template = re.sub(r'\[\d+\]', '[{}]', part1, count=1)
        self._log_info(f'已学习到XPath模式: 在第{diff_idx}段，索引从{start_index}递增到{next_index}', 'green')
        
        return {'diff_segment_index': diff_idx, 'template': template, 'start_index': start_index}

    def _generate_xpath_for_item(self, base_xpath, loop_counter, pattern):
        """参考脚本的循环XPath生成"""
        import re
        parts = base_xpath.split('/')
        if pattern and 'diff_segment_index' in pattern:
            idx = pattern['diff_segment_index']
            if len(parts) > idx:
                actual_dom_index = pattern['start_index'] + loop_counter - 1
                original_segment = parts[idx]
                if re.search(r'\[\d+\]', original_segment):
                    parts[idx] = re.sub(r'\[\d+\]', f'[{actual_dom_index}]', original_segment, count=1)
                    return '/'.join(parts)
        return base_xpath

    def run_actions_loop(self, manual_order_count=None):
        """参考脚本的循环主流程，兼容main.py操作配置"""
        if not self.driver:
            self._log_info('循环模式错误: 浏览器未连接。', 'red')
            return
            
        # 确保order_clipboard_contents字典已初始化
        if not hasattr(self, 'order_clipboard_contents'):
            self.order_clipboard_contents = {}
            self._log_info("初始化订单ID与收货信息映射字典", "blue")
            print("DEBUG-RUN-ACTIONS-LOOP: 初始化订单ID与收货信息映射字典")
        
        # 检查是否有订单数量元素被标记为自动检测
        order_count_elements = [op for op in self.operation_sequence if op.get("is_order_count", False)]
        
        if order_count_elements:
            # 有订单数量元素，使用自动检测模式
            count_action = order_count_elements[0]
            num_items = 0
            try:
                self._log_info(f"执行: {count_action['name']} (自动获取总数)")
                # 使用智能查找元素
                element = self._find_element_smart(count_action['name'], count_action['xpath'])
                if not element:
                    self._log_info('无法找到总数元素，请检查XPath。', 'red')
                    return
                    
                count_text = element.text.strip()
                import re
                numbers = re.findall(r'\d+', count_text)
                if not numbers:
                    self._log_info(f'错误: 在 "{count_text}" 中未找到数字。', 'red')
                    return
                num_items = int(numbers[0])
                self._log_info(f'自动检测到项目总数: {num_items}')
            except Exception as e:
                self._log_info(f'自动获取总数失败: {e}', 'red')
                return
            if num_items == 0:
                self._log_info('项目总数为0，无需执行。')
                return
            
            # 过滤掉订单数量元素，只处理其他元素
            actions_to_loop = [op for op in self.operation_sequence if not op.get("is_order_count", False)]
        else:
            # 没有订单数量元素，使用手动输入模式
            self._log_info("未选择订单数量元素，使用手动输入的订单数量", "blue")
            if manual_order_count is None:
                self._log_info('错误：未提供手动输入的订单数量', 'red')
                return
            num_items = manual_order_count
            self._log_info(f"使用手动输入的订单数量: {num_items}", "blue")
            
            # 跳过名为"待发货的数量"的元素，处理其他所有元素
            actions_to_loop = [op for op in self.operation_sequence if op.get('name') != '待发货的数量']
            self._log_info(f"手动输入模式：跳过'待发货的数量'元素，将处理{len(actions_to_loop)}个其他元素", "blue")
            
        # 限制处理数量（如果启用了调试模式）
        if hasattr(self, 'confirm_click') and self.confirm_click.get():
            debug_limit = 3  # 调试模式下限制处理前3个订单
            if num_items > debug_limit:
                self._log_info(f'调试模式已启用，仅处理前{debug_limit}个订单。', 'blue')
                num_items = debug_limit
                
        # 确保actions_to_loop不为空
        if not actions_to_loop:
            self._log_info('错误：没有可执行的操作元素', 'red')
            return
            
        first_action_xpath = actions_to_loop[0]['xpath']
        xpath_pattern = None
        
        # 尝试使用参照点学习XPath模式
        if hasattr(self, 'ref2_xpath') and self.ref2_xpath:
            xpath_pattern = self._learn_xpath_pattern(first_action_xpath, self.ref2_xpath)
            if not xpath_pattern:
                self._log_info('无法从参照XPath中学习到规律，将回退到基本模式。', 'orange')
                
        # 如果无法从参照点学习，尝试基本模式
        if not xpath_pattern:
            self._log_info('使用基本循环模式（递增最后一个数字）...', 'info')
            import re
            base_index = -1
            matches = list(re.finditer(r'\[(\d+)\]', first_action_xpath))
            if matches:
                base_index = int(matches[-1].group(1))
                self._log_info(f'检测到列表起始索引为: {base_index}')
                xpath_pattern = {'diff_segment_index': len(first_action_xpath.split('/'))-1, 'template': '', 'start_index': base_index}
            else:
                self._log_info('循环模式错误: 无法在第一个操作的XPath中找到列表索引（如 [1], [2]）。无法继续循环。', 'red')
                return
                
        # 设置进度条最大值
        self.progress_bar["maximum"] = num_items
        self.progress_bar["value"] = 0
        self.progress_label.config(text=f"0/{num_items}")
        self._log_info(f"设置进度条最大值为: {num_items}", "blue")
        
        self.collected_data = []
        processed_order_ids = set()  # 用于检测重复订单
        consecutive_same_order = 0   # 连续重复订单计数
        
        for i in range(1, num_items+1):
            if not self.is_running:
                break
                
            self._log_info(f"[循环] 正在处理第 {i}/{num_items} 个订单", "blue")
            order_data = {}
            
            # 处理当前订单的所有操作
            for op in actions_to_loop:
                # 首先检查是否已终止操作
                if not self.is_running:
                    self._log_info("操作已终止，停止处理当前订单", "orange")
                    break
                    
                # 检查验证码检测状态
                if hasattr(self, 'force_stop_flag') and self.force_stop_flag:
                    self._log_info("检测到验证码，暂停操作执行", "red")
                    # 等待验证码消失
                    while hasattr(self, 'force_stop_flag') and self.force_stop_flag and self.is_running:
                        time.sleep(0.5)
                        self.root.update()
                    if not self.is_running:
                        break
                    self._log_info("验证码已消失，继续执行操作", "green")
                
                # 特殊处理"查看2"元素，先微弱下滑几个像素，再执行原地点击而不查找元素
                if op['name'] == "查看2":
                    self._log_info(f"特殊处理元素: '查看2'，跳过查找直接执行原地点击", "blue")
                    # 确保浏览器窗口有焦点
                    self._switch_focus_to_browser()
                    time.sleep(0.3)
                    
                    # 先执行微弱下滑几个像素
                    try:
                        # 使用JavaScript执行微弱下滑
                        self.driver.execute_script("window.scrollBy(0, 5);")
                        self._log_info(f"已执行微弱下滑5像素", "blue")
                        time.sleep(0.2)  # 给滚动一点时间生效
                    except Exception as e:
                        self._log_info(f"微弱下滑失败: {str(e)}", "orange")
                    
                    # 执行原地点击
                    pyautogui.click()
                    self._log_info(f"已执行'查看2'的原地点击", "blue")
                    time.sleep(0.5)
                    
                    # 恢复焦点到采集工具窗口
                    self._manage_focus()
                    
                    # 添加延迟
                    if not self.confirm_click.get():
                        time.sleep(self.auto_action_interval)
                        self._log_info(f"自动执行模式：操作间延迟 {self.auto_action_interval} 秒", "blue")
                    continue
                
                # 正常处理其他元素
                op_xpath = op.get('xpath', '')
                op_item_xpath = self._generate_xpath_for_item(op_xpath, i, xpath_pattern)
                op_copy = op.copy()
                op_copy['xpath'] = op_item_xpath
                try:
                    result = self._execute_operation(op_copy)
                    if result is not None:
                        order_data[op['name']] = result
                    # 如果"点击前确认"未勾选，每个操作后添加延迟
                    if not self.confirm_click.get():
                        time.sleep(self.auto_action_interval)
                        self._log_info(f"自动执行模式：操作间延迟 {self.auto_action_interval} 秒", "blue")
                except Exception as e:
                    self._log_info(f"执行'{op['name']}'操作失败: {str(e)}", "red")
            
            # 检查是否成功采集了订单数据
            if order_data:
                # 检查是否是重复订单
                current_order_id = None
                for key in ['订单编号', '订单ID', 'order_id', 'orderid']:
                    if key in order_data:
                        current_order_id = order_data[key]
                        break
                
                if current_order_id:
                    if current_order_id in processed_order_ids:
                        consecutive_same_order += 1
                        self._log_info(f"检测到重复订单: {current_order_id}，这是第{consecutive_same_order}次重复", "orange")
                        
                        # 如果连续3次重复，尝试不同的滚动策略
                        if consecutive_same_order >= 3:
                            self._log_info("连续多次重复订单，尝试不同的滚动策略", "red")
                            # 检查是否已终止操作
                            if not self.is_running:
                                self._log_info("操作已终止，停止处理重复订单", "orange")
                                break
                            if not self._scroll_to_next_order():
                                self._log_info("所有滚动策略都失败，无法继续处理", "red")
                                break
                            consecutive_same_order = 0  # 重置计数器
                            continue  # 跳过当前订单，重新尝试
                    else:
                        processed_order_ids.add(current_order_id)
                        consecutive_same_order = 0  # 重置计数器
                        self.collected_data.append(order_data)
                else:
                    # 没有找到订单ID，但仍然添加数据
                    self.collected_data.append(order_data)
            
            # 更新进度条
            self.progress_bar["value"] = i
            self.progress_label.config(text=f"{i}/{num_items}")
            self.root.update()
            
            # 如果不是最后一个订单，滚动到下一个
            if i < num_items:
                # 检查是否已终止操作
                if not self.is_running:
                    self._log_info("操作已终止，停止滚动到下一个订单", "orange")
                    break
                    
                if not self._scroll_to_next_order():
                    self._log_info("无法滚动到下一个订单，尝试继续处理", "orange")
                    # 即使滚动失败，也尝试继续处理
                
                # 等待页面加载前再次检查状态
                if not self.is_running:
                    self._log_info("操作已终止，停止等待页面加载", "orange")
                    break
                time.sleep(1.5)
        
        self._log_info(f"[循环] 已完成所有 {len(self.collected_data)} 个订单的处理", "green")
        self._stop_collection()
        if len(self.collected_data) > 0:
            self.excel_button.config(state=tk.NORMAL)
            self.word_button.config(state=tk.NORMAL)

    def _find_element_smart(self, name, original_xpath):
        """智能元素查找，使用多种策略定位元素"""
        if not self.driver:
            self._log_info(f"浏览器未连接，无法查找元素 '{name}'", "red")
            return None
            
        self._log_info(f"智能查找元素: '{name}'", "blue")
        element = None
        
        # 策略1: 使用原始XPath
        try:
            element = self.driver.find_element(By.XPATH, original_xpath)
            self._log_info(f"使用原始XPath找到元素 '{name}'", "green")
            # 给元素添加一个属性，标记它的名称，用于后续应用偏移量
            self.driver.execute_script("arguments[0].setAttribute('data-element-name', arguments[1]);", element, name)
            
            # 检查元素位置
            rect = self.driver.execute_script('return arguments[0].getBoundingClientRect();', element)
            if rect['width'] <= 0 or rect['height'] <= 0:
                self._log_info(f"警告: 使用原始XPath找到的元素'{name}'尺寸异常: width={rect['width']}, height={rect['height']}", "orange")
                # 尝试滚动到元素
                self.driver.execute_script("arguments[0].scrollIntoView({behavior: 'auto', block: 'center'});", element)
                time.sleep(0.5)
            
            return element
        except Exception as e:
            self._log_info(f"原始XPath未找到元素 '{name}': {str(e)}", "orange")
        
        # 策略2: 使用相对XPath
        try:
            # 尝试生成更健壮的相对XPath
            relative_xpath = self._generate_relative_xpath(original_xpath)
            if relative_xpath:
                element = self.driver.find_element(By.XPATH, relative_xpath)
                self._log_info(f"使用相对XPath找到元素 '{name}'", "green")
                # 给元素添加一个属性，标记它的名称，用于后续应用偏移量
                self.driver.execute_script("arguments[0].setAttribute('data-element-name', arguments[1]);", element, name)
                
                # 检查元素位置并确保元素可见
                rect = self.driver.execute_script('return arguments[0].getBoundingClientRect();', element)
                self._log_info(f"相对XPath找到的元素'{name}'位置: left={rect['left']}, top={rect['top']}, width={rect['width']}, height={rect['height']}", "blue")
                
                if rect['width'] <= 0 or rect['height'] <= 0 or rect['left'] < 0 or rect['top'] < 0:
                    self._log_info(f"警告: 使用相对XPath找到的元素'{name}'位置或尺寸异常，尝试滚动使其可见", "orange")
                    # 尝试滚动到元素
                    self.driver.execute_script("arguments[0].scrollIntoView({behavior: 'auto', block: 'center'});", element)
                    time.sleep(0.5)
                    
                    # 重新获取位置
                    rect = self.driver.execute_script('return arguments[0].getBoundingClientRect();', element)
                    self._log_info(f"滚动后元素'{name}'位置: left={rect['left']}, top={rect['top']}, width={rect['width']}, height={rect['height']}", "blue")
                
                # 尝试获取原始元素的位置和尺寸作为参考
                try:
                    # 如果原始元素存在，获取其位置信息作为参考
                    original_elements = self.driver.find_elements(By.XPATH, original_xpath)
                    if original_elements:
                        original_rect = self.driver.execute_script('return arguments[0].getBoundingClientRect();', original_elements[0])
                        self._log_info(f"原始XPath元素位置参考: left={original_rect['left']}, top={original_rect['top']}, width={original_rect['width']}, height={original_rect['height']}", "blue")
                        
                        # 如果相对XPath找到的元素位置异常，但原始元素位置正常，尝试使用原始元素
                        if (rect['width'] <= 0 or rect['height'] <= 0) and original_rect['width'] > 0 and original_rect['height'] > 0:
                            self._log_info(f"使用原始XPath元素作为备选，因为相对XPath元素位置异常", "orange")
                            return original_elements[0]
                except Exception:
                    pass
                
                return element
        except Exception as e:
            self._log_info(f"相对XPath未找到元素 '{name}': {str(e)}", "orange")
        
        # 策略3: 使用文本内容查找
        try:
            element = self._find_by_text_content(name)
            if element:
                self._log_info(f"使用文本内容找到元素 '{name}'", "green")
                # 给元素添加一个属性，标记它的名称，用于后续应用偏移量
                self.driver.execute_script("arguments[0].setAttribute('data-element-name', arguments[1]);", element, name)
                
                # 检查元素位置
                rect = self.driver.execute_script('return arguments[0].getBoundingClientRect();', element)
                if rect['width'] <= 0 or rect['height'] <= 0:
                    self._log_info(f"警告: 使用文本内容找到的元素'{name}'尺寸异常: width={rect['width']}, height={rect['height']}", "orange")
                    # 尝试滚动到元素
                    self.driver.execute_script("arguments[0].scrollIntoView({behavior: 'auto', block: 'center'});", element)
                    time.sleep(0.5)
                
                return element
        except Exception as e:
            self._log_info(f"文本内容未找到元素 '{name}': {str(e)}", "orange")
        
        # 策略4: 使用CSS选择器
        try:
            # 尝试从XPath转换为CSS选择器
            css_selector = self._xpath_to_css(original_xpath)
            if css_selector:
                element = self.driver.find_element(By.CSS_SELECTOR, css_selector)
                self._log_info(f"使用CSS选择器找到元素 '{name}'", "green")
                # 给元素添加一个属性，标记它的名称，用于后续应用偏移量
                self.driver.execute_script("arguments[0].setAttribute('data-element-name', arguments[1]);", element, name)
                
                # 检查元素位置
                rect = self.driver.execute_script('return arguments[0].getBoundingClientRect();', element)
                if rect['width'] <= 0 or rect['height'] <= 0:
                    self._log_info(f"警告: 使用CSS选择器找到的元素'{name}'尺寸异常: width={rect['width']}, height={rect['height']}", "orange")
                    # 尝试滚动到元素
                    self.driver.execute_script("arguments[0].scrollIntoView({behavior: 'auto', block: 'center'});", element)
                    time.sleep(0.5)
                
                return element
        except Exception as e:
            self._log_info(f"CSS选择器未找到元素 '{name}': {str(e)}", "orange")
        
        # 策略5: 使用JavaScript查找
        try:
            js_script = f"""
            function findElementByContent(text) {{
                const elements = Array.from(document.querySelectorAll('*'));
                return elements.find(el => {{
                    const content = el.textContent || '';
                    return content.includes(text) && content.length < 100;
                }});
            }}
            
            const element = findElementByContent('{name}');
            if (element) {{
                element.style.border = '2px solid red';
                element.scrollIntoView({{behavior: 'smooth', block: 'center'}});
                element.setAttribute('data-element-name', '{name}');
                return true;
            }}
            return false;
            """
            
            found = self.driver.execute_script(js_script)
            if found:
                # 使用JavaScript找到元素后，我们需要重新定位它以返回WebElement对象
                # 这里使用一个简单的策略：查找带有红色边框的元素
                element = self.driver.find_element(By.XPATH, "//*[contains(@style, 'border: 2px solid red')]")
                self._log_info(f"使用JavaScript找到元素 '{name}'", "green")
                
                # 检查元素位置
                rect = self.driver.execute_script('return arguments[0].getBoundingClientRect();', element)
                if rect['width'] <= 0 or rect['height'] <= 0:
                    self._log_info(f"警告: 使用JavaScript找到的元素'{name}'尺寸异常: width={rect['width']}, height={rect['height']}", "orange")
                    # 尝试滚动到元素
                    self.driver.execute_script("arguments[0].scrollIntoView({behavior: 'auto', block: 'center'});", element)
                    time.sleep(0.5)
                
                return element
        except Exception as e:
            self._log_info(f"JavaScript未找到元素 '{name}': {str(e)}", "orange")
        
        self._log_info(f"所有策略都未能找到元素 '{name}'", "red")
        return None
        
    def _generate_relative_xpath(self, absolute_xpath):
        """从绝对XPath生成更健壮的相对XPath"""
        try:
            # 提取最后几个节点，通常包含更多特定信息
            parts = absolute_xpath.split('/')
            if len(parts) <= 3:  # XPath太短，无法生成有意义的相对路径
                return None
                
            # 从最后一个节点开始构建
            last_part = parts[-1]
            
            # 尝试使用文本内容、ID、class等属性构建相对XPath
            relative_xpaths = []
            
            # 1. 如果最后一个节点包含文本，使用文本构建
            if '[text()' in last_part or 'contains(text()' in last_part:
                text_match = re.search(r'text\(\)\s*=\s*[\'"]([^\'"]+)[\'"]', last_part)
                contains_match = re.search(r'contains\(text\(\),\s*[\'"]([^\'"]+)[\'"]', last_part)
                
                if text_match:
                    text = text_match.group(1)
                    relative_xpaths.append(f"//*[text()='{text}']")
                elif contains_match:
                    text = contains_match.group(1)
                    relative_xpaths.append(f"//*[contains(text(),'{text}')]")
            
            # 2. 如果包含ID属性，使用ID构建
            id_match = re.search(r'@id\s*=\s*[\'"]([^\'"]+)[\'"]', last_part)
            if id_match:
                id_value = id_match.group(1)
                relative_xpaths.append(f"//*[@id='{id_value}']")
            
            # 3. 如果包含class属性，使用class构建
            class_match = re.search(r'@class\s*=\s*[\'"]([^\'"]+)[\'"]', last_part)
            contains_class = re.search(r'contains\(@class,\s*[\'"]([^\'"]+)[\'"]', last_part)
            
            if class_match:
                class_value = class_match.group(1)
                relative_xpaths.append(f"//*[@class='{class_value}']")
            elif contains_class:
                class_value = contains_class.group(1)
                relative_xpaths.append(f"//*[contains(@class,'{class_value}')]")
            
            # 4. 如果是特定标签，添加标签限制
            tag_match = re.match(r'([a-z]+)(\[|$)', last_part)
            if tag_match:
                tag = tag_match.group(1)
                # 为每个已生成的XPath添加标签限制
                tagged_xpaths = []
                for xpath in relative_xpaths:
                    tagged_xpath = xpath.replace('*', tag)
                    tagged_xpaths.append(tagged_xpath)
                relative_xpaths.extend(tagged_xpaths)
            
            # 5. 如果上述方法都没有产生有效的相对XPath，尝试使用最后两个节点
            if not relative_xpaths and len(parts) >= 2:
                last_two_parts = '/'.join(parts[-2:])
                relative_xpaths.append(f"//{last_two_parts}")
            
            return relative_xpaths[0] if relative_xpaths else None
            
        except Exception as e:
            self._log_info(f"生成相对XPath失败: {str(e)}", "red")
            return None
            
    def _xpath_to_css(self, xpath):
        """尝试将XPath转换为CSS选择器（简化版本）"""
        try:
            # 这只是一个简化的转换，不能处理所有XPath
            if not xpath or xpath.startswith('//'):
                return None
                
            # 移除绝对路径前缀
            if xpath.startswith('/html/'):
                xpath = xpath[6:]
                
            parts = xpath.split('/')
            css_parts = []
            
            for part in parts:
                if not part:
                    continue
                    
                # 处理ID
                if '@id' in part:
                    match = re.search(r'@id\s*=\s*[\'"]([^\'"]+)[\'"]', part)
                    if match:
                        css_parts.append(f"#{match.group(1)}")
                        continue
                
                # 处理class
                if '@class' in part:
                    match = re.search(r'@class\s*=\s*[\'"]([^\'"]+)[\'"]', part)
                    if match:
                        classes = match.group(1).split()
                        css_parts.append(f".{'.'.join(classes)}")
                        continue
                
                # 处理基本标签和索引
                tag_match = re.match(r'([a-z]+)(?:\[(\d+)\])?', part)
                if tag_match:
                    tag = tag_match.group(1)
                    index = tag_match.group(2)
                    if index:
                        # CSS选择器索引从1开始，XPath也是从1开始
                        css_parts.append(f"{tag}:nth-child({index})")
                    else:
                        css_parts.append(tag)
                        
            return ' > '.join(css_parts) if css_parts else None
            
        except Exception:
            return None
            
    def _find_by_text_content(self, text):
        """通过文本内容查找元素"""
        try:
            # 尝试多种文本匹配策略
            selectors = [
                f"//*[text()='{text}']",
                f"//*[contains(text(),'{text}')]",
                f"//*[contains(normalize-space(text()),'{text}')]",
                f"//span[contains(text(),'{text}')]",
                f"//div[contains(text(),'{text}')]",
                f"//a[contains(text(),'{text}')]",
                f"//button[contains(text(),'{text}')]",
                f"//label[contains(text(),'{text}')]",
                f"//*[contains(@title,'{text}')]",
                f"//*[contains(@aria-label,'{text}')]"
            ]
            
            for selector in selectors:
                try:
                    elements = self.driver.find_elements(By.XPATH, selector)
                    if elements:
                        # 优先选择文本长度接近的元素
                        elements.sort(key=lambda e: abs(len(e.text) - len(text)))
                        return elements[0]
                except:
                    continue
                    
            return None
            
        except Exception as e:
            self._log_info(f"通过文本内容查找元素失败: {str(e)}", "red")
            return None

    def _get_clipboard_content(self):
        """
        获取剪贴板内容，添加焦点保护
        """
        # 检查pyperclip依赖
        if pyperclip is None:
            self._log_info("pyperclip模块未正确导入，请安装: pip install pyperclip", "red")
            return ""
            
        try:
            # 尝试将焦点返回到采集工具窗口
            self._ensure_focus_for_clipboard()
            
            # 增加短暂延时，确保剪贴板内容已更新
            time.sleep(0.3)
            
            # 获取剪贴板内容
            clipboard_content = pyperclip.paste()
            content_length = len(clipboard_content) if clipboard_content else 0
            
            self._log_info(f"获取剪贴板内容，长度: {content_length}", "blue")
            return clipboard_content
        except Exception as e:
            self._log_info(f"获取剪贴板内容失败: {str(e)}", "red")
            return ""
    
    def _wait_for_clipboard_content(self, timeout=10.0, check_interval=0.2, min_length=10):
        """
        等待剪贴板内容更新，带超时机制，增加初始内容过滤
        
        参数:
        - timeout: 最大等待时间（秒）
        - check_interval: 检查间隔（秒）
        - min_length: 有效内容最小长度
        
        返回:
        - 更新的剪贴板内容或None
        """
        # 检查pyperclip依赖
        if pyperclip is None:
            self._log_info("pyperclip模块未正确导入，请安装: pip install pyperclip", "red")
            return None
            
        start_time = time.time()
        last_content = pyperclip.paste()
        initial_length = len(last_content) if last_content else 0
        
        # 获取初始剪贴板状态进行比较
        initial_content = getattr(self, 'initial_clipboard_content', "")
        
        self._log_info(f"开始等待剪贴板更新，初始内容长度: {initial_length}", "blue")
        
        # 清空剪贴板
        pyperclip.copy("")
        time.sleep(0.2)
        
        while time.time() - start_time < timeout:
            # 检查是否已终止操作
            if not self.is_running:
                self._log_info("操作已终止，停止等待剪贴板更新", "orange")
                return None
                
            # 确保焦点回到采集工具
            if (time.time() - start_time) > 2.0:  # 2秒后开始尝试恢复焦点
                self._manage_focus()
                
            current_content = pyperclip.paste()
            current_length = len(current_content) if current_content else 0
            
            # 检查内容是否等于初始剪贴板内容
            is_initial_content = current_content == initial_content
            if is_initial_content and current_length > 0:
                self._log_info(f"检测到剪贴板内容与脚本启动前相同，继续等待...", "orange")
                time.sleep(check_interval)
                continue
                
            # 检查内容长度是否超过1000字符，如果是则很可能是错误内容
            if current_length > 1000:
                self._log_info(f"检测到异常长度的剪贴板内容({current_length}字符)，继续等待...", "red")
                time.sleep(check_interval)
                continue
                
            # 内容发生变化且达到最小长度要求且不等于初始内容
            if current_content != last_content and current_length >= min_length and not is_initial_content:
                self._log_info(f"检测到有效剪贴板内容更新，长度: {current_length}", "green")
                return current_content
                
            last_content = current_content
            time.sleep(check_interval)
        
        self._log_info(f"等待剪贴板更新超时，最终内容长度: {len(last_content) if last_content else 0}", "orange")
        
        # 确保不返回初始内容
        if last_content == initial_content:
            self._log_info("最终内容与脚本启动前相同，返回空内容", "red")
            return None
            
        return last_content if len(last_content) >= min_length else None
    
    def _store_clipboard_content(self, element_name, content, order_id=None):
        """存储剪贴板内容，与订单ID建立映射关系，增加改进的内容有效性验证"""
        if not content or not isinstance(content, str) or not content.strip():
            self._log_info(f"剪贴板内容为空或无效，跳过存储", "orange")
            return False
            
        # 如果内容等于初始剪贴板内容，拒绝存储
        if hasattr(self, 'initial_clipboard_content') and content == self.initial_clipboard_content:
            self._log_info(f"检测到剪贴板内容与脚本启动前相同，拒绝存储", "red")
            # 立即触发人工输入
            self._manually_associate_clipboard_with_order_id(content)
            return False
        
        # 如果没有提供订单ID，尝试使用当前订单ID
        if not order_id:
            order_id = getattr(self, 'last_captured_order_id', None) or self.current_order_id
        
        # 必须有明确订单ID，禁止 None、空字符串、temp_id 写入
        if not order_id or not str(order_id).strip() or str(order_id).startswith("temp_order"):
            self._log_info("未提供有效订单ID，拒绝写入映射", "red")
            return False
            
        # 清理订单ID格式
        clean_order_id = order_id.replace("订单编号：", "") if isinstance(order_id, str) else str(order_id)
        
        # 验证内容有效性，但不修改内容
        is_valid, confidence, reason = self._is_valid_shipping_info(content)
        
        # 记录验证结果，便于后续分析
        if not hasattr(self, 'content_validation_results'):
            self.content_validation_results = {}
        
        self.content_validation_results[clean_order_id] = {
            'is_valid': is_valid,
            'confidence': confidence,
            'reason': reason,
            'timestamp': time.time()
        }
        
        # 检查是否已存在此订单ID的映射
        if clean_order_id in self.order_clipboard_contents:
            existing_content = self.order_clipboard_contents[clean_order_id]
            existing_valid, existing_confidence, existing_reason = self._is_valid_shipping_info(existing_content)
            
            # 决策逻辑：
            # 1. 如果新内容有效而现有内容无效，则替换
            # 2. 如果两者都有效，保留可信度更高的
            # 3. 如果可信度相同或非常接近(差距<10)，保留更长的有效内容
            # 4. 如果两者都无效，标记为需要人工检查
            
            if is_valid and not existing_valid:
                self._log_info(f"用有效内容({confidence}分)替换无效内容: {reason}", "green")
                self.order_clipboard_contents[clean_order_id] = content
                return True
            elif not is_valid and existing_valid:
                self._log_info(f"保留现有有效内容({existing_confidence}分)，忽略无效内容: {reason}", "orange")
                return False
            elif is_valid and existing_valid:
                # 两者都有效，比较可信度
                if abs(confidence - existing_confidence) < 10:
                    # 可信度接近，比较长度
                    if len(content) > len(existing_content):
                        self._log_info(f"可信度接近，保留更长的原始内容 ({len(content)}字符 vs {len(existing_content)}字符)", "green")
                        self.order_clipboard_contents[clean_order_id] = content
                        return True
                    else:
                        self._log_info(f"可信度接近，保留现有更长或相同长度的原始内容", "orange")
                        return False
                elif confidence > existing_confidence:
                    self._log_info(f"保留可信度更高的内容 ({confidence}分 vs {existing_confidence}分)", "green")
                    self.order_clipboard_contents[clean_order_id] = content
                    return True
                else:
                    self._log_info(f"保留现有可信度更高的内容 ({existing_confidence}分 vs {confidence}分)", "orange")
                    return False
            else:
                # 两者都无效，标记为需要人工检查
                # 仍然需要选择一个版本暂时保存
                if not hasattr(self, 'orders_need_review'):
                    self.orders_need_review = set()
                
                self.orders_need_review.add(clean_order_id)
                self._log_info(f"订单 {clean_order_id} 的收货信息可能无效，标记为需要人工检查", "red")
                
                # 暂时选择一个相对较好的版本，等待人工审核
                if abs(confidence - existing_confidence) < 10:
                    if len(content) > len(existing_content):
                        self._log_info(f"两者均需人工检查，暂时保留更长的原始内容", "orange")
                        self.order_clipboard_contents[clean_order_id] = content
                        return True
                    else:
                        self._log_info(f"两者均需人工检查，暂时保留现有更长的原始内容", "orange")
                        return False
                elif confidence > existing_confidence:
                    self._log_info(f"两者均需人工检查，暂时保留可信度较高的原始内容 ({confidence}分 vs {existing_confidence}分)", "orange")
                    self.order_clipboard_contents[clean_order_id] = content
                    return True
                else:
                    self._log_info(f"两者均需人工检查，暂时保留现有可信度较高的原始内容 ({existing_confidence}分 vs {confidence}分)", "orange")
                    return False
        else:
            # 新订单ID，直接存储原始内容
            self.order_clipboard_contents[clean_order_id] = content
            
            # 如果内容可能无效，标记为需要人工检查
            if not is_valid:
                if not hasattr(self, 'orders_need_review'):
                    self.orders_need_review = set()
                
                self.orders_need_review.add(clean_order_id)
                self._log_info(f"订单 {clean_order_id} 的收货信息可能无效，标记为需要人工检查", "red")
            
            status = "有效" if is_valid else "可能无效"
            self._log_info(f"新建映射: 订单ID {clean_order_id} <-> 收货信息 ({status}, {confidence}分)", "green" if is_valid else "orange")
            return True
    
    def _save_clipboard_mappings(self):
        """
        保存订单ID与剪贴板内容的映射到JSON文件
        """
        try:
            if not hasattr(self, 'order_clipboard_contents') or not self.order_clipboard_contents:
                self._log_info("没有剪贴板映射数据可保存", "orange")
                return
                
            # 准备要保存的数据
            save_data = {
                "updated_at": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                "order_clipboard_mappings": self.order_clipboard_contents
            }
            
            # 保存到JSON文件
            with open('clipboard_mappings.json', 'w', encoding='utf-8') as f:
                json.dump(save_data, f, ensure_ascii=False, indent=2)
            
            self._log_info(f"已保存 {len(self.order_clipboard_contents)} 个订单ID与剪贴板内容的映射", "green")
        except Exception as e:
            self._log_info(f"保存剪贴板映射失败: {str(e)}", "red")
    
    def _load_clipboard_mappings(self):
        """
        从JSON文件加载订单ID与剪贴板内容的映射
        """
        # 确保字典已初始化
        if not hasattr(self, 'order_clipboard_contents'):
            self.order_clipboard_contents = {}
        
        try:
            if os.path.exists('clipboard_mappings.json'):
                with open('clipboard_mappings.json', 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                if 'order_clipboard_mappings' in data:
                    # 将已加载的映射合并到现有字典
                    loaded_mappings = data['order_clipboard_mappings']
                    self.order_clipboard_contents.update(loaded_mappings)
                    self._log_info(f"已加载 {len(loaded_mappings)} 个订单ID与剪贴板内容的映射", "green")
                    return True
            
            self._log_info("未找到剪贴板映射文件或文件格式不正确", "orange")
            return False
        except Exception as e:
            self._log_info(f"加载剪贴板映射失败: {str(e)}", "red")
            return False

    def _clean_existing_clipboard_mappings(self):
        """清理现有的剪贴板映射，移除异常长度或无效的内容"""
        if not hasattr(self, 'order_clipboard_contents') or not self.order_clipboard_contents:
            return 0
            
        cleaned_count = 0
        for order_id, content in list(self.order_clipboard_contents.items()):
            # 检查内容长度
            if len(content) > 1000:
                self._log_info(f"清理异常长度的收货信息: 订单ID {order_id}, 长度 {len(content)}字符", "orange")
                del self.order_clipboard_contents[order_id]
                cleaned_count += 1
                continue
                
            # 验证内容有效性
            is_valid, confidence, reason = self._is_valid_shipping_info(content)
            if not is_valid:
                self._log_info(f"清理无效的收货信息: 订单ID {order_id}, 原因: {reason}", "orange")
                del self.order_clipboard_contents[order_id]
                cleaned_count += 1
                
        if cleaned_count > 0:
            self._log_info(f"已清理 {cleaned_count} 条异常收货信息映射", "green")
            # 保存清理后的映射
            self._save_clipboard_mappings()
            
        return cleaned_count
            
    def _manage_focus(self):
        """
        焦点管理器，确保在关键操作后焦点回到采集工具
        """
        try:
            # 获取主窗口句柄
            hwnd = ctypes.windll.user32.FindWindowW(None, "收货信息自动采集工具")
            if hwnd:
                # 尝试将焦点设置回采集工具窗口
                ctypes.windll.user32.SetForegroundWindow(hwnd)
                self._log_info("已将焦点返回到采集工具", "blue")
                return True
            else:
                self._log_info("未找到采集工具窗口，无法重置焦点", "orange")
                return False
        except Exception as e:
            self._log_info(f"焦点管理失败: {str(e)}", "red")
            return False
    
    def _ensure_focus_for_clipboard(self, max_attempts=3, delay=0.5):
        """
        确保在获取剪贴板内容前采集工具有焦点
        
        参数:
        - max_attempts: 最大尝试次数
        - delay: 每次尝试后的延时(秒)
        
        返回:
        - 是否成功将焦点设置到采集工具
        """
        for attempt in range(max_attempts):
            # 先尝试恢复焦点
            self._manage_focus()
            
            # 等待焦点切换完成
            time.sleep(delay)
            
            try:
                # 检查当前活动窗口是否为采集工具
                foreground_window = ctypes.windll.user32.GetForegroundWindow()
                window_text_length = ctypes.windll.user32.GetWindowTextLengthW(foreground_window)
                buffer = ctypes.create_unicode_buffer(window_text_length + 1)
                ctypes.windll.user32.GetWindowTextW(foreground_window, buffer, window_text_length + 1)
                
                if "收货信息自动采集工具" in buffer.value:
                    self._log_info(f"已确认焦点在采集工具上 (尝试 {attempt+1}/{max_attempts})", "green")
                    return True
                    
                self._log_info(f"焦点仍在其他窗口: {buffer.value} (尝试 {attempt+1}/{max_attempts})", "orange")
            except Exception as e:
                self._log_info(f"检查焦点失败: {str(e)} (尝试 {attempt+1}/{max_attempts})", "red")
        
        self._log_info("无法确保焦点回到采集工具，将尝试继续操作", "red")
        return False
    
    def _switch_focus_to_browser(self):
        """
        主动将焦点切换到浏览器窗口
        """
        if not self.driver:
            self._log_info("浏览器未连接，无法切换焦点", "orange")
            return False
        
        try:
            # 尝试执行一个简单的JavaScript来激活浏览器窗口
            self.driver.execute_script("window.focus()")
            
            # 获取浏览器窗口标题
            title = self.driver.title
            
            # 尝试通过标题查找窗口并激活
            if title:
                hwnd = ctypes.windll.user32.FindWindowW(None, title)
                if hwnd:
                    ctypes.windll.user32.SetForegroundWindow(hwnd)
                    self._log_info(f"已将焦点切换到浏览器: {title}", "blue")
                    return True
                
            self._log_info("无法通过标题切换焦点到浏览器", "orange")
            return False
        except Exception as e:
            self._log_info(f"切换焦点到浏览器失败: {str(e)}", "red")
            return False
            
    def _switch_focus_between_browser_and_tool(self, action):
        """
        在浏览器和采集工具之间切换焦点
        
        参数:
        - action: 要执行的操作，如"click"、"getText"等
        """
        if action in ["click", "clickAndGetClipboard"]:
            # 点击前确保浏览器有焦点
            self._switch_focus_to_browser()
            time.sleep(0.2)
        else:
            # 其他操作确保采集工具有焦点
            self._manage_focus()
            time.sleep(0.2)
    
    def _start_clipboard_monitor(self):
        """
        启动剪贴板内容监听线程
        """
        if self.clipboard_monitor_active:
            return
            
        self.clipboard_monitor_active = True
        self._log_info("启动剪贴板监听...", "blue")
        
        def monitor_clipboard():
            last_content = ""
            while self.clipboard_monitor_active:
                try:
                    current_content = pyperclip.paste()
                    if current_content != last_content and current_content.strip():
                        last_content = current_content
                        self.last_known_clipboard = current_content
                        
                        # 尝试提取当前订单ID
                        current_order_id = self.current_order_id or self._extract_current_order_id()
                        
                        # 记录日志但不打印剪贴板内容以避免日志过大
                        content_length = len(current_content)
                        self._log_info(f"[监听器] 检测到剪贴板变化，长度: {content_length}，当前订单ID: {current_order_id}", "blue")
                        
                        # 如果有订单ID，则建立映射
                        if current_order_id and content_length > 10:
                            self._store_clipboard_content("监听器捕获", current_content, current_order_id)
                except:
                    pass
                    
                # 减少CPU使用
                time.sleep(0.5)
        
        # 启动监听线程
        self.clipboard_monitor_thread = threading.Thread(target=monitor_clipboard, daemon=True)
        self.clipboard_monitor_thread.start()
    
    def _stop_clipboard_monitor(self):
        """
        停止剪贴板监听线程
        """
        self.clipboard_monitor_active = False
        if self.clipboard_monitor_thread:
            # 等待线程结束
            self.clipboard_monitor_thread.join(1.0)
            self.clipboard_monitor_thread = None
            self._log_info("停止剪贴板监听", "blue")
            
    def _manually_associate_clipboard_with_order_id(self, clipboard_content=None):
        """
        手动关联剪贴板内容与订单ID的对话框
        
        参数:
        - clipboard_content: 可选的剪贴板内容，如果不提供将自动获取
        
        返回:
        - 是否成功关联
        """
        if not clipboard_content:
            clipboard_content = self._get_clipboard_content()
        
        if not clipboard_content or not clipboard_content.strip():
            messagebox.showerror("错误", "剪贴板内容为空，无法关联")
            return False
            
        # 创建对话框
        dialog = tk.Toplevel(self.root)
        dialog.title("手动关联订单ID")
        dialog.geometry("500x400")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 设置为模态窗口
        dialog.focus_set()
        
        # 预览剪贴板内容
        ttk.Label(dialog, text="剪贴板内容预览:").pack(pady=5)
        preview_frame = ttk.Frame(dialog)
        preview_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        preview_text = ScrolledText(preview_frame, height=10)
        preview_text.pack(fill=tk.BOTH, expand=True)
        preview_text.insert(tk.END, clipboard_content[:500] + ("..." if len(clipboard_content) > 500 else ""))
        preview_text.config(state=tk.DISABLED)
        
        # 订单ID输入
        id_frame = ttk.Frame(dialog)
        id_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Label(id_frame, text="输入订单ID:").pack(side=tk.LEFT)
        order_id_var = tk.StringVar()
        order_id_entry = ttk.Entry(id_frame, textvariable=order_id_var, width=30)
        order_id_entry.pack(side=tk.LEFT, padx=5)
        
        # 最近使用的订单ID
        if self.last_order_ids:
            recent_frame = ttk.LabelFrame(dialog, text="最近使用的订单ID")
            recent_frame.pack(fill=tk.X, padx=10, pady=5)
            
            for order_id in reversed(self.last_order_ids):
                def set_id(oid=order_id):
                    order_id_var.set(oid)
                
                ttk.Button(
                    recent_frame,
                    text=order_id,
                    command=set_id
                ).pack(side=tk.LEFT, padx=5, pady=5)
        
        # 按钮区域
        btn_frame = ttk.Frame(dialog)
        btn_frame.pack(fill=tk.X, pady=10)
        
        result = {"confirmed": False, "order_id": ""}
        
        def on_confirm():
            input_id = order_id_var.get().strip()
            if not input_id:
                messagebox.showerror("错误", "请输入订单ID")
                return
                
            result["confirmed"] = True
            result["order_id"] = input_id
            dialog.destroy()
            
        def on_cancel():
            dialog.destroy()
        
        ttk.Button(btn_frame, text="确认关联", command=on_confirm).pack(side=tk.RIGHT, padx=10)
        ttk.Button(btn_frame, text="取消", command=on_cancel).pack(side=tk.RIGHT, padx=10)
        
        # 等待对话框关闭
        dialog.wait_window()
        
        # 处理结果
        if result["confirmed"] and result["order_id"]:
            self._store_clipboard_content("手动关联", clipboard_content, result["order_id"])
            
            # 将ID添加到最近使用列表
            if result["order_id"] not in self.last_order_ids:
                self.last_order_ids.append(result["order_id"])
                # 保留最近的5个ID
                if len(self.last_order_ids) > 5:
                    self.last_order_ids.pop(0)
            
            # 更新当前订单ID
            self.current_order_id = result["order_id"]
            
            self._log_info(f"已手动关联订单ID: {result['order_id']} 与剪贴板内容", "green")
            self._save_clipboard_mappings()
            return True
            
        return False

    def _is_valid_shipping_info(self, content):
        """
        验证内容是否为有效的收货信息，采用更灵活的规则
        返回 (is_valid, confidence_score, reason)
        """
        import re
        
        if not content or not isinstance(content, str):
            return False, 0, "内容为空或类型错误"
        
        content = content.strip()
        if len(content) < 5:  # 内容过短，肯定不是有效收货信息
            return False, 0, "内容过短"
            
        # 添加最大长度限制 - 正常收货信息不应超过1000字符
        if len(content) > 1000:
            return False, 0, f"内容过长 ({len(content)}字符)，超出正常收货信息范围"
        
        # 1. 首先检查是否包含明显的错误日志内容
        error_indicators = [
            r"\[\d{2}:\d{2}:\d{2}\].*错误",
            r"\[\d{2}:\d{2}:\d{2}\].*异常",
            r"\[\d{2}:\d{2}:\d{2}\].*Exception",
            r"\[\d{2}:\d{2}:\d{2}\].*Error",
            r"Traceback \(most recent call last\)",
            r"Exception in thread",
            r"cannot access local variable",
            r"ImportError:",
            r"AttributeError:",
            r"TypeError:",
            r"ValueError:"
        ]
        
        # 增加代码特征检测
        code_indicators = [
            r"def\s+\w+\(.*\):", 
            r"class\s+\w+\(.*\):",
            r"import\s+\w+",
            r"from\s+\w+\s+import",
            r"return\s+",
            r"if\s+.*:",
            r"else:",
            r"elif\s+.*:",
            r"for\s+.*\s+in\s+.*:",
            r"while\s+.*:",
            r"try:",
            r"except",
            r"finally:",
            r"```python",
            r"```"
        ]
        
        # 增加JSON格式检测 - 防止将元素配置数据误认为收货信息
        json_indicators = [
            r'"full_text"\s*:',
            r'"chinese_only"\s*:',
            r'"xpath"\s*:',
            r'"action"\s*:',
            r'"custom_name"\s*:',
            r'"element_type"\s*:',
            r'"captured_at"\s*:',
            r'\[\s*\{.*"full_text"',
            r'\{.*"xpath".*\}',
            r'/html/body/div\[\d+\]'
        ]
        
        for pattern in error_indicators:
            if re.search(pattern, content):
                return False, 0, f"内容包含错误日志特征: {pattern}"
        
        # 检查代码特征
        for pattern in code_indicators:
            if re.search(pattern, content):
                return False, 0, f"内容包含代码特征: {pattern}"
        
        # 检查JSON格式特征 - 防止将元素配置数据误认为收货信息
        for pattern in json_indicators:
            if re.search(pattern, content):
                return False, 0, f"内容包含JSON/元素配置特征: {pattern}"
        
        # 2. 检查是否包含手机号码 (这是收货信息的强特征)
        # 支持手机号和座机号两种格式
        phone_patterns = [
            r"1[3-9]\d{9}",  # 手机号
            r"\d{3,4}-\d{7,8}"  # 座机号 如 021-53395199
        ]
        
        has_phone = any(re.search(pattern, content) for pattern in phone_patterns)
        
        # 3. 检查是否包含地址特征 (省/市/区/县等)
        address_patterns = [
            r"(省|市|区|县|镇|乡|村|路|街|号楼|单元|室)",
            r"[东南西北中]门",
            r"[A-Za-z0-9#@\-]+号?库"  # 匹配仓库编号，如"2号库@DX-5E74D2M6D-F#"
        ]
        
        has_address = any(re.search(pattern, content) for pattern in address_patterns)
        
        # 4. 检查行数 - 有效的收货信息通常至少有2-3行
        lines = [line for line in content.split('\n') if line.strip()]
        has_multiple_lines = len(lines) >= 2
        
        # 5. 检查是否包含典型的收货信息格式特征
        # - 第一行通常是姓名（不做严格验证，因为可能是任何字符）
        # - 第二行通常是电话号码
        # - 后面几行是地址
        
        format_confidence = 0
        if has_multiple_lines:
            # 检查第二行是否包含电话号码
            if len(lines) >= 2 and any(re.search(pattern, lines[1]) for pattern in phone_patterns):
                format_confidence += 30
            
            # 检查第三行开始是否包含地址特征
            if len(lines) >= 3 and any(re.search(pattern, '\n'.join(lines[2:])) for pattern in address_patterns):
                format_confidence += 30
        
        # 计算总体可信度
        confidence = 0
        
        # 手机号是强特征
        if has_phone:
            confidence += 40
        
        # 地址特征也很重要
        if has_address:
            confidence += 30
        
        # 多行格式增加可信度
        if has_multiple_lines:
            confidence += 10
        
        # 典型格式特征
        confidence += format_confidence * 0.2  # 权重较低
        
        # 检查是否包含错误日志常见词汇（降低可信度）
        suspicious_terms = [
            "error", "exception", "failed", "undefined", "null", 
            "错误", "异常", "失败", "未定义", "空值"
        ]
        
        for term in suspicious_terms:
            if term.lower() in content.lower():
                confidence -= 15
                break
        
        # 最终判断
        is_valid = confidence >= 50  # 降低阈值，更宽松的判断
        
        reason = (
            f"电话:{has_phone}, 地址:{has_address}, "
            f"多行格式:{has_multiple_lines}, 格式特征:{format_confidence}, "
            f"可信度:{confidence}"
        )
        
        return is_valid, confidence, reason
    
    def _store_clipboard_content(self, element_name, content, order_id=None):
        if not content or not isinstance(content, str) or not content.strip():
            self._log_info(f"剪贴板内容为空，未存储", "orange")
            return False
        # 必须有明确订单ID，禁止 None、空字符串、temp_id 写入
        if not order_id or not str(order_id).strip() or str(order_id).startswith("temp_order"):
            self._log_info("未提供有效订单ID，拒绝写入映射", "red")
            return False
        clean_order_id = order_id.replace("订单编号：", "") if isinstance(order_id, str) else str(order_id)
        if clean_order_id in self.order_clipboard_contents:
            existing_content = self.order_clipboard_contents[clean_order_id]
            if len(existing_content) > len(content):
                self._log_info(f"保留现有更长的收货信息映射 (现有:{len(existing_content)}字符 vs 新内容:{len(content)}字符)", "orange")
                print(f"DEBUG-STORE-SKIP: 订单ID '{clean_order_id}' 已有更长的收货信息 ({len(existing_content)} > {len(content)})")
                return False
        self.order_clipboard_contents[clean_order_id] = content
        print(f"DEBUG-STORE-MAP: 订单ID '{clean_order_id}' -> 收货信息长度: {len(content)}")
        self._log_info(f"已建立映射: 订单ID {clean_order_id} <-> 收货信息", "green")
        return True

    def _batch_review_shipping_info(self, orders_to_review):
        """提供批量审核收货信息的界面"""
        from tkinter import Toplevel, Label, Button, Text, Scrollbar, Frame, StringVar, Entry
        import tkinter as tk
        
        if not orders_to_review:
            return
        
        # 创建审核窗口
        review_window = Toplevel(self.root)
        review_window.title("收货信息审核")
        review_window.geometry("800x600")
        review_window.transient(self.root)
        review_window.grab_set()
        
        # 创建说明标签
        Label(review_window, text="以下订单的收货信息可能存在问题，请逐一审核:", font=("Arial", 12)).pack(pady=10)
        
        # 创建滚动区域
        main_frame = Frame(review_window)
        main_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        canvas = tk.Canvas(main_frame)
        scrollbar = Scrollbar(main_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # 存储修改后的内容
        modified_contents = {}
        
        # 为每个需要审核的订单创建编辑区域
        for i, (order_id, content, confidence, reason, flag) in enumerate(orders_to_review):
            frame = Frame(scrollable_frame, bd=2, relief="groove")
            frame.pack(fill="x", expand=True, pady=5)
            
            # 订单信息
            Label(frame, text=f"订单 {i+1}/{len(orders_to_review)}: {order_id}", font=("Arial", 11, "bold")).pack(anchor="w")
            Label(frame, text=f"可信度: {confidence}分 | 原因: {reason}", fg="orange").pack(anchor="w")
            Label(frame, text=f"标记: {flag}", fg="red").pack(anchor="w")
            
            # 内容编辑区
            Label(frame, text="收货信息:").pack(anchor="w")
            text_widget = Text(frame, height=5, width=80)
            text_widget.pack(fill="x", padx=5, pady=5)
            text_widget.insert("1.0", content)
            
            # 保存修改内容的回调
            def save_content(order_id=order_id, text_widget=text_widget):
                modified_contents[order_id] = text_widget.get("1.0", "end-1c")
            
            # 操作按钮
            btn_frame = Frame(frame)
            btn_frame.pack(fill="x", pady=5)
            
            Button(btn_frame, text="确认无误", command=lambda oid=order_id: self._confirm_valid_content(oid, review_window)).pack(side="left", padx=5)
            Button(btn_frame, text="保存修改", command=lambda: save_content()).pack(side="left", padx=5)
            Button(btn_frame, text="删除此映射", command=lambda oid=order_id: self._delete_mapping(oid, review_window)).pack(side="left", padx=5)
        
        # 底部按钮
        bottom_frame = Frame(review_window)
        bottom_frame.pack(fill="x", pady=10)
        
        def on_finish():
            # 应用所有修改
            for order_id, new_content in modified_contents.items():
                if order_id in self.order_clipboard_contents and new_content.strip():
                    self.order_clipboard_contents[order_id] = new_content
                    self._log_info(f"已更新订单 {order_id} 的收货信息", "green")
            
            # 保存修改
            self._save_clipboard_mappings()
            
            # 清除需要审核的标记
            if hasattr(self, 'orders_need_review'):
                self.orders_need_review.clear()
            
            review_window.destroy()
        
        Button(bottom_frame, text="完成审核", command=on_finish, font=("Arial", 12)).pack(side="right", padx=10)
        Button(bottom_frame, text="取消", command=review_window.destroy, font=("Arial", 12)).pack(side="right", padx=10)
        
        # 等待窗口关闭
        self.root.wait_window(review_window)

    def _confirm_valid_content(self, order_id, parent_window=None):
        """确认订单内容有效，从需要审核的列表中移除"""
        if hasattr(self, 'orders_need_review') and order_id in self.orders_need_review:
            self.orders_need_review.remove(order_id)
            self._log_info(f"已确认订单 {order_id} 的收货信息有效", "green")
        
        # 更新界面反馈
        if parent_window:
            from tkinter import Label
            feedback = Label(parent_window, text=f"已确认订单 {order_id} 有效", fg="green")
            feedback.pack()
            parent_window.after(2000, feedback.destroy)

    def _delete_mapping(self, order_id, parent_window=None):
        """删除指定订单的映射"""
        if order_id in self.order_clipboard_contents:
            del self.order_clipboard_contents[order_id]
            self._log_info(f"已删除订单 {order_id} 的收货信息映射", "orange")
            
            # 从需要审核的列表中移除
            if hasattr(self, 'orders_need_review') and order_id in self.orders_need_review:
                self.orders_need_review.remove(order_id)
            
            # 保存修改
            self._save_clipboard_mappings()
        
        # 更新界面反馈
        if parent_window:
            from tkinter import Label
            feedback = Label(parent_window, text=f"已删除订单 {order_id} 的映射", fg="orange")
            feedback.pack()
            parent_window.after(2000, feedback.destroy)
            
    def _check_shipping_info_before_export(self):
        """检查并尝试修复收货信息字段，增加人工审核机制"""
        if not hasattr(self, 'collected_data') or not self.collected_data:
            self._log_info("没有收货信息可以导出", "red")
            from tkinter import messagebox
            messagebox.showwarning("导出失败", "没有收货信息可以导出，请先采集数据。")
            return False
            
        field_name = '复制完整收货信息'  # 收货信息字段名
        fixed_count = 0
        
        # 检查所有订单的收货信息
        orders_to_review = []
        
        # 确保order_clipboard_contents字典存在
        if not hasattr(self, 'order_clipboard_contents'):
            self.order_clipboard_contents = {}
            self._log_info("警告: 订单ID与收货信息映射字典不存在，已创建空字典", "red")
            print("DEBUG-EXPORT-ERROR: 订单ID与收货信息映射字典不存在，已创建空字典")
        else:
            # 导出前先确保所有映射都已保存到文件
            self._save_clipboard_mappings()
            self._log_info("已将当前所有映射保存到文件", "blue")
            
        # 输出当前订单ID与收货信息的映射关系
        self._log_info(f"导出前订单ID与收货信息映射关系：", "blue")
        print(f"DEBUG-EXPORT-MAP-COUNT: 映射字典中包含 {len(self.order_clipboard_contents)} 个订单")
        for order_id, content in self.order_clipboard_contents.items():
            content_preview = content[:30] + "..." if content else "空"
            print(f"DEBUG-EXPORT-MAP: 订单ID '{order_id}' -> 收货信息: '{content_preview}'")
            self._log_info(f"订单ID: {order_id}, 收货信息: {content_preview}", "blue")
        
        # 检查每个订单的收货信息是否唯一
        shipping_info_set = set()
        duplicate_count = 0
        
        for order_data in self.collected_data:
            # 获取当前订单ID
            order_id = order_data.get('订单编号', '')
            if isinstance(order_id, str) and order_id.startswith('订单编号：'):
                order_id = order_id.replace('订单编号：', '')
            
            self._log_info(f"处理订单: {order_id}", "blue")
            
            # 检查是否存在此字段
            if field_name in order_data:
                # 检查内容是否为布尔值或空
                if order_data[field_name] is True or order_data[field_name] is False or not order_data[field_name]:
                    self._log_info(f"检测到订单 {order_id} 的无效收货信息字段值: {order_data[field_name]}", "orange")
                    
                    # 优先使用订单专属的收货信息
                    if order_id and order_id in self.order_clipboard_contents:
                        order_data[field_name] = self.order_clipboard_contents[order_id]
                        fixed_count += 1
                        self._log_info(f"使用订单专属收货信息修复: '{self.order_clipboard_contents[order_id][:30]}...'", "green")
                        
                        # 检查是否是重复的收货信息
                        if order_data[field_name] in shipping_info_set:
                            duplicate_count += 1
                            self._log_info(f"警告: 订单 {order_id} 的收货信息与其他订单重复", "red")
                        else:
                            shipping_info_set.add(order_data[field_name])
                        
                        continue
                    
                    # 其次尝试使用全局变量
                    if hasattr(self, 'last_clipboard_content') and self.last_clipboard_content:
                        order_data[field_name] = self.last_clipboard_content
                        fixed_count += 1
                        self._log_info(f"使用全局变量修复了收货信息: '{self.last_clipboard_content[:30]}...'", "green")
                        continue
                    
                    # 最后尝试从剪贴板获取最新内容
                    import pyperclip
                    clipboard_content = pyperclip.paste()
                    if clipboard_content and clipboard_content.strip():
                        order_data[field_name] = clipboard_content
                        fixed_count += 1
                        self._log_info(f"使用剪贴板内容修复了收货信息: '{clipboard_content[:30]}...'", "green")
                        continue
                        
                    # 如果都失败了，记录一个明确的错误信息
                    order_data[field_name] = "【收货信息获取失败】"
                    self._log_info(f"无法修复订单 {order_id} 的收货信息字段，已标记为失败", "red")
            else:
                # 如果字段不存在，优先使用订单专属的收货信息
                if order_id and order_id in self.order_clipboard_contents:
                    order_data[field_name] = self.order_clipboard_contents[order_id]
                    fixed_count += 1
                    self._log_info(f"添加了订单 {order_id} 的专属收货信息: '{self.order_clipboard_contents[order_id][:30]}...'", "green")
                    continue
                
                # 其次尝试使用全局变量
                elif hasattr(self, 'last_clipboard_content') and self.last_clipboard_content:
                    order_data[field_name] = self.last_clipboard_content
                    fixed_count += 1
                    self._log_info(f"添加了缺失的收货信息字段: '{self.last_clipboard_content[:30]}...'", "green")
                else:
                    # 最后尝试从剪贴板获取
                    import pyperclip
                    clipboard_content = pyperclip.paste()
                    if clipboard_content and clipboard_content.strip():
                        order_data[field_name] = clipboard_content
                        fixed_count += 1
                        self._log_info(f"添加了缺失的收货信息字段: '{clipboard_content[:30]}...'", "green")
        
        if fixed_count > 0:
            self._log_info(f"导出前共修复了 {fixed_count} 条收货信息记录", "green")
        
        if duplicate_count > 0:
            self._log_info(f"警告: 检测到 {duplicate_count} 个订单的收货信息重复", "red")
            
        # 导出前检查收货信息字段内容
        self._log_info("导出前检查收货信息字段内容:", "blue")
        for i, order_data in enumerate(self.collected_data):
            order_id = order_data.get('订单编号', '')
            if isinstance(order_id, str) and order_id.startswith('订单编号：'):
                order_id = order_id.replace('订单编号：', '')
            else:
                order_id = str(order_id)
                
            shipping_info = order_data.get(field_name, "未设置")
            shipping_info_type = type(shipping_info).__name__
            shipping_info_preview = shipping_info[:30] + "..." if isinstance(shipping_info, str) and shipping_info != "未设置" else str(shipping_info)
            
            print(f"DEBUG-EXPORT-CHECK-{i+1}: 订单ID '{order_id}', {field_name} = {shipping_info_preview} (类型: {shipping_info_type})")
            self._log_info(f"记录 {i+1}: 订单ID={order_id}, {field_name} = {shipping_info_preview} (类型: {shipping_info_type})", "blue")
            
        return fixed_count > 0

    # ==================== 验证码检测相关方法 ====================
    
    def _show_captcha_manager(self):
        """显示验证码管理窗口"""
        manager_window = tk.Toplevel(self.root)
        manager_window.title("验证码检测管理")
        manager_window.geometry("600x500")
        manager_window.transient(self.root)
        manager_window.grab_set()
        
        # 模板管理区域
        template_frame = ttk.LabelFrame(manager_window, text="模板管理")
        template_frame.pack(fill="both", expand=True, padx=10, pady=5)
        
        # 模板列表
        ttk.Label(template_frame, text="当前模板:").pack(anchor="w", pady=(10, 5))
        self.template_listbox = tk.Listbox(template_frame, height=5)
        self.template_listbox.pack(fill="x", padx=10, pady=5)
        
        # 更新模板列表
        self._update_template_list()
        
        # 模板操作按钮
        template_btn_frame = ttk.Frame(template_frame)
        template_btn_frame.pack(fill="x", padx=10, pady=5)
        
        ttk.Button(template_btn_frame, text="添加模板", command=self._add_captcha_template).pack(side="left", padx=5)
        ttk.Button(template_btn_frame, text="清除所有模板", command=self._clear_captcha_templates).pack(side="left", padx=5)
        ttk.Button(template_btn_frame, text="自动加载模板", command=self._auto_load_captcha_templates).pack(side="left", padx=5)
        ttk.Button(template_btn_frame, text="截图保存为模板", command=self._capture_window_as_template).pack(side="left", padx=5)
        
        # 高级设置区域
        settings_frame = ttk.LabelFrame(manager_window, text="高级设置")
        settings_frame.pack(fill="x", padx=10, pady=5)
        
        # 检测间隔设置
        interval_frame = ttk.Frame(settings_frame)
        interval_frame.pack(fill="x", padx=10, pady=5)
        ttk.Label(interval_frame, text="检测间隔(秒):").pack(side="left")
        self.detection_interval_var = tk.DoubleVar(value=self.detection_interval)
        interval_spinbox = ttk.Spinbox(interval_frame, from_=0.1, to=5.0, increment=0.1, 
                                     textvariable=self.detection_interval_var, width=10)
        interval_spinbox.pack(side="left", padx=5)
        
        # 相似度阈值设置
        similarity_frame = ttk.Frame(settings_frame)
        similarity_frame.pack(fill="x", padx=10, pady=5)
        ttk.Label(similarity_frame, text="相似度阈值:").pack(side="left")
        self.similarity_threshold_var = tk.DoubleVar(value=self.similarity_threshold)
        similarity_scale = ttk.Scale(similarity_frame, from_=0.1, to=1.0, 
                                   variable=self.similarity_threshold_var, orient="horizontal")
        similarity_scale.pack(side="left", fill="x", expand=True, padx=5)
        self.similarity_label = ttk.Label(similarity_frame, text=f"{self.similarity_threshold:.2f}")
        self.similarity_label.pack(side="left")
        
        def update_similarity_label(*args):
            self.similarity_label.config(text=f"{self.similarity_threshold_var.get():.2f}")
        
        self.similarity_threshold_var.trace('w', update_similarity_label)
        
        # 遮罩层检测设置
        mask_frame = ttk.Frame(settings_frame)
        mask_frame.pack(fill="x", padx=10, pady=5)
        self.use_mask_detection_var = tk.BooleanVar(value=self.use_mask_detection)
        ttk.Checkbutton(mask_frame, text="启用遮罩层检测", variable=self.use_mask_detection_var).pack(side="left")
        
        # 控制按钮
        control_frame = ttk.Frame(manager_window)
        control_frame.pack(fill="x", padx=10, pady=10)
        
        def apply_settings():
            self.detection_interval = self.detection_interval_var.get()
            self.similarity_threshold = self.similarity_threshold_var.get()
            self.use_mask_detection = self.use_mask_detection_var.get()
            self._log_info("验证码检测设置已更新", "green")
            manager_window.destroy()
        
        ttk.Button(control_frame, text="应用设置", command=apply_settings).pack(side="right", padx=5)
        ttk.Button(control_frame, text="取消", command=manager_window.destroy).pack(side="right", padx=5)
    
    def _update_template_list(self):
        """更新模板列表显示"""
        if hasattr(self, 'template_listbox'):
            self.template_listbox.delete(0, tk.END)
            for i, template in enumerate(self.template_images):
                self.template_listbox.insert(tk.END, f"模板 {i+1} ({template['name']})")
    
    def _select_captcha_target_window(self):
        """选择验证码检测的目标窗口"""
        try:
            import win32gui
            
            def enum_windows_callback(hwnd, windows):
                if win32gui.IsWindowVisible(hwnd) and win32gui.GetWindowText(hwnd):
                    windows.append((hwnd, win32gui.GetWindowText(hwnd)))
                return True
            
            windows = []
            win32gui.EnumWindows(enum_windows_callback, windows)
            
            # 创建窗口选择对话框
            select_window = tk.Toplevel(self.root)
            select_window.title("选择目标窗口")
            select_window.geometry("400x300")
            select_window.transient(self.root)
            select_window.grab_set()
            
            ttk.Label(select_window, text="请选择要监控的窗口:").pack(pady=10)
            
            window_listbox = tk.Listbox(select_window)
            window_listbox.pack(fill="both", expand=True, padx=10, pady=5)
            
            for hwnd, title in windows:
                window_listbox.insert(tk.END, f"{title} (句柄: {hwnd})")
            
            def on_select():
                selection = window_listbox.curselection()
                if selection:
                    selected_window = windows[selection[0]]
                    self.target_window_handle = selected_window[0]
                    self.target_window_title = selected_window[1]
                    self._update_captcha_status_display()
                    self._log_info(f"已选择目标窗口: {self.target_window_title}", "green")
                    select_window.destroy()
            
            ttk.Button(select_window, text="确定", command=on_select).pack(pady=10)
            
        except Exception as e:
            self._log_info(f"选择目标窗口失败: {e}", "error")
    
    def _add_captcha_template(self):
        """添加验证码模板"""
        try:
            from tkinter import filedialog
            file_path = filedialog.askopenfilename(
                title="选择验证码模板图片",
                filetypes=[("图片文件", "*.jpg *.jpeg *.png *.bmp")]
            )
            
            if file_path:
                # 使用cv2.imdecode处理中文路径
                try:
                    with open(file_path, 'rb') as f:
                        image_data = f.read()
                    image_array = np.frombuffer(image_data, np.uint8)
                    template_image = cv2.imdecode(image_array, cv2.IMREAD_COLOR)
                except Exception as e:
                    template_image = None
                    self._log_info(f"读取文件失败: {e}", "error")
                
                if template_image is not None:
                    template_name = os.path.basename(file_path)
                    self.template_images.append({
                        'name': template_name,
                        'image': template_image,
                        'path': file_path
                    })
                    self._update_template_list()
                    self._update_captcha_status_display()
                    self._log_info(f"已添加模板: {template_name}", "green")
                else:
                    self._log_info("无法读取图片文件", "error")
        except Exception as e:
            self._log_info(f"添加模板失败: {e}", "error")
    
    def _clear_captcha_templates(self):
        """清除所有验证码模板"""
        self.template_images.clear()
        self._update_template_list()
        self._update_captcha_status_display()
        self._log_info("已清除所有模板", "orange")
    
    def _auto_load_captcha_templates(self):
        """自动加载验证码模板"""
        try:
            current_dir = os.path.dirname(os.path.abspath(__file__))
            templates_dir = os.path.join(current_dir, "captcha_templates")
            template_pattern = os.path.join(templates_dir, "模板*.jpg")
            template_files = glob.glob(template_pattern)
            
            loaded_count = 0
            for file_path in template_files:
                # 使用cv2.imdecode处理中文路径
                try:
                    with open(file_path, 'rb') as f:
                        image_data = f.read()
                    image_array = np.frombuffer(image_data, np.uint8)
                    template_image = cv2.imdecode(image_array, cv2.IMREAD_COLOR)
                except Exception as e:
                    template_image = None
                    self._log_info(f"读取模板文件失败 {os.path.basename(file_path)}: {e}", "error")
                
                if template_image is not None:
                    template_name = os.path.basename(file_path)
                    # 检查是否已存在同名模板
                    if not any(t['name'] == template_name for t in self.template_images):
                        self.template_images.append({
                            'name': template_name,
                            'image': template_image,
                            'path': file_path
                        })
                        loaded_count += 1
            
            self._update_template_list()
            self._update_captcha_status_display()
            
            if loaded_count > 0:
                self._log_info(f"自动加载了 {loaded_count} 个模板", "green")
            else:
                self._log_info("未找到新的模板文件", "orange")
                
        except Exception as e:
            self._log_info(f"自动加载模板失败: {e}", "error")
    
    def _capture_window_as_template(self):
        """截图目标窗口并保存为模板"""
        try:
            if not self.target_window_handle:
                self._log_info("请先选择目标窗口", "error")
                return
            
            # 截取目标窗口
            screenshot = self._capture_target_window()
            if screenshot is None:
                self._log_info("截取窗口失败", "error")
                return
            
            # 转换为PIL图像
            screenshot_pil = Image.fromarray(cv2.cvtColor(screenshot, cv2.COLOR_BGR2RGB))
            
            # 获取程序所在目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            templates_dir = os.path.join(current_dir, "captcha_templates")
            
            # 确保captcha_templates文件夹存在
            if not os.path.exists(templates_dir):
                os.makedirs(templates_dir)
            
            # 查找现有模板文件，确定下一个编号
            template_pattern = os.path.join(templates_dir, "模板*.jpg")
            existing_templates = glob.glob(template_pattern)
            
            # 确定新模板的编号
            if not existing_templates:
                template_name = "模板.jpg"
            else:
                # 提取现有模板的编号
                numbers = []
                for template_path in existing_templates:
                    filename = os.path.basename(template_path)
                    if filename == "模板.jpg":
                        numbers.append(1)
                    elif filename.startswith("模板") and filename.endswith(".jpg"):
                        try:
                            num_str = filename[2:-4]  # 去掉"模板"和".jpg"
                            if num_str.isdigit():
                                numbers.append(int(num_str))
                        except:
                            pass
                
                # 找到下一个可用编号
                if numbers:
                    next_num = max(numbers) + 1
                    if next_num == 2:
                        template_name = "模板2.jpg"
                    else:
                        template_name = f"模板{next_num}.jpg"
                else:
                    template_name = "模板.jpg"
            
            # 保存模板文件
            template_path = os.path.join(templates_dir, template_name)
            screenshot_pil.save(template_path, "JPEG", quality=95)
            
            # 加载到模板列表中
            template_image = cv2.cvtColor(screenshot, cv2.COLOR_BGR2RGB)
            template_image = cv2.cvtColor(template_image, cv2.COLOR_RGB2BGR)
            
            self.template_images.append({
                'name': template_name,
                'image': template_image,
                'path': template_path
            })
            
            self._update_template_list()
            self._update_captcha_status_display()
            self._log_info(f"已截图并保存为模板: {template_name}", "green")
            
        except Exception as e:
            self._log_info(f"截图保存为模板失败: {e}", "error")
    
    def _start_captcha_detection(self):
        """启动验证码检测"""
        if self.captcha_running:
            return
        
        if not self.template_images and not self.use_mask_detection:
            self._log_info("请先添加模板或启用遮罩层检测", "error")
            return
        
        if not self.target_window_handle:
            self._log_info("请先选择目标窗口", "error")
            return
        
        self.captcha_running = True
        self.captcha_detected = False
        self._update_captcha_status_display()
        
        # 启动检测线程
        self.captcha_thread = threading.Thread(target=self._captcha_detection_loop, daemon=True)
        self.captcha_thread.start()
        
        self._log_info("验证码检测已启动", "green")
    
    def _stop_captcha_detection(self):
        """停止验证码检测"""
        self.captcha_running = False
        self.captcha_detected = False
        self._update_captcha_status_display()
        self._log_info("验证码检测已停止", "orange")
    
    def _captcha_detection_loop(self):
        """验证码检测循环"""
        frame_count = 0  # 帧计数器，用于跳过初始几帧
        while self.captcha_running:
            try:
                # 截取目标窗口
                screenshot = self._capture_target_window()
                if screenshot is None:
                    time.sleep(self.detection_interval)
                    continue
                
                # 跳过初始几帧以避免误报
                if frame_count < 3:  # 跳过前3帧
                    frame_count += 1
                    time.sleep(self.detection_interval)
                    continue
                
                detected = False
                
                # 模板匹配检测
                if self.template_images:
                    detected = self._template_match_detection(screenshot)
                
                # 遮罩层检测
                if not detected and self.use_mask_detection:
                    detected = self._mask_layer_detection(screenshot)
                
                # 处理检测结果
                if detected and not self.captcha_detected:
                    self.captcha_detected = True
                    self.root.after(0, self._on_captcha_detected)
                elif not detected and self.captcha_detected:
                    self.captcha_detected = False
                    self.root.after(0, self._on_captcha_disappeared)
                
                # 更新状态显示
                self.root.after(0, self._update_captcha_status_display)
                
            except Exception as e:
                self._log_info(f"验证码检测出错: {e}", "error")
            
            time.sleep(self.detection_interval)
    
    def _capture_target_window(self):
        """截取目标窗口"""
        try:
            import win32gui, win32ui, win32con
            from PIL import Image
            
            # 检查窗口是否存在
            if not win32gui.IsWindow(self.target_window_handle):
                return None
            
            # 获取窗口客户区矩形（不包括标题栏和边框）
            client_rect = win32gui.GetClientRect(self.target_window_handle)
            client_width = client_rect[2]
            client_height = client_rect[3]
            
            if client_width <= 0 or client_height <= 0:
                return None
            
            # 获取窗口设备上下文
            hwnd_dc = win32gui.GetWindowDC(self.target_window_handle)
            mfc_dc = win32ui.CreateDCFromHandle(hwnd_dc)
            save_dc = mfc_dc.CreateCompatibleDC()
            
            # 创建位图对象
            save_bitmap = win32ui.CreateBitmap()
            save_bitmap.CreateCompatibleBitmap(mfc_dc, client_width, client_height)
            save_dc.SelectObject(save_bitmap)
            
            # 复制窗口内容到位图
            # 使用PrintWindow API，这对于某些应用程序更可靠
            import ctypes
            from ctypes import wintypes
            user32 = ctypes.windll.user32
            result = user32.PrintWindow(self.target_window_handle, save_dc.GetSafeHdc(), 3)  # PW_RENDERFULLCONTENT = 3
            
            if result == 0:
                # PrintWindow失败，尝试BitBlt方法
                result = save_dc.BitBlt((0, 0), (client_width, client_height), mfc_dc, (0, 0), win32con.SRCCOPY)
            
            if result == 0:
                # 两种方法都失败
                return None
            
            # 获取位图数据
            bmp_info = save_bitmap.GetInfo()
            bmp_str = save_bitmap.GetBitmapBits(True)
            
            # 转换为PIL图像
            im = Image.frombuffer(
                'RGB',
                (bmp_info['bmWidth'], bmp_info['bmHeight']),
                bmp_str, 'raw', 'BGRX', 0, 1
            )
            
            # 转换为OpenCV格式
            screenshot = cv2.cvtColor(np.array(im), cv2.COLOR_RGB2BGR)
            
            # 清理资源
            win32gui.DeleteObject(save_bitmap.GetHandle())
            save_dc.DeleteDC()
            mfc_dc.DeleteDC()
            win32gui.ReleaseDC(self.target_window_handle, hwnd_dc)
            
            return screenshot
            
        except Exception as e:
            self._log_info(f"截取窗口失败: {e}", "error")
            return None
    
    def _template_match_detection(self, screenshot):
        """模板匹配检测"""
        try:
            for template_info in self.template_images:
                template = template_info['image']
                result = cv2.matchTemplate(screenshot, template, cv2.TM_CCOEFF_NORMED)
                _, max_val, _, _ = cv2.minMaxLoc(result)
                
                if max_val >= self.similarity_threshold:
                    return True
            return False
        except Exception as e:
            self._log_info(f"模板匹配检测出错: {e}", "error")
            return False
    
    def _mask_layer_detection(self, screenshot):
        """遮罩层检测"""
        try:
            # 转换为灰度图
            gray = cv2.cvtColor(screenshot, cv2.COLOR_BGR2GRAY)
            
            # 计算图像的平均亮度
            mean_brightness = np.mean(gray)
            
            # 如果平均亮度过低，可能存在遮罩层
            if mean_brightness < 50:  # 可调阈值
                return True
            
            # 检测大面积的暗色区域
            _, binary = cv2.threshold(gray, 30, 255, cv2.THRESH_BINARY)
            dark_pixels = np.sum(binary == 0)
            total_pixels = gray.shape[0] * gray.shape[1]
            dark_ratio = dark_pixels / total_pixels
            
            # 如果暗色像素比例过高，可能存在遮罩层
            if dark_ratio > 0.7:  # 可调阈值
                return True
            
            return False
        except Exception as e:
            self._log_info(f"遮罩层检测出错: {e}", "error")
            return False
    
    def _update_captcha_status_display(self):
        """更新验证码状态显示"""
        try:
            # 更新验证码状态
            if self.captcha_detected:
                status_text = "检测到验证码"
                status_color = "red"
                indicator_color = "red"
            elif self.captcha_running:
                status_text = "检测中..."
                status_color = "green"
                indicator_color = "green"
            else:
                status_text = "未启动"
                status_color = "gray"
                indicator_color = "gray"
            
            if hasattr(self, 'captcha_status_label'):
                self.captcha_status_label.config(text=status_text, foreground=status_color)
            
            if hasattr(self, 'captcha_indicator'):
                self.captcha_indicator.delete("all")
                self.captcha_indicator.create_oval(2, 2, 18, 18, fill=indicator_color, outline="black")
            
            # 更新目标窗口状态
            if hasattr(self, 'target_window_label'):
                if self.target_window_handle:
                    window_text = f"目标窗口: {self.target_window_title[:20]}..."
                else:
                    window_text = "目标窗口: 未选择"
                self.target_window_label.config(text=window_text)
            
            # 更新模板数量
            if hasattr(self, 'template_count_label'):
                template_count = len(self.template_images)
                self.template_count_label.config(text=f"模板数量: {template_count}")
                
        except Exception as e:
            self._log_info(f"更新验证码状态显示出错: {e}", "error")
    
    def _on_captcha_detected(self):
        """检测到验证码时的处理"""
        try:
            # 停止当前操作
            if self.is_running and not self.is_paused:
                self.is_paused = True
                self.force_stop_flag = True
                self._log_info("检测到验证码，自动暂停操作", "red")
            
            # 播放警报音
            try:
                winsound.Beep(1000, 500)  # 1000Hz, 500ms
            except:
                pass
            
            # 显示警告框
            from tkinter import messagebox
            messagebox.showwarning("验证码检测", "检测到验证码！\n\n操作已自动暂停，请手动处理验证码后继续。")
            
            self._update_captcha_status_display()
            
        except Exception as e:
            self._log_info(f"处理验证码检测事件出错: {e}", "error")
    
    def _on_captcha_disappeared(self):
        """验证码消失时的处理"""
        try:
            # 清除强制停止标志
            if hasattr(self, 'force_stop_flag'):
                self.force_stop_flag = False
            
            self._log_info("验证码已消失", "green")
            self._update_captcha_status_display()
            
            # 如果之前因验证码暂停，现在可以准备重新执行当前订单
            if self.is_paused and hasattr(self, 'current_order_index'):
                self._log_info("验证码已处理，可以继续操作", "green")
            
        except Exception as e:
            self._log_info(f"处理验证码消失事件出错: {e}", "error")
    
    def collect_page_turn_element(self):
        """采集翻页元素"""
        try:
            if not self.driver:
                self._log_info("请先连接浏览器", "error")
                return
            
            # 设置翻页采集模式标志
            self.collecting_page_turn = True
            
            self._log_info("请将鼠标悬停在翻页按钮上，然后按\".\"键进行采集...", "blue")
            
            # 更新按钮状态
            if hasattr(self, 'collect_page_btn'):
                self.collect_page_btn.config(text="等待采集中...", state=tk.DISABLED)
                
        except Exception as e:
            self._log_info(f"启动翻页元素采集出错: {e}", "error")
            messagebox.showerror("错误", f"启动翻页元素采集出错: {e}")
    
    def _handle_page_turn_collection(self):
        """处理翻页元素采集 - 使用text.py中成功的悬停监听机制"""
        try:
            if not self.driver:
                self._log_info("浏览器未连接，无法采集", "red")
                messagebox.showerror("错误", "浏览器未连接，无法采集")
                # 恢复按钮状态
                if hasattr(self, 'collect_page_btn'):
                    self.collect_page_btn.config(text="采集翻页元素", state=tk.NORMAL)
                # 清除翻页采集模式标志
                self.collecting_page_turn = False
                return
            
            self._log_info("正在采集翻页按钮XPath...", "blue")
            
            # 使用悬停监听获取元素XPath
            self._log_info("正在通过悬停监听获取元素XPath...", "blue")
            xpath = self._get_hovered_xpath_recursive()
            
            if not xpath:
                self._log_info("✗ 未检测到悬停元素", "red")
                messagebox.showerror("采集失败", "未能获取到元素信息，请确保鼠标悬停在浏览器页面内的有效元素上。")
                # 恢复按钮状态
                if hasattr(self, 'collect_page_btn'):
                    self.collect_page_btn.config(text="采集翻页元素", state=tk.NORMAL)
                # 清除翻页采集模式标志
                self.collecting_page_turn = False
                return
                
            self._log_info("✓ 成功通过悬停监听方式定位到元素", "green")
            
            # 获取元素文本和元素对象
            element = None
            element_text = "未知元素"
            
            try:
                element = self.driver.find_element(By.XPATH, xpath)
                element_text = element.text or element.get_attribute('title') or element.get_attribute('aria-label') or element.tag_name
                self._log_info(f"成功获取元素: {element_text}")
                                    
            except Exception as e:
                self._log_info(f"无法通过XPath找到元素: {str(e)}", "orange")
                # 即使找不到元素，我们仍然保存XPath，因为它可能在点击时有效
                element_text = "SVG元素或动态元素"
            
            self._log_info(f"最终XPath: {xpath}")
            
            # 设置为翻页按钮XPath
            self.next_page_xpath = xpath
            self.next_page_collected = True
            
            # 高亮显示采集到的元素
            if element:
                try:
                    self._highlight_element_with_script(element)
                    self._log_info("元素高亮显示成功", "green")
                except Exception as e:
                    self._log_info(f"无法高亮显示元素: {str(e)}", "orange")
            else:
                self._log_info("跳过高亮显示：元素可能是SVG或动态元素", "orange")
            
            self._log_info(f"翻页元素采集成功: {xpath}", "green")
            
            # 恢复按钮状态
            if hasattr(self, 'collect_page_btn'):
                self.collect_page_btn.config(text="重新采集翻页元素", state=tk.NORMAL)
            
            messagebox.showinfo("成功", f"翻页元素采集成功！\n元素: {element_text}")
            
            # 清除翻页采集模式标志
            self.collecting_page_turn = False
                
        except Exception as e:
            self._log_info(f"采集翻页元素出错: {e}", "error")
            messagebox.showerror("错误", f"采集翻页元素出错: {e}")
            
            # 恢复按钮状态
            if hasattr(self, 'collect_page_btn'):
                self.collect_page_btn.config(text="采集翻页元素", state=tk.NORMAL)
            
            # 清除翻页采集模式标志
            self.collecting_page_turn = False
    
    def _highlight_element_with_script(self, element):
        """使用JavaScript高亮显示元素"""
        try:
            script = """
            var element = arguments[0];
            var originalStyle = element.getAttribute('style') || '';
            element.setAttribute('style', originalStyle + '; border: 3px solid red !important; background-color: yellow !important;');
            setTimeout(function() {
                element.setAttribute('style', originalStyle);
            }, 3000);
            """
            self.driver.execute_script(script, element)
        except Exception as e:
            self._log_info(f"高亮显示元素失败: {str(e)}", "orange")
    
    def _get_hovered_xpath(self, mouse_x, mouse_y):
        """获取鼠标悬停位置的元素XPath"""
        try:
            # 将屏幕坐标转换为浏览器内坐标
            browser_x, browser_y = self._screen_to_browser_coords(mouse_x, mouse_y)
            
            if browser_x is None or browser_y is None:
                return None
            
            # 执行JavaScript获取元素
            script = f"""
            var element = document.elementFromPoint({browser_x}, {browser_y});
            if (!element) return null;
            
            function getXPath(element) {{
                if (element.id !== '') {{
                    return '//*[@id="' + element.id + '"]';
                }}
                if (element === document.body) {{
                    return '/html/body';
                }}
                
                var ix = 0;
                var siblings = element.parentNode.childNodes;
                for (var i = 0; i < siblings.length; i++) {{
                    var sibling = siblings[i];
                    if (sibling === element) {{
                        return getXPath(element.parentNode) + '/' + element.tagName.toLowerCase() + '[' + (ix + 1) + ']';
                    }}
                    if (sibling.nodeType === 1 && sibling.tagName === element.tagName) {{
                        ix++;
                    }}
                }}
            }}
            
            return getXPath(element);
            """
            
            xpath = self.driver.execute_script(script)
            return xpath
            
        except Exception as e:
            self._log_info(f"获取元素XPath失败: {e}", "error")
            return None
    
    def _screen_to_browser_coords(self, screen_x, screen_y):
        """将屏幕坐标转换为浏览器内坐标"""
        try:
            if not self.target_window_handle:
                return None, None
            
            # 获取窗口位置
            rect = win32gui.GetWindowRect(self.target_window_handle)
            window_x, window_y, window_right, window_bottom = rect
            
            # 计算浏览器内坐标
            browser_x = screen_x - window_x
            browser_y = screen_y - window_y
            
            # 考虑浏览器标题栏和工具栏的偏移
            browser_y -= 100  # 大概的标题栏和工具栏高度
            
            return browser_x, browser_y
            
        except Exception as e:
            self._log_info(f"坐标转换失败: {e}", "error")
            return None, None
    
    def _highlight_element(self, xpath):
        """高亮显示元素"""
        try:
            script = f"""
            var element = document.evaluate('{xpath}', document, null, XPathResult.FIRST_ORDERED_NODE_TYPE, null).singleNodeValue;
            if (element) {{
                element.style.border = '3px solid red';
                element.style.backgroundColor = 'yellow';
                element.style.opacity = '0.8';
                
                setTimeout(function() {{
                    element.style.border = '';
                    element.style.backgroundColor = '';
                    element.style.opacity = '';
                }}, 3000);
            }}
            """
            
            self.driver.execute_script(script)
            
        except Exception as e:
            self._log_info(f"高亮元素失败: {e}", "error")
    
    def check_page_turn_needed(self):
        """检查是否需要翻页"""
        try:
            # 获取当前翻页页数设置
            if hasattr(self, 'page_count_var'):
                self.target_page_count = int(self.page_count_var.get())
            
            # 检查是否达到翻页条件
            if hasattr(self, 'current_order_index') and self.current_order_index > 0:
                if (self.current_order_index + 1) % self.target_page_count == 0:
                    return True
            
            return False
            
        except Exception as e:
            self._log_info(f"检查翻页条件失败: {e}", "error")
            return False
    
    def execute_page_turn(self):
        """执行翻页操作"""
        try:
            if not self.next_page_collected or not self.next_page_xpath:
                self._log_info("请先采集翻页元素", "error")
                return False
            
            self._log_info("开始执行翻页操作...", "blue")
            
            # 翻页前截图
            before_screenshot = self._save_page_screenshot("before")
            if not before_screenshot:
                self._log_info("翻页前截图失败", "error")
                return False
            
            # 检查翻页按钮是否存在
            try:
                element = self.driver.find_element(By.XPATH, self.next_page_xpath)
                if not element.is_enabled():
                    self._log_info("翻页按钮不可用", "error")
                    return False
            except Exception as e:
                self._log_info(f"找不到翻页按钮: {e}", "error")
                return False
            
            # 点击翻页按钮
            try:
                element.click()
                self._log_info("翻页按钮点击成功", "green")
            except Exception as e:
                self._log_info(f"点击翻页按钮失败: {e}", "error")
                return False
            
            # 等待页面加载
            time.sleep(3)
            
            # 翻页后截图
            after_screenshot = self._save_page_screenshot("after")
            if not after_screenshot:
                self._log_info("翻页后截图失败", "error")
                return False
            
            # 比较截图验证翻页是否成功
            if self._compare_screenshots(before_screenshot, after_screenshot):
                self.page_turn_count += 1
                self._log_info(f"翻页成功！当前已翻页 {self.page_turn_count} 次", "green")
                
                # 清理截图文件
                self._clear_screenshots()
                
                return True
            else:
                self._log_info("翻页失败：页面内容未发生变化", "error")
                return False
                
        except Exception as e:
            self._log_info(f"翻页操作失败: {e}", "error")
            return False
    
    def _save_page_screenshot(self, prefix):
        """保存页面截图"""
        try:
            if not self.target_window_handle:
                return None
            
            screenshot = self._capture_window_screenshot()
            if screenshot is None:
                return None
            
            timestamp = int(time.time() * 1000)
            filename = f"{prefix}_{timestamp}.png"
            filepath = os.path.join(self.screenshot_dir, filename)
            
            cv2.imwrite(filepath, screenshot)
            return filepath
            
        except Exception as e:
            self._log_info(f"保存截图失败: {e}", "error")
            return None
    
    def _compare_screenshots(self, before_path, after_path):
        """比较两张截图是否不同"""
        try:
            # 读取图片
            before_img = cv2.imread(before_path)
            after_img = cv2.imread(after_path)
            
            if before_img is None or after_img is None:
                return False
            
            # 确保图片尺寸相同
            if before_img.shape != after_img.shape:
                return True  # 尺寸不同说明页面发生了变化
            
            # 计算图片差异
            diff = cv2.absdiff(before_img, after_img)
            gray_diff = cv2.cvtColor(diff, cv2.COLOR_BGR2GRAY)
            
            # 计算非零像素的比例
            non_zero_count = cv2.countNonZero(gray_diff)
            total_pixels = gray_diff.shape[0] * gray_diff.shape[1]
            diff_ratio = non_zero_count / total_pixels
            
            # 如果差异比例大于阈值，认为翻页成功
            threshold = 0.1  # 10%的像素发生变化
            return diff_ratio > threshold
            
        except Exception as e:
            self._log_info(f"比较截图失败: {e}", "error")
            return False
    
    def _clear_screenshots(self):
        """清理截图文件"""
        try:
            for filename in os.listdir(self.screenshot_dir):
                if filename.endswith('.png'):
                    filepath = os.path.join(self.screenshot_dir, filename)
                    os.remove(filepath)
        except Exception as e:
            self._log_info(f"清理截图文件失败: {e}", "error")
    
    def on_page_count_changed(self, *args):
        """翻页页数改变时的回调"""
        try:
            if hasattr(self, 'page_count_var'):
                self.target_page_count = int(self.page_count_var.get())
                self._log_info(f"翻页页数设置为: {self.target_page_count}", "blue")
        except Exception as e:
            self._log_info(f"设置翻页页数失败: {e}", "error")

def main():
    """主函数"""
    # 检查依赖是否完整
    if not dependencies_ok:
        print("\n程序无法启动，请先安装缺失的依赖包。")
        input("按回车键退出...")
        return
    
    # 设置高DPI感知（Windows系统）
    try:
        ctypes.windll.shcore.SetProcessDpiAwareness(1)
    except Exception as e:
        logging.warning(f"无法设置DPI感知: {str(e)}")
    
    # 创建Tk根窗口
    try:
        root = tk.Tk()
    except Exception as e:
        logging.error(f"Tk初始化失败: {str(e)}")
        print(f"GUI初始化失败: {str(e)}")
        print("请确认已正确安装Python和Tkinter，或尝试重新安装Python。")
        print(f"Python版本: {sys.version}")
        print(f"TCL_LIBRARY环境变量: {os.environ.get('TCL_LIBRARY', '未设置')}")
        print(f"TK_LIBRARY环境变量: {os.environ.get('TK_LIBRARY', '未设置')}")
        input("按回车键退出...")
        return  # 退出程序
    
    # 创建应用实例
    app = ShippingInfoCollector(root)
    
    # 启动应用主循环
    root.mainloop()

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        # 捕获并记录未处理的异常
        error_msg = f"程序发生未处理的异常: {str(e)}\n{traceback.format_exc()}"
        logging.error(error_msg)
        try:
            messagebox.showerror("错误", f"程序发生错误:\n{str(e)}\n\n详细信息已记录到日志文件中。")
        except Exception:
            # 如果messagebox也失败，至少确保错误被记录
            print(f"程序发生严重错误: {str(e)}\n详细信息已记录到日志文件中。")